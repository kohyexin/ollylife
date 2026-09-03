from copy import deepcopy
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.16.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.17.docx"


REPLACEMENTS = {
    "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 0.16": "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 0.17",
    "Purpose  Define who must build, expose, consume, operate and test each component of the programme journey, where an OlyLife member starts in VCCHUB. This 0.16 edition adds VCCHUB email-OTP verification before OlyLife membership validation, OlyLife Full Name carry-over, and OlyLife-configured programme fees charged atomically to the member Wallet. The preceding document revision remains retained. Final production configuration and contractual approvals remain subject to joint sign-off.": "Purpose  Define who must build, expose, consume, operate and test each component of the programme journey, where an OlyLife member starts in VCCHUB. This 0.17 edition adds mandatory initial card funding at creation: the member chooses an opening card balance of at least 20.00 in the Wallet/card currency, and VCCHUB atomically deducts that amount together with any configured creation fee. The preceding document revision remains retained. Final production configuration and contractual approvals remain subject to joint sign-off.",
    "Card creation with one active virtual and one active physical card, OlyLife-configured Wallet fee charging, wallet-to-card top-up, authenticated cancellation, transaction history and Issuing capability mapping.": "Card creation with one active virtual and one active physical card, mandatory initial card funding, OlyLife-configured Wallet fee charging, wallet-to-card top-up, authenticated cancellation, transaction history and Issuing capability mapping.",
    "Create the cardholder and zero-balance wallet after Sumsub approval; show member/cardholder/address data read-only during card creation; enforce one active virtual plus one active physical card; calculate and display any OlyLife-configured card-creation fee; atomically debit that fee from Wallet only when issuance succeeds; show every card separately; use the configured OlyLife office address for physical fulfilment; and provide authenticated cancellation.": "Create the cardholder and zero-balance wallet after Sumsub approval; show member/cardholder/address data read-only during card creation; enforce one active virtual plus one active physical card; require the member to choose an initial card balance of at least 20.00 in the Wallet/card currency; calculate and display that funding amount, any OlyLife-configured creation fee and the resulting Wallet balance; then atomically create the card, debit fee plus funding from Wallet and credit the funding amount to the card. Show every card separately, use the configured OlyLife office address for physical fulfilment and provide authenticated cancellation.",
    "VCCHUB wallet and Star SaaS Issuing: reuse existing wallet, cardholder and card capabilities behind VCCHUB internal services. Extend fee accounting because the current platform posts fees to Merchant balance: programme fees must instead debit the member Wallet, produce a fee-ledger entry and reverse atomically when the triggering operation fails. Carry the registered address into the cardholder/card flow, but do not expose editable address fields on card creation.": "VCCHUB Wallet and Star SaaS Issuing: reuse existing Wallet, cardholder and card capabilities behind VCCHUB internal services. Extend fee accounting because the current platform posts fees to Merchant balance: programme fees must instead debit the member Wallet. Extend card creation to post a separate Wallet-to-card initial-funding entry, enforce the 20.00 minimum and commit card issuance, fee debit and funding transfer atomically. Carry the registered address into the cardholder/card flow, but do not expose editable address fields on card creation.",
    "Demo limitation  The demonstration uses Sumsub Sandbox, browser-held demo state, a mock email OTP (123456), a mock 2FA code (123456) for login and card cancellation, simulated OlyLife approval, a SGD 10.00 sample card-creation fee and a configurable placeholder office address. It illustrates the journey but is not a production email-delivery, identity, ledger, approval, fee-engine or fulfilment implementation.": "Demo limitation  The demonstration uses Sumsub Sandbox, browser-held demo state, a mock email OTP (123456), a mock 2FA code (123456) for login and card cancellation, simulated OlyLife approval, a SGD 10.00 sample card-creation fee, a SGD 20.00 minimum initial card balance and a configurable placeholder office address. It illustrates the journey but is not a production email-delivery, identity, ledger, approval, fee-engine or fulfilment implementation.",
    "Card creation shows member/cardholder and registered-address information read-only. The programme permits one active virtual card and one active physical card; both appear as separate rows with independent action menus. Before confirmation, VCCHUB shows the applicable OlyLife-configured creation fee and resulting Wallet balance. Successful creation posts the fee and card atomically. Physical creation collects no recipient and uses the configured OlyLife office address.": "Card creation shows member/cardholder and registered-address information read-only. The programme permits one active virtual card and one active physical card; both appear as separate rows with independent action menus. The member enters an initial card balance of at least 20.00 in the Wallet/card currency. Before confirmation, VCCHUB shows the applicable OlyLife-configured creation fee, initial funding amount, total Wallet deduction and resulting Wallet balance. Successful creation atomically creates the card, posts the fee, debits the funding from Wallet and credits it to the card. Physical creation collects no recipient and uses the configured OlyLife office address.",
    "Confirmed programme decisions  PhotonPay remains the issuer. The programme starts in VCCHUB, verifies email control, validates OlyLife membership before registration, collects the member's full address, creates a zero-balance Wallet with no card after KYC, and funds Wallets only through OlyLife's approved top-up process. Each Wallet is assigned a policy allowing one active virtual and one active physical card, with one default BIN per type and no user BIN selection. Policy changes apply prospectively unless a separate existing-wallet migration is approved. Physical cards are sent to the configured OlyLife office. Card cancellation releases the cancelled type slot. Issuer/card-transaction fees remain governed by PhotonPay. Additional programme fees requested by OlyLife are configured in VCCHUB and deducted from Wallet when their trigger succeeds; this requires a VCCHUB extension because current fee postings use Merchant balance.": "Confirmed programme decisions  PhotonPay remains the issuer. The programme starts in VCCHUB, verifies email control, validates OlyLife membership before registration, collects the member's full address, creates a zero-balance Wallet with no card after KYC, and funds Wallets only through OlyLife's approved top-up process. Each new card must receive an opening balance of at least 20.00 in the Wallet/card currency; VCCHUB deducts that funding together with any creation fee and credits the funding amount to the new card in one atomic operation. Each Wallet is assigned a policy allowing one active virtual and one active physical card, with one default BIN per type and no user BIN selection. Policy changes apply prospectively unless a separate existing-wallet migration is approved. Physical cards are sent to the configured OlyLife office. Card cancellation releases the cancelled type slot. Issuer/card-transaction fees remain governed by PhotonPay. Additional programme fees requested by OlyLife are configured in VCCHUB and deducted from Wallet when their trigger succeeds; this requires a VCCHUB extension because current fee postings use Merchant balance.",
    "Definition of ready  Programme development is ready when email OTP, OLY-01 member validation, account mapping, OlyLife approval/debit, VCCHUB Wallet credit/status, Wallet fee configuration/posting/reversal, Sumsub access, international-name mapping and PhotonPay compatibility evidence, office-address configuration, authenticated card cancellation and acceptance owners are documented and testable.": "Definition of ready  Programme development is ready when email OTP, OLY-01 member validation, account mapping, OlyLife approval/debit, VCCHUB Wallet credit/status, minimum initial card funding, Wallet fee configuration/posting/reversal, atomic card-creation ledger treatment, Sumsub access, international-name mapping and PhotonPay compatibility evidence, office-address configuration, authenticated card cancellation and acceptance owners are documented and testable.",
    "Handoff checkpoint  The first programme workshop should close email-OTP controls, member-status eligibility, the OLY-01 Full Name contract, OlyLife-to-Sumsub name precedence, account-ready mapping, Wallet-credit idempotency/reversal, the programme-fee catalogue and approval process, Wallet fee posting/reversal and Merchant-balance exclusion, the real OlyLife office address, physical fulfilment controls, default BINs, prospective card-policy versioning and cancellation authentication/audit details. Each NEW or EXTEND item can then become an owned backlog item and contract test.": "Handoff checkpoint  The first programme workshop should close email-OTP controls, member-status eligibility, the OLY-01 Full Name contract, OlyLife-to-Sumsub name precedence, account-ready mapping, Wallet-credit idempotency/reversal, the 20.00 initial-card-balance minimum, card-creation fee and funding ledger order, atomic failure/reversal treatment, Merchant-balance exclusion, the real OlyLife office address, physical fulfilment controls, default BINs, prospective card-policy versioning and cancellation authentication/audit details. Each NEW or EXTEND item can then become an owned backlog item and contract test.",
}


TABLE_REPLACEMENTS = {
    "Draft - email OTP and Wallet fee extension": "Draft - minimum initial card funding",
    "0.16": "0.17",
    "Show cardholder and registered address read-only. Display the OlyLife-configured card-creation fee and resulting Wallet balance; on success, atomically create the selected card and debit the fee from Wallet. Permit one active virtual plus one active physical card; physical uses the configured OlyLife office address.": "Show cardholder and registered address read-only. Require an initial card balance of at least 20.00 in the Wallet/card currency. Display the initial funding, OlyLife-configured creation fee, total deduction and resulting Wallet balance; atomically create the card, debit fee plus funding from Wallet and credit funding to the card. Permit one active virtual plus one active physical card; physical uses the configured OlyLife office address.",
    "Require approved KYC, an active cardholder and a Wallet balance sufficient for the configured fee; allow one active virtual and one active physical card under the policy version assigned at wallet creation": "Require approved KYC, an active cardholder, an initial card balance of at least 20.00 and a Wallet balance sufficient for initial funding plus the configured fee; allow one active virtual and one active physical card under the policy version assigned at wallet creation",
    "Reuse issuance; show cardholder/address read-only; preview and atomically debit the configured Wallet fee on success": "Reuse issuance; show cardholder/address read-only; preview initial funding, fee, total deduction and resulting balances; atomically create, charge and fund the card",
    "OlyLife owns each programme-fee decision and approves fee code, trigger, amount/rate, currency, tax, effective dates and refundability. VCCHUB implements a controlled fee configuration and extends its fee engine/ledger to debit Wallet rather than Merchant balance. The UI shows the fee and resulting Wallet balance before confirmation. The trigger action and fee posting commit atomically; failure produces no retained fee or an exactly-once reversal. Example: a USD 10.00 card-creation fee deducts USD 10.00 from a USD Wallet when the card is successfully created.": "OlyLife owns each programme-fee decision and approves fee code, trigger, amount/rate, currency, tax, effective dates and refundability. VCCHUB implements controlled fee configuration and extends its fee engine/ledger to debit Wallet rather than Merchant balance. For card creation, the UI shows the fee, chosen initial funding, total deduction and resulting balances. Example: with a USD 500.00 Wallet, USD 10.00 creation fee and USD 30.00 initial card balance, successful creation leaves USD 460.00 in Wallet and USD 30.00 on the card. The action commits atomically; failure retains neither fee nor funding debit.",
    "Issue only when no active virtual card exists and Wallet covers the configured creation fee; registered data remains read-only; post fee and issuance atomically.": "Issue only when no active virtual card exists, initial funding is at least 20.00 and Wallet covers initial funding plus the configured fee; registered data remains read-only; atomically issue, charge and fund the card.",
    "Issue only when no active physical card exists and Wallet covers the configured creation fee; use the OlyLife office address; post fee and issuance atomically; collect no member recipient.": "Issue only when no active physical card exists, initial funding is at least 20.00 and Wallet covers initial funding plus the configured fee; use the OlyLife office address; atomically issue, charge and fund the card; collect no member recipient.",
    "Reuse internally after confirming Wallet covers the configured creation fee. Permit one active card of each type, auto-use the default BIN, and atomically post the Wallet fee with successful issuance. List cards separately by card ID. Physical fulfilment uses the OlyLife office address. Cancellation releases only the cancelled type slot.": "Reuse internally after confirming the selected initial card balance is at least 20.00 and Wallet covers that funding plus the configured creation fee. Permit one active card of each type, auto-use the default BIN, and atomically create the card, post the Wallet fee, debit initial funding from Wallet and credit it to the card. Return separate fee/funding transaction IDs and final balances. List cards separately by card ID. Physical fulfilment uses the OlyLife office address. Cancellation releases only the cancelled type slot.",
    "card_id, card_type, card_status, last4, currency, card_balance, delivery_destination": "card_id, card_type, card_status, last4, currency, card_balance, delivery_destination, initial_funding_transaction_id, initial_card_balance, minimum_initial_balance, creation_fee, total_wallet_deduction, wallet_balance_after",
    "Funding and fee gate; OlyLife-controlled fee schedule; Wallet rather than Merchant fee posting; fee preview, atomic debit, idempotency and reversal; one active virtual plus one active physical card; default BIN per type; PhotonPay charset/name compatibility; separate card actions; read-only identity/address; OlyLife-office delivery; Wallet-to-card top-up; authenticated cancellation with Wallet refund/type-slot release.": "Initial-funding and fee gate; minimum 20.00 opening card balance; Wallet coverage of funding plus fee; separate Wallet-fee and Wallet-to-card ledger entries; preview, atomic debit/credit, idempotency and reversal; one active virtual plus one active physical card; default BIN per type; PhotonPay charset/name compatibility; separate card actions; read-only identity/address; OlyLife-office delivery; later Wallet-to-card top-up; authenticated cancellation with Wallet refund/type-slot release.",
    "Configure a SGD 10.00 card-creation fee (and representative zero/percentage/future-effective fees), then create cards with sufficient and insufficient Wallet balances. Confirm the equivalent USD 10.00 case in a USD Wallet.": "Configure a SGD 10.00 card-creation fee and SGD 20.00 minimum initial card balance. Test initial funding at 0.00, 19.99, 20.00 and 30.00 with sufficient and insufficient Wallet balances. Confirm the equivalent USD 10.00 fee plus USD 20.00 minimum in a USD Wallet.",
    "The applicable fee and post-action balance are shown before confirmation. On success, VCCHUB debits the fee exactly once from Wallet—not Merchant balance—and records fee code/version, trigger, request/card IDs and pre/post balances. Insufficient balance blocks the action with no posting. PhotonPay issuer/card-transaction fees remain separate.": "The fee, selected initial funding, total Wallet deduction and resulting Wallet/card balances are shown before confirmation. Values below 20.00 are rejected. On success, VCCHUB debits the fee and funding exactly once from Wallet, credits only funding to the card and records separate transaction IDs, fee code/version, trigger, request/card IDs and pre/post balances. Merchant balance is unchanged. Insufficient balance blocks the action with no posting.",
    "Attempt virtual and physical card creation with zero balance, with less than the configured fee, and with sufficient Wallet balance; then repeat after an approved Wallet top-up.": "Attempt virtual and physical card creation with initial card balances of 0.00, 19.99, 20.00 and 30.00. For valid amounts, test Wallet balances below, equal to and above initial funding plus the configured creation fee; repeat after an approved Wallet top-up.",
    "Creation is blocked unless Wallet covers the full configured fee. Funded creation proceeds only after authoritative Wallet checks; successful issuance deducts the fee exactly once and the card starts at zero card balance.": "Creation rejects an initial card balance below 20.00. It is also blocked unless Wallet covers initial funding plus the full configured fee. Successful issuance deducts both exactly once, credits only the initial funding to the card and returns final balances. Example: 500.00 minus 10.00 fee minus 30.00 funding leaves 460.00 in Wallet and 30.00 on card.",
    "No unsupported active card, retained programme fee or unmatched Wallet debit remains. Retry/query/reversal uses stable IDs and reconciles card issuance and fee posting to one final state.": "No unsupported active card, retained programme fee, orphan initial-funding credit or unmatched Wallet debit remains. Retry/query/reversal uses stable IDs and reconciles card issuance, fee posting and initial funding to one final state.",
    "Initial fee catalogue; fee code/trigger; fixed or percentage calculation; Wallet currency; tax; effective dates; member disclosure; refundability (including cancellation); configuration approvers; Merchant-balance exclusion; ledger/event fields; idempotency; reversal and reconciliation. Proposed current-release default is same-currency charging and no refund after a successfully delivered service unless the fee version explicitly says otherwise.": "Initial fee catalogue; fee code/trigger; fixed or percentage calculation; Wallet currency; tax; effective dates; member disclosure; refundability (including cancellation); configuration approvers; Merchant-balance exclusion; ledger/event fields; idempotency; reversal and reconciliation. Card creation additionally requires at least 20.00 initial funding in the Wallet/card currency. Proposed current-release default is same-currency charging and no fee refund after a successfully delivered service unless the fee version explicitly says otherwise.",
    "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH schemas; implement email OTP, wallet onboarding, Full Name carry-over, registration address, Sumsub, direct 2FA, zero-balance Wallet creation, read-only card data, international-name mapping, PhotonPay charset validation, prospective card-policy versioning, default BINs, separate card actions, OlyLife-office fulfilment, authenticated cancellation and the Wallet-fee configuration/ledger/reversal extension.": "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH schemas; implement email OTP, wallet onboarding, Full Name carry-over, registration address, Sumsub, direct 2FA, zero-balance Wallet creation, read-only card data, international-name mapping, PhotonPay charset validation, prospective card-policy versioning, default BINs, separate card actions, OlyLife-office fulfilment, authenticated cancellation, the Wallet-fee extension and atomic minimum initial-card-funding flow.",
}


def all_paragraphs(doc):
    seen = set()
    for paragraph in doc.paragraphs:
        if paragraph._p not in seen:
            seen.add(paragraph._p)
            yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph._p not in seen:
                        seen.add(paragraph._p)
                        yield paragraph


def set_text(paragraph, value):
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def replace_exact(doc):
    found = set()
    for paragraph in all_paragraphs(doc):
        replacement = REPLACEMENTS.get(paragraph.text)
        if replacement is None:
            replacement = TABLE_REPLACEMENTS.get(paragraph.text)
        if replacement is not None:
            found.add(paragraph.text)
            set_text(paragraph, replacement)
    missing = (set(REPLACEMENTS) | set(TABLE_REPLACEMENTS)) - found
    if missing:
        raise RuntimeError("Expected text not found:\n" + "\n".join(sorted(missing)))

    mapping = next(p for p in doc.paragraphs if p.text.startswith("Mapping conclusion  Reuse VCCHUB"))
    set_text(mapping, "Mapping conclusion  Reuse VCCHUB's cardholder, Wallet and virtual/physical card capabilities internally. Add VCCHUB email-OTP verification before the OlyLife member-validation call, extend VCCHUB fee processing so configured programme fees post to Wallet rather than Merchant balance, and extend card creation to require and atomically post minimum initial funding from Wallet to the new card. Retain the VCCHUB Wallet-credit/status interfaces consumed by OlyLife after top-up approval. Card creation, fee posting, initial funding and cancellation remain internal to VCCHUB. Invitation and SSO APIs are not required for the programme journey.")


def append_styled_row(table, values):
    new_tr = deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        set_text(cell.paragraphs[0], value)
        for paragraph in cell.paragraphs[1:]:
            set_text(paragraph, "")


def update_json(doc):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith('{\n  "event_id": "EVT-20260903-000060"'):
            value = paragraph.text.replace(
                '"card_balance": "0.00",\n    "wallet_balance": "490.00",',
                '"card_balance": "30.00",\n    "wallet_balance": "460.00",\n    "minimum_initial_balance": "20.00",\n    "total_wallet_deduction": "40.00",',
            ).replace(
                '"delivery_destination": "OLYLIFE_OFFICE",\n    "programme_fee": {',
                '"delivery_destination": "OLYLIFE_OFFICE",\n    "initial_funding": {\n      "transaction_id": "CTX-82000060",\n      "amount": "30.00",\n      "currency": "SGD",\n      "source": "WALLET",\n      "destination": "CARD"\n    },\n    "programme_fee": {',
            )
            set_text(paragraph, value)


def main():
    doc = Document(SOURCE)
    replace_exact(doc)

    append_styled_row(doc.tables[6], [
        "9. Minimum initial card funding",
        "Every new virtual or physical card must receive at least 20.00 in the Wallet/card currency at creation. The member may choose a higher amount. VCCHUB displays fee, funding, total deduction and resulting balances, validates Wallet coverage, then atomically creates the card, debits fee plus funding from Wallet, posts separate ledger entries and credits funding to the card. Failure retains no fee or funding debit and creates no usable card.",
        "Confirm or amend the 20.00 minimum, applicable currencies, maximum initial funding (if any), user-facing wording and whether future minimum changes are prospective or apply to existing Wallet policies.",
    ])
    append_styled_row(doc.tables[7], [
        "11",
        "A card cannot be created with a zero opening balance. The minimum initial card balance is 20.00 in the Wallet/card currency; the member may choose a higher amount.",
        "VCCHUB checks Wallet balance against initial funding plus every applicable fee. Card issuance, Wallet fee debit, Wallet-to-card funding debit/credit and ledger entries commit atomically. Example: Wallet 500.00 - fee 10.00 - initial funding 30.00 = Wallet 460.00 and card 30.00.",
    ])
    append_styled_row(doc.tables[36], [
        "Initial card funding",
        "Confirmed minimum 20.00 in the Wallet/card currency; whether a maximum applies; allowed decimal precision; UI disclosure; Wallet sufficiency check; separate funding transaction ID; atomic card/fee/funding posting; idempotency; failure reversal; webhook fields; cancellation refund treatment and reconciliation.",
    ])

    update_json(doc)

    props = doc.core_properties
    props.subject = "Implementation Responsibility Specification - Revision 0.17"
    props.last_modified_by = "Star SaaS Limited"
    props.comments = "Adds mandatory minimum initial card funding and atomic fee plus funding treatment."
    props.keywords = "OlyLife, VCCHUB, revision 0.17, minimum initial card balance, Wallet fees, atomic card funding, PhotonPay"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
