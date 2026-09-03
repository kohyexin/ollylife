from pathlib import Path
import re
from io import BytesIO
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.4.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15.docx"
BRAND_SOURCE = ROOT / "deliverables" / "Ollylife_VCCHUB_Implementation_Responsibility_Specification.docx"


# Longest and most specific phrases come first. These edits intentionally keep
# the document's existing styles, pagination controls, branding and content.
REPLACEMENTS = [
    ("PARTNER IMPLEMENTATION SPECIFICATION · VERSION 2.4", "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 0.15"),
    ("OlyLife–VCCHUB Wallet Activation & Card Program (Version 2.4)", "OlyLife–VCCHUB Wallet Activation & Card Program"),
    (
        "Purpose  Define who must build, expose, consume, operate and test each component of the Version 2 journey, where an OlyLife member starts in VCCHUB. This V2.4 edition adds the proposed mandatory end-to-end UAT suite, including Korean and Japanese non-Latin-script identity and PhotonPay cardholder-name compatibility tests. V2.3 remains retained as the preceding reference. Final production configuration and contractual approvals remain subject to joint sign-off.",
        "Purpose  Define who must build, expose, consume, operate and test each component of the programme journey, where an OlyLife member starts in VCCHUB. This 0.15 edition includes the proposed mandatory end-to-end UAT suite, including Korean and Japanese non-Latin-script identity and PhotonPay cardholder-name compatibility tests. The preceding document revision remains retained. Final production configuration and contractual approvals remain subject to joint sign-off.",
    ),
    ("0.15 (V2.4 mandatory UAT and international-name tests)", "0.15"),
    ("Draft - V2.4 mandatory UAT suite", "Draft - mandatory UAT suite"),
    ("3.5 Confirmed V2 programme assumptions and release constraints", "3.5 Confirmed programme assumptions and release constraints"),
    ("4.3 V2 partner interface catalogue", "4.3 Partner interface catalogue"),
    ("Confirmed V2 position", "Confirmed programme position"),
    ("V2 route independence", "Programme route independence"),
    ("Version 2 demonstration", "demonstration"),
    ("Version 2 journey", "programme journey"),
    ("V2 partner interface catalogue", "partner interface catalogue"),
    ("V2 partner interfaces", "partner interfaces"),
    ("V2 delivery baseline", "programme delivery baseline"),
    ("V2 wallet-user entry point", "wallet-user entry point"),
    ("V2 onboarding entry point", "wallet onboarding entry point"),
    ("V2 onboarding", "wallet onboarding"),
    ("V2 base journey", "programme journey"),
    ("V2 base flow", "programme flow"),
    ("V2 base currency", "programme base currency"),
    ("V2 card creation", "card creation"),
    ("V2 requirements", "programme requirements"),
    ("V2 physical card", "physical card"),
    ("V2 development", "Programme development"),
    ("VCCHUB V2 configuration scope", "VCCHUB configuration scope"),
    ("VCCHUB V2", "VCCHUB"),
    ("V2 mapping rule", "Mapping rule"),
    ("V2 programme", "programme"),
    ("V2 workshop", "programme workshop"),
    ("V2 position", "programme position"),
    ("V2 configuration", "programme configuration"),
    ("V2 release", "current release"),
    ("in V2", "in this release"),
    ("for V2", "for this release"),
    ("the V2", "the current release"),
    (" V2 ", " current release "),
    (" V2.", " current release."),
    (" V2,", " current release,"),
    ("The paths below are logical current release contracts", "The paths below are logical programme contracts"),
    ("used by the current release member-validation", "used by the programme's member-validation"),
    ("Approved current release journey", "Approved programme journey"),
    ("Collected during current release registration", "Collected during programme registration"),
    ("The full current release journey", "The full programme journey"),
    ("outside current release", "outside this release"),
    ("current release fixes each new Wallet", "The programme fixes each new Wallet"),
    ("current release has no multi-BIN", "This release has no multi-BIN"),
    ("current release starts in VCCHUB", "The programme starts in VCCHUB"),
    ("current release does not require", "This release does not require"),
    ("current release does not depend", "This release does not depend"),
    ("current release does not use", "This release does not use"),
]


VERSION_PATTERN = re.compile(r"(?i)(?:\bV2(?:\.\d+)?\b|Version\s+2(?:\.\d+)?)")


def iter_paragraphs(doc):
    seen = set()

    def emit(paragraph):
        key = paragraph._p
        if key not in seen:
            seen.add(key)
            yield paragraph

    for paragraph in doc.paragraphs:
        yield from emit(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield from emit(paragraph)
    for section in doc.sections:
        for story in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            for paragraph in story.paragraphs:
                yield from emit(paragraph)
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield from emit(paragraph)


def replace_in_run_text(text):
    for old, new in REPLACEMENTS:
        text = text.replace(old, new)
    return text


def replace_paragraph(paragraph):
    # Preserve run-level formatting whenever the target text is contained in a
    # run, which is true for the source document's version labels and prose.
    for run in paragraph.runs:
        run.text = replace_in_run_text(run.text)

    # Repair any context-sensitive phrase that happened to span runs. Retain
    # the first run's formatting; this fallback is needed only for a few table
    # cells where the old version tag was isolated in its own run.
    current_text = paragraph.text
    full_text = replace_in_run_text(current_text)
    if full_text != current_text or VERSION_PATTERN.search(full_text):
        if VERSION_PATTERN.search(full_text):
            full_text = VERSION_PATTERN.sub("current release", full_text)
            full_text = replace_in_run_text(full_text)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(full_text)


def add_logo(paragraph, logo_blob, *, space_after_pt):
    """Restore the body logo using the exact dimensions from the branded source."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    run = paragraph.add_run()
    inline = run.add_picture(
        BytesIO(logo_blob),
        width=Emu(1856231),
        height=Emu(327004),
    )
    inline._inline.docPr.set("descr", "Star SaaS logo")


def restore_front_matter_branding(doc):
    if not BRAND_SOURCE.exists():
        raise FileNotFoundError(BRAND_SOURCE)
    with ZipFile(BRAND_SOURCE) as source_package:
        logo_blob = source_package.read("word/media/image1.png")

    # Paragraph 0 is the reserved cover-logo position. The blank paragraph
    # immediately before the disclaimer already carries its page break and is
    # the reserved second-page logo position.
    add_logo(doc.paragraphs[0], logo_blob, space_after_pt=10)
    disclaimer_anchor = next(
        p for p in doc.paragraphs
        if p.paragraph_format.page_break_before and not p.text.strip()
    )
    add_logo(disclaimer_anchor, logo_blob, space_after_pt=28)

    # Keep page 2 exclusively for the proprietary/confidential disclaimer.
    section_one = next(
        p for p in doc.paragraphs if p.text.strip() == "1. Recommended ownership model"
    )
    section_one.paragraph_format.page_break_before = True


def normalize_numbering_order(doc):
    """Keep custom list definitions schema-valid for Microsoft Word.

    The source document appended its custom abstract numbering definitions
    after existing ``w:num`` instances. Lightweight previewers tolerated that
    ordering, but Word repaired the numbering part on open and displayed the
    intended bullet lists as one continuing decimal sequence. OOXML requires
    all abstract definitions to precede their concrete list instances.
    """
    numbering = doc.part.numbering_part.element
    abstract_tag = qn("w:abstractNum")
    num_tag = qn("w:num")
    children = list(numbering)
    prefix = [child for child in children if child.tag not in (abstract_tag, num_tag)]
    abstracts = [child for child in children if child.tag == abstract_tag]
    nums = [child for child in children if child.tag == num_tag]
    numbering[:] = prefix + abstracts + nums


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    doc = Document(SOURCE)
    for paragraph in iter_paragraphs(doc):
        replace_paragraph(paragraph)

    restore_front_matter_branding(doc)
    normalize_numbering_order(doc)

    props = doc.core_properties
    props.title = "OlyLife-VCCHUB Wallet Activation & Card Program"
    props.subject = "Implementation Responsibility Specification - Revision 0.15"
    props.author = "Star SaaS Limited"
    props.last_modified_by = "Star SaaS Limited"
    props.comments = "Single-version programme specification, revision 0.15."
    props.keywords = (
        "OlyLife, VCCHUB, revision 0.15, programme assumptions, card policy, "
        "default BIN, PhotonPay fees, Wallet top-up, cancellation"
    )

    remaining = [p.text for p in iter_paragraphs(doc) if VERSION_PATTERN.search(p.text)]
    if remaining:
        raise RuntimeError("Version 2 references remain: " + " | ".join(remaining[:5]))

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
