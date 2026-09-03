from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15_ZH-CN.docx"


def set_east_asia_font(run, name: str) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def main() -> None:
    document = Document(TARGET)

    disclaimer = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("本文件及其中所载信息均归 Star SaaS Limited")
    )
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    disclaimer.paragraph_format.space_before = Pt(0)
    disclaimer.paragraph_format.space_after = Pt(18)
    disclaimer.paragraph_format.line_spacing = 1.3
    disclaimer.paragraph_format.keep_together = True
    disclaimer.paragraph_format.keep_with_next = True

    for run in disclaimer.runs:
        run.font.size = Pt(11)
        run.font.bold = False
        run.font.italic = False
        run.font.color.rgb = RGBColor(0x16, 0x27, 0x38)
        set_east_asia_font(run, "Microsoft YaHei")

    control = next(
        paragraph for paragraph in document.paragraphs if paragraph.text.startswith("文件控制")
    )
    control.paragraph_format.keep_together = True

    section_one = next(
        paragraph for paragraph in document.paragraphs if paragraph.text == "1. 建议的责任归属模型"
    )
    section_one.paragraph_format.page_break_before = True

    document.save(TARGET)


if __name__ == "__main__":
    main()
