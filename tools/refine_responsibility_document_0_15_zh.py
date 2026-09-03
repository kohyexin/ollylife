from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15.docx"
TARGET = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15_ZH-CN.docx"


def paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from paragraphs(cell)


MANUAL = {
    "Owns the authoritative member identity record, Sumsub contract/account, commission ledger and top-up approval process.": "负责维护权威会员身份记录、Sumsub 合同/账户、佣金账本及充值审批流程。",
    "Return an active-member result with Member ID, First Name, Last Name and the matched email, plus phone data where approved.": "会员有效时，返回会员编号、名字、姓氏及匹配的电子邮箱；经批准时一并返回电话号码资料。",
    "Provide the OlyLife top-up request, Admin/Support approval, atomic commission debit/reversal and VCCHUB wallet-credit orchestration.": "提供 OlyLife 充值申请、管理员/运营人员审批、佣金原子扣减/冲正，以及 VCCHUB 钱包入账编排。",
    "Create the cardholder and zero-balance wallet after Sumsub approval; show member/cardholder/address data read-only during card creation, enforce one active virtual plus one active physical card, show every card as a separate actionable record, use the configured OlyLife office address for physical fulfilment and provide authenticated card cancellation.": "Sumsub 审批通过后创建持卡人及零余额钱包；创建卡片时以只读方式显示会员、持卡人及地址资料；限制为一张有效虚拟卡及一张有效实体卡；每张卡片均作为可独立操作的记录显示；实体卡配送使用已配置的 OlyLife 办公室地址；并提供经身份验证的卡片取消功能。",
    "4.5 VCC-01 — Credit member wallet (POST /wallet/topup)": "4.5 VCC-01——会员钱包入账（POST /wallet/topup）",
    "VCC-01 is provided by VCCHUB and called by OlyLife only after an approved, atomic commission debit. Monetary values use decimal strings and ISO 4217 currency codes. request_id is the exactly-once key across debit, API response, webhook and reconciliation.": "VCC-01 由 VCCHUB 提供。OlyLife 仅可在充值获批且佣金已原子扣减后调用。金额使用十进制字符串，币种使用 ISO 4217 代码。request_id 是贯穿佣金扣减、API 响应、Webhook 及对账流程的恰好一次幂等键。",
    "VCC-02 is the authoritative reconciliation endpoint when an account-ready or lifecycle webhook is delayed, duplicated or missed. OlyLife may query by external_member_id, wallet_id, or both; if both are supplied they must map to the same wallet.": "当钱包账户就绪或生命周期 Webhook 延迟、重复或遗漏时，VCC-02 是权威对账查询接口。OlyLife 可使用 external_member_id、wallet_id 或两者进行查询；若同时提供，两者必须映射至同一钱包。",
    "4.7 Partner webhook message contract": "4.7 合作方 Webhook 消息约定",
    "4.7.2 Common webhook envelope": "4.7.2 通用 Webhook 消息封装",
    "4.7.4 Illustrative webhook messages": "4.7.4 Webhook 消息示例",
    "4.7.5 OlyLife acknowledgement example": "4.7.5 OlyLife 确认响应示例",
    "4.8 OlyLife internal top-up request and approval workflow": "4.8 OlyLife 内部充值申请及审批流程",
    "Sumsub: OlyLife owns the contract/account and securely provisions named access/API credentials; VCCHUB implements token creation, WebSDK configuration, signed webhook verification, review-status retrieval and approved applicant-data retrieval.": "Sumsub：合同及账户归 OlyLife 所有，并由 OlyLife 安全提供实名访问权限/API 凭据；VCCHUB 负责实现访问令牌创建、WebSDK 配置、已签名 Webhook 验证、审核状态查询及获批申请人资料查询。",
    "PhotonPay issuer: VCCHUB reuses existing virtual/physical issuance and cancellation capabilities. For a physical card, VCCHUB uses the configured OlyLife office delivery address and does not collect a member recipient address.": "PhotonPay 发卡机构：VCCHUB 复用现有虚拟卡/实体卡发行及取消能力。实体卡使用已配置的 OlyLife 办公室配送地址，不向会员收集收件地址。",
    "Why this section remains  Keep a concise reference trail for the signing, envelope and internal Issuing conventions used by the programme's member-validation, wallet-credit, card top-up and card-cancellation operations. It is not an OlyLife checklist for VCCHUB-internal card APIs.": "保留本节的原因  本节简要记录会员验证、钱包入账、卡片充值及卡片取消操作所采用的签名、消息封装及内部发卡规范；它并非要求 OlyLife 对 VCCHUB 内部卡片 API 逐项实施的清单。",
    "Existing OlyLife member signs up in VCCHUB; OLY-01 returns active status, Member ID, First Name, Last Name and matched email; registration continues with those fields shown read-only.": "现有 OlyLife 会员在 VCCHUB 注册；OLY-01 返回有效状态、会员编号、名字、姓氏及匹配的电子邮箱；注册继续进行，并以只读方式显示这些字段。",
    "OlyLife top-up remains pending with no balance change until Admin/Support approval; approval debits commission and credits VCCHUB exactly once; rejection changes neither balance.": "OlyLife 充值在管理员/运营人员审批前保持待处理，余额不变；审批通过后仅扣减一次佣金并仅向 VCCHUB 钱包入账一次；拒绝则两个余额均不变。",
    "Card creation shows member/cardholder and registered-address information read-only. The programme permits one active virtual card and one active physical card; both appear as separate rows with independent action menus. Physical creation collects no recipient and uses the configured OlyLife office address.": "创建卡片时，会员/持卡人及注册地址资料均以只读方式显示。本项目允许一张有效虚拟卡及一张有效实体卡，两者分别显示为独立记录并具有各自的操作菜单。创建实体卡时不收集收件人资料，系统使用已配置的 OlyLife 办公室地址。",
    "Wallet-to-card top-up shows Wallet balance, transfers exactly once and records balances/transactions against the selected card. Card cancellation requires fresh 2FA, is irreversible, returns the eligible card balance to Wallet, records the event and releases only the cancelled card-type slot so that type can be replaced.": "钱包向卡片充值时显示钱包余额，仅执行一次资金转移，并在所选卡片下记录余额及交易。取消卡片必须重新完成 2FA，且不可撤销；符合条件的卡片余额退回钱包，系统记录该事件，并仅释放已取消卡片类型的名额，以便重新创建同类型卡片。",
    "8.2.4 Card creation, funding, cancellation and entitlement": "8.2.4 卡片创建、充值、取消及额度规则",
    "Confirmed programme decisions  PhotonPay remains the issuer. The programme starts in VCCHUB, validates OlyLife membership before registration, collects the member's full address, creates a zero-balance Wallet with no card after KYC, and funds Wallets only through OlyLife's approved top-up process. Each Wallet is assigned a policy allowing one active virtual and one active physical card, with one default BIN per type and no user BIN selection. Policy changes apply prospectively unless a separate existing-wallet migration is approved. Physical cards are sent to the configured OlyLife office. Card cancellation releases the cancelled type slot. OlyLife handles programme/card fee charging directly with PhotonPay; VCCHUB does not deduct such fees from Wallet balance in this release.": "已确认的项目决定  PhotonPay 继续担任发卡机构。流程从 VCCHUB 开始，注册前验证 OlyLife 会员资格并收集会员完整地址；KYC 通过后创建零余额且无卡的钱包；钱包只能通过 OlyLife 已审批的充值流程入金。每个钱包适用一项卡片政策，允许一张有效虚拟卡及一张有效实体卡；每种类型使用一个默认 BIN，会员不可选择 BIN。除非另行批准现有钱包迁移，否则政策变更仅适用于之后新建的钱包。实体卡寄送至已配置的 OlyLife 办公室。取消卡片会释放相应卡片类型的名额。OlyLife 直接与 PhotonPay 处理项目/卡片费用，本版本中 VCCHUB 不从钱包余额扣除此类费用。",
    "Definition of ready  Programme development is ready when OLY-01 member validation, account mapping, OlyLife approval/debit, VCCHUB Wallet credit/status, Sumsub access, international-name mapping and PhotonPay compatibility evidence, office-address configuration, authenticated card cancellation and acceptance owners are documented and testable.": "就绪定义  当 OLY-01 会员验证、账户映射、OlyLife 审批/扣款、VCCHUB 钱包入账/状态、Sumsub 访问权限、国际姓名映射及 PhotonPay 兼容性证据、办公室地址配置、经身份验证的卡片取消，以及验收负责人均已形成文件并可测试时，项目开发方可视为就绪。",
    "Handoff checkpoint  The first programme workshop should close the member-status eligibility rules, OlyLife-to-Sumsub name precedence, account-ready mapping, wallet-credit idempotency/reversal, real OlyLife office address, physical fulfilment controls, default BINs, prospective card-policy versioning and cancellation authentication/audit details. Each NEW or EXTEND item can then become an owned backlog item and contract test.": "交接检查点  首次项目工作会议应确认会员状态资格规则、OlyLife 与 Sumsub 姓名取值优先级、账户就绪映射、钱包入账幂等/冲正、实际 OlyLife 办公室地址、实体卡履约控制、默认 BIN、仅面向新钱包的卡片政策版本管理，以及卡片取消的身份验证/审计细节。随后，每个 NEW 或 EXTEND 项目均可转化为有明确负责方的待办事项及契约测试。",
    "Top-up approval & funding": "充值审批及入金",
    "OlyLife owns request approval and commission debit; VCCHUB owns exactly-once wallet credit and wallet ledger.": "OlyLife 负责申请审批及佣金扣减；VCCHUB 负责钱包恰好一次入账及钱包账本。",
    "Top-up request": "充值申请",
    "Approval & wallet credit": "审批及钱包入账",
    "Apply each action to the selected card. Cancellation requires fresh 2FA, refunds eligible card balance to Wallet and releases only that card-type slot.": "每项操作仅作用于所选卡片。取消卡片须重新完成 2FA，将符合条件的卡片余额退回钱包，并仅释放该卡片类型的名额。",
    "Use Sumsub-approved first/last name and DOB as authoritative cardholder identity": "使用 Sumsub 审批通过的名字、姓氏及出生日期作为权威持卡人身份资料",
    "No commission top-up button in VCCHUB": "VCCHUB 不提供佣金充值按钮",
    "Admin/Support approval": "管理员/运营人员审批",
    "Credit wallet exactly once and return transaction reference": "仅向钱包入账一次，并返回交易参考号",
    "Reverse/adjust commission when an approved debit cannot produce a final wallet credit": "若已批准的佣金扣减未能最终完成钱包入账，则冲正/调整佣金",
    "Check Wallet balance, debit Wallet and credit the selected card atomically. No VCCHUB-configured programme/card fee is deducted from Wallet in this release": "检查钱包余额，并以原子方式扣减钱包及向所选卡片入账。本版本不从钱包扣除任何由 VCCHUB 配置的项目/卡片费用",
    "After KYC approval and zero-balance wallet creation, VCCHUB sends a signed wallet_account_ready event containing external_member_id, vcchub_user_id, cardholder_id and wallet_id. OlyLife stores one active member-to-wallet mapping and uses POST /wallet/status to reconcile uncertain events. Top-up stays disabled until the mapping is unambiguous and ready.": "KYC 审批通过并创建零余额钱包后，VCCHUB 发送已签名的 wallet_account_ready 事件，其中包含 external_member_id、vcchub_user_id、cardholder_id 及 wallet_id。OlyLife 保存唯一有效的会员与钱包映射，并使用 POST /wallet/status 对账不确定事件。在映射明确且状态就绪前，充值功能保持禁用。",
    "5. Top-up approval, debit and reversal": "5. 充值审批、扣减及冲正",
    "The member submits a request in OlyLife and balances remain unchanged while pending. One authorised Admin/Support reviewer may approve normal requests; maker-checker applies above OlyLife's risk threshold and to manual adjustments/reversals. Approval atomically rechecks eligibility and commission, debits once, then calls POST /wallet/topup with the same idempotency chain. A permanent credit failure triggers an idempotent commission reversal or controlled repair.": "会员在 OlyLife 提交申请，待处理期间各项余额保持不变。一般申请可由一名获授权的管理员/运营审核人员审批；超过 OlyLife 风险阈值的申请及人工调整/冲正须采用经办复核双人机制。审批时以原子方式重新检查资格及佣金余额，仅扣减一次佣金，并沿用同一幂等链调用 POST /wallet/topup。若钱包最终入账失败，则触发幂等佣金冲正或受控修复。",
    "Wallet top-up cannot be initiated from VCCHUB. Every wallet top-up starts in OlyLife and follows OlyLife's internal approval window.": "不得从 VCCHUB 发起钱包充值。每笔钱包充值均从 OlyLife 发起，并遵循 OlyLife 的内部审批流程。",
    "Balances remain unchanged while approval is pending. After approval, OlyLife debits commission once and calls VCCHUB POST /wallet/topup; VCCHUB credits the Wallet exactly once and returns or publishes the final result.": "审批待处理期间余额保持不变。审批通过后，OlyLife 仅扣减一次佣金并调用 VCCHUB 的 POST /wallet/topup；VCCHUB 仅向钱包入账一次，并返回或发布最终结果。",
    "The production login/MFA controls belong to VCCHUB. Funds credited to the Wallet are used within the card programme; any future wallet-to-commission withdrawal or SSO flow requires separate scope and approval.": "生产环境登录/MFA 控制由 VCCHUB 负责。已入账钱包的资金仅用于卡片项目；未来如需钱包资金退回佣金余额或提供 SSO 流程，须另行确定范围并审批。",
    "Create a pending member request; no commission debit or wallet credit.": "创建待处理的会员充值申请；不扣减佣金，也不向钱包入账。",
    "Approve/reject top-up": "批准/拒绝充值",
    "Reverse an eligible debit if wallet credit cannot complete.": "若钱包无法完成入账，则冲正符合条件的佣金扣款。",
    "Wallet top-up result": "钱包充值结果",
    "Final or failed asynchronous wallet-credit result.": "钱包异步入账的最终成功或失败结果。",
    "Synchronize successful or failed wallet-to-card funding.": "同步钱包向卡片充值的成功或失败结果。",
    "Card cancellation": "卡片取消",
    "Require fresh 2FA and explicit irreversible confirmation; refund eligible card balance to Wallet exactly once, mark the selected card permanently cancelled, release only that card-type slot and record the cancellation audit/transaction. The other card type remains occupied. Wallet withdrawal remains unsupported.": "须重新完成 2FA 并明确确认该操作不可撤销；将符合条件的卡片余额仅退回钱包一次，将所选卡片永久标记为已取消，仅释放该卡片类型的名额，并记录取消审计及交易。另一卡片类型的名额仍被占用。仍不支持钱包提现。",
    "For Korean and Japanese approved applicants, call the PhotonPay sandbox cardholder/card-creation flow for both configured BINs using the agreed name field and UTF-8 encoding. Test documented maximum lengths and physical-card embossing constraints.": "针对已通过审核的韩国及日本申请人，使用约定的姓名字段及 UTF-8 编码，对两个已配置 BIN 调用 PhotonPay 沙盒持卡人/卡片创建流程，并测试文件规定的最大长度及实体卡压印限制。",
    "Exactly one cardholder and SGD 0.00 wallet are created with no card. Mapping IDs are returned by status/webhook and top-up eligibility is unambiguous.": "仅创建一个持卡人及一个余额为 SGD 0.00 的无卡钱包。状态接口/Webhook 返回映射标识，且充值资格明确无歧义。",
    "No VCCHUB top-up action is exposed. While pending, commission and wallet balances remain unchanged and a complete approval record exists.": "VCCHUB 不提供充值操作。待审批期间，佣金及钱包余额保持不变，并须保留完整审批记录。",
    "Approve a request with sufficient commission; OlyLife atomically debits once and calls POST /wallet/topup.": "批准佣金余额充足的申请；OlyLife 以原子方式仅扣减一次佣金，并调用 POST /wallet/topup。",
    "Only valid 2FA cancels. Eligible card balance returns to Wallet exactly once; card is permanently Cancelled, event is emitted and only its type slot is released.": "只有有效的 2FA 才能完成取消。符合条件的卡片余额仅退回钱包一次；卡片永久标记为“已取消”，系统发出相应事件，并仅释放该卡片类型的名额。",
    "OlyLife commission, VCCHUB Wallet, selected card balances and lifecycle states reconcile with no orphan, duplicate or unexplained amount.": "OlyLife 佣金、VCCHUB 钱包、所选卡片余额及生命周期状态均可完成对账，不存在孤立、重复或无法解释的金额。",
    "Exactly-once money effects and single-resource transitions hold; processing resumes or reconciles without manual data corruption.": "资金影响保持恰好一次语义，单一资源状态转换保持一致；处理可恢复或完成对账，且不会造成手工数据损坏。",
    "Support tiers, escalation contacts, monitoring ownership, incident notification, DR/RTO/RPO, audit and reporting cadence.": "支持级别、升级联系人、监控责任、事件通知、灾难恢复/RTO/RPO、审计及报告频率。",
    "Cardholder, zero-balance wallet, registration address, cards, balances, top-ups, cancellations and transaction history.": "持卡人、零余额钱包、注册地址、卡片、余额、充值、卡片取消及交易记录。",
    "Show cardholder and registered address read-only. Permit one active virtual plus one active physical card and list both separately; physical uses the configured OlyLife office delivery address.": "以只读方式显示持卡人及注册地址。允许一张有效虚拟卡及一张有效实体卡，并分别列出；实体卡使用已配置的 OlyLife 办公室配送地址。",
    "Joint UAT covers matched/not-found/unavailable member checks, address validation, KYC outcomes, ready mapping, approval/rejection/insufficient commission, exactly-once wallet credit, reversal, one card per type, office delivery, card top-up and authenticated cancellation. Go-live requires all critical/high scenarios passed, no open Severity 1 or 2 defect, production credentials/configuration validated and signed Product, Technology and Operations/Compliance approval from both parties.": "联合 UAT 涵盖会员匹配/未找到/服务不可用、地址验证、KYC 结果、账户就绪映射、审批/拒绝/佣金不足、钱包恰好一次入账、冲正、每种类型限一张卡、办公室配送、卡片充值及经身份验证的卡片取消。上线前须通过所有关键及高优先级场景，不得存在未解决的严重等级 1 或 2 缺陷，生产凭据及配置须完成验证，并取得双方产品、技术及运营/合规负责人的签署批准。",
    "Reuse internally for Wallet-to-card top-up; first check Wallet balance, then pair request_id with VCCHUB idempotency, Wallet debit and selected-card credit.": "内部复用于钱包向卡片充值；先检查钱包余额，再将 request_id 与 VCCHUB 幂等控制、钱包扣款及所选卡片入账关联。",
    "0000 means success. Any other code must be handled as a business/technical outcome, not inferred from message text.": "0000 表示成功。任何其他代码均须按明确的业务/技术结果处理，不得仅根据消息文本推断。",
    "VCCHUB wallet-ledger reference persisted by OlyLife.": "由 OlyLife 持久保存的 VCCHUB 钱包账本参考号。",
    "Delivery is at least once and may be out of order. resource_version increases for each resource; OlyLife ignores an older version after acknowledging it.": "Webhook 采用至少一次交付语义，事件可能乱序到达。每个资源的 resource_version 递增；OlyLife 确认后应忽略较旧版本。",
    "SHA-256 signature over the full envelope except sign.": "对除 sign 字段以外的完整消息封装计算 SHA-256 签名。",
    "Rejection leaves both balances unchanged. If debit succeeds but wallet credit fails, retry safely or initiate OLY-04 reversal/repair.": "拒绝申请时两个余额均保持不变。若佣金扣减成功但钱包入账失败，应安全重试，或启动 OLY-04 冲正/修复。",
    "Pending request, Admin/Support decision, atomic debit, exactly-once wallet credit, reversal and reconciliation pass.": "待处理申请、管理员/运营人员决定、原子扣减、钱包恰好一次入账、冲正及对账均通过。",
    "Funding gate; one active virtual plus one active physical card; prospective card-policy versioning; default BIN per type; PhotonPay charset/name and embossing compatibility; separate card rows/actions; read-only identity/address; OlyLife-office delivery; Wallet-to-card top-up; authenticated cancellation with Wallet refund and type-slot release; no VCCHUB Wallet-fee deduction.": "入金门槛；一张有效虚拟卡及一张有效实体卡；仅面向新钱包的卡片政策版本管理；每种卡片类型使用默认 BIN；验证 PhotonPay 字符集、姓名及压印兼容性；卡片分别显示且可独立操作；身份及地址只读；配送至 OlyLife 办公室；钱包向卡片充值；经身份验证的取消、余额退回钱包及相应类型名额释放；VCCHUB 不从钱包扣除项目/卡片费用。",
    "Credit failure and reversal — Joint": "入账失败及冲正——双方",
    "Simulate successful OlyLife debit followed by final VCCHUB wallet-credit failure.": "模拟 OlyLife 成功扣减佣金后，VCCHUB 钱包最终入账失败。",
    "Authenticated cancellation/refund — VCCHUB": "经身份验证的卡片取消/退款——VCCHUB",
    "Reviewer roles/dual control, limits, atomic debit, VCC-01 idempotency, settlement, reversal, reconciliation and manual adjustment controls.": "审核人员角色/双人控制、限额、原子扣减、VCC-01 幂等控制、结算、冲正、对账及人工调整控制。",
    "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH schemas; implement wallet onboarding, member carry-over, registration address, Sumsub, direct 2FA, zero-balance Wallet creation, read-only card data, international-name mapping, PhotonPay charset validation, prospective card-policy versioning, one default BIN per type, separate per-card actions, OlyLife-office physical fulfilment and authenticated cancellation with Wallet refund/type-slot release.": "发布 OLY-01 调用方要求及 VCC-01/02/VCC-WH 数据架构；实现钱包开户、会员资料带入、注册地址、Sumsub、直接登录 2FA、零余额钱包创建、只读卡片资料、国际姓名映射、PhotonPay 字符集验证、仅面向新钱包的卡片政策版本管理、每种类型一个默认 BIN、每张卡片独立操作、配送至 OlyLife 办公室，以及经身份验证的卡片取消并退回钱包余额/释放相应类型名额。",
    "Owns the wallet-user entry point, onboarding journey and wallet/card product orchestration.": "负责钱包用户入口、开户流程及钱包/卡片产品编排。",
    "Receiver algorithm: verify HTTPS source controls and signature; validate event_version/schema; check event_id; durably store the event and processing state; apply only when resource_version is newer; commit the local update; then return HTTP 200. A duplicate valid event returns the same successful acknowledgement without reapplying the state change.": "接收方处理逻辑：验证 HTTPS 来源控制及签名；验证 event_version/数据架构；检查 event_id；持久保存事件及处理状态；仅在 resource_version 较新时应用；提交本地更新；随后返回 HTTP 200。对于重复但有效的事件，应返回相同的成功确认，不得再次应用状态变更。",
    "Admin/Support approves, OlyLife atomically debits commission and calls VCCHUB to credit the wallet exactly once.": "管理员/运营人员审批后，OlyLife 以原子方式扣减佣金，并调用 VCCHUB 使钱包恰好一次入账。",
    "No onboarding UI required": "无需开户界面",
    "Provide and confirm the legal delivery address, recipient/contact, telephone, operating hours, courier instructions, change approvers and failed-delivery owner.": "提供并确认合法配送地址、收件人/联系人、电话、办公时间、快递指示、地址变更审批人及配送失败责任人。",
    "wallet onboarding starts in VCCHUB. OlyLife validates membership by email and returns Member ID, First Name, Last Name and the matched email for an eligible member.": "钱包开户从 VCCHUB 开始。OlyLife 按电子邮箱验证会员资格，并针对符合条件的会员返回会员编号、名字、姓氏及匹配的电子邮箱。",
    "Build wallet onboarding, username=email, member fields, Terms, phone/password and registered-address capture.": "构建钱包开户流程，包括用户名=电子邮箱、会员资料、条款、电话/密码及注册地址收集。",
    "Not exposed in the programme journey": "不在项目用户流程中开放",
    "Consumer handling": "调用方处理要求",
    "Echo the request_id for end-to-end correlation.": "原样返回 request_id，以支持端到端关联。",
    "Echoed exactly-once/correlation key.": "原样返回的恰好一次/关联键。",
    "Echoed mapping values; reject/reconcile any mismatch.": "原样返回映射值；任何不匹配均须拒绝或对账处理。",
    "Retry the identical payload with the same request_id, then reconcile. Never issue a second commission debit.": "使用相同 request_id 重试完全相同的请求内容，随后进行对账。不得再次扣减佣金。",
    "When true, return non-sensitive card summaries for lifecycle reconciliation.": "为 true 时，返回非敏感卡片摘要，以进行生命周期对账。",
    "Echoed correlation key.": "原样返回的关联键。",
    "Any non-2xx or timeout triggers the Star SaaS webhook retry pattern: first retry after about one second, exponential backoff with jitter and a delivery window of up to 36 hours. Exhausted events enter operational reconciliation.": "任何非 2xx 响应或超时均触发 Star SaaS Webhook 重试机制：约一秒后首次重试，随后采用带随机抖动的指数退避，最长交付窗口为 36 小时。重试耗尽的事件进入运营对账流程。",
    "Use the agreed Star SaaS signing convention over the canonical JSON body: recursively sort object keys, exclude sign, append the environment signkey and calculate lowercase SHA-256 hex. Enforce the timestamp replay window, event_id idempotency, HTTPS and source allowlisting as appropriate.": "对规范化 JSON 请求体采用约定的 Star SaaS 签名规则：递归排序对象键，排除 sign 字段，附加对应环境的 signkey，并计算小写 SHA-256 十六进制值。应适当强制执行时间戳防重放窗口、event_id 幂等控制、HTTPS 及来源白名单。",
    "Recheck at approval, decline atomically without debit and create no wallet credit.": "审批时重新检查；若拒绝，则以原子方式处理，不扣款也不向钱包入账。",
    "Complete onboarding and login without a V1 invitation email, OlyLife-to-VCCHUB SSO or return-to-OlyLife control.": "在无需邀请邮件、OlyLife 至 VCCHUB SSO 或返回 OlyLife 控件的情况下完成开户及登录。",
    "No unsupported active card or unmatched final debit remains. Retry/query/reversal uses stable IDs and reconciles to one final state.": "不得遗留不受支持的有效卡片或未匹配的最终扣款。重试、查询及冲正均使用稳定标识，并对账至唯一最终状态。",
    "No password, signing key, raw biometric/document image, PAN or CVV is exposed outside approved controls; masking, retention and role access are evidenced.": "密码、签名密钥、原始生物识别资料/证件图像、PAN 或 CVV 均不得暴露于已批准的控制范围之外；须提供数据掩码、保留期限及角色访问控制证据。",
    "Reconcile a complete member journey using Member ID, wallet/cardholder/card IDs, request_id, event_id and both ledger references.": "使用会员编号、钱包/持卡人/卡片标识、request_id、event_id 及双方账本参考号，对完整会员流程进行对账。",
    "Actual OlyLife office address, authorised change process, courier, tracking, failed delivery, replacement and office receiving owner. OlyLife agrees card/programme fee charging directly with PhotonPay; VCCHUB does not deduct those fees from Wallet in this release.": "确认实际 OlyLife 办公室地址、获授权的变更流程、快递、追踪、配送失败、补发及办公室收件责任人。OlyLife 直接与 PhotonPay 约定卡片/项目费用的收取方式；本版本中 VCCHUB 不从钱包扣除此类费用。",
    "New committed or final operation status.": "新提交或最终操作状态。",
    "KYC-approved cardholder and zero-balance wallet are committed": "已提交 KYC 审批通过的持卡人及零余额钱包",
    "Approved commission-funded wallet credit commits": "已提交经审批的佣金资金钱包入账",
    "Wallet-to-card transfer commits": "已提交钱包向卡片的资金转移",
    "Cancellation, eligible Wallet refund and slot release commit": "已提交卡片取消、符合条件的钱包退款及名额释放",
    "Internal VCCHUB specification / implementation": "VCCHUB 内部规范/实施",
    "Approved programme journey; responsibilities; member-data/KYC/address purpose; API/event schemas; security and test plan.": "已批准的项目流程、职责、会员资料/KYC/地址用途、API/事件数据架构、安全及测试计划。",
    "The full programme journey succeeds using VCCHUB sign-up, direct sign-in and 2FA only.": "完整项目流程仅通过 VCCHUB 注册、直接登录及 2FA 即可成功完成。",
}

REPLACEMENTS = {
    "上调": "充值", "上加": "充值", "上架": "充值", "顶端": "充值",
    "物理卡": "实体卡", "活动虚拟卡": "有效虚拟卡", "活动物理卡": "有效实体卡",
    "委员会借项": "佣金扣减", "佣金借项": "佣金扣减", "wallet信用": "钱包入账",
    "Wallet信用": "钱包入账", "wallet 信用": "钱包入账", "Wallet 信用": "钱包入账",
    "wallet信贷": "钱包入账", "Wallet信贷": "钱包入账", "借记Wallet": "向钱包入账",
    "调节": "对账", "同位素": "幂等", "倒置": "冲正", "反转": "冲正",
    "注销牌": "卡片取消", "注销卡": "卡片取消", "卡片注销": "卡片取消",
    "取消牌": "取消卡片", "插槽": "名额", "类型槽": "类型名额", "移徙": "迁移",
    "合伙人webhook电文合同": "合作方 Webhook 消息约定", "webhook信封": "Webhook 消息封装",
    "webhook电文": "Webhook 消息", "标准响应信封": "标准响应封装", "精确的一键": "恰好一次幂等键",
    "节目": "项目", "发行中": "本版本中", "支助": "运营", "核准": "审批",
    "物理发行": "实体卡发行", "物理履行": "实体卡履约", "实际履行": "实体卡履约",
    "物理卡化": "实体卡压印", "物理卡模具": "实体卡压印", "物理 BIN": "实体卡 BIN",
    "虚拟 BIN": "虚拟卡 BIN", "Wallet到卡": "钱包向卡片", "wallet到卡": "钱包向卡片",
    "Wallet-to-card": "钱包向卡片", "wallet- 入卡": "钱包向卡片", "wallet- 退出": "钱包提现",
    "Wallet 撤回": "钱包提现", "Wallet 退出": "钱包提现", "Wallet余额": "钱包余额",
    "wallet余额": "钱包余额", "Wallet创建": "钱包创建", "wallet创建": "钱包创建",
    "零余额wallet": "零余额钱包", "零余额 wallet": "零余额钱包", "成员Wallet": "会员钱包",
    "卡持有者": "持卡人", "成员身份": "会员编号", "姓、 姓": "名字、姓氏",
    "合伙人/公司": "合作方/公司",
    "姓名、姓名": "名字、姓氏", "第一/最后姓名": "名字/姓氏", "名字,姓": "名字、姓氏",
    "承认的例子": "确认响应示例", "在永久接受之前不要承认": "在持久化接收前不得返回确认",
}


def replace_text(paragraph, text):
    if paragraph.text == text:
        return
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


source_doc = Document(SOURCE)
target_doc = Document(TARGET)
source_paragraphs = list(paragraphs(source_doc))
target_paragraphs = list(paragraphs(target_doc))

for source_paragraph, target_paragraph in zip(source_paragraphs, target_paragraphs):
    if source_paragraph.text in MANUAL:
        replace_text(target_paragraph, MANUAL[source_paragraph.text])
        continue
    revised = target_paragraph.text
    for old, new in REPLACEMENTS.items():
        revised = revised.replace(old, new)
    if revised != target_paragraph.text:
        replace_text(target_paragraph, revised)

target_doc.save(TARGET)
print(TARGET)
