from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from build_responsibility_document import (
    LIGHT_GREEN,
    GREEN,
    add_callout,
    format_table,
    set_cant_split,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.1.docx"


def delete_between(start_paragraph, end_paragraph) -> None:
    element = start_paragraph._p.getnext()
    while element is not None and element is not end_paragraph._p:
        next_element = element.getnext()
        element.getparent().remove(element)
        element = next_element


def move_before(element, anchor_paragraph) -> None:
    anchor_paragraph._p.addprevious(element)


def update_control_table(doc: Document) -> None:
    control = doc.tables[0]
    control.rows[0].cells[3].text = "Draft - V2.1 confirmation matrix"
    control.rows[2].cells[1].text = "0.12 (V2.1 shared-delivery proposals)"
    control.rows[2].cells[3].text = "3 September 2026"


def compact_cover(doc: Document) -> None:
    """Keep all four scope items on the cover so the disclaimer stays on page 2."""
    paragraphs = doc.paragraphs
    for index, size, after, line_spacing in [
        (2, 26.0, 4, 1.0),
        (3, 12.0, 10, 1.05),
        (5, 9.4, 6, 1.05),
        (7, 9.8, 4, 1.08),
        (9, 9.6, 2, 1.05),
        (10, 9.6, 2, 1.05),
        (11, 9.6, 2, 1.05),
        (12, 9.6, 2, 1.05),
    ]:
        paragraph = paragraphs[index]
        paragraph.paragraph_format.space_after = Pt(after)
        paragraph.paragraph_format.line_spacing = line_spacing
        for run in paragraph.runs:
            run.font.size = Pt(size)


def build_v2_1() -> None:
    doc = Document(SOURCE)
    paragraphs = doc.paragraphs
    section_heading = next(p for p in paragraphs if p.text.strip() == "3.4 Shared delivery responsibilities")
    next_heading = next(p for p in paragraphs if p.text.strip() == "4. External interface catalogue")
    delete_between(section_heading, next_heading)

    intro = doc.add_paragraph()
    intro.add_run(
        "The following positions are VCCHUB's proposed implementation defaults. They are written as decisions rather "
        "than open questions so OlyLife can approve them quickly. OlyLife should mark each row Confirmed or return "
        "replacement wording, an owner and a target date."
    )
    intro.paragraph_format.keep_with_next = True
    move_before(intro._p, next_heading)

    headers = ["Decision area", "VCCHUB proposed default", "OlyLife action / response"]
    rows = [
        [
            "1. Member eligibility and response",
            "Only an active and programme-eligible OlyLife member may register. Match a normalized email case-insensitively. "
            "A successful OLY-01 response returns external_member_id, first_name, last_name, matched_email and approved phone fields. "
            "Not-found, inactive or incomplete records create no VCCHUB account.",
            "Confirm the eligible member statuses, required fields and member-support contact; otherwise provide replacement rules and error codes.",
        ],
        [
            "2. API security, identifiers and idempotency",
            "Use server-to-server TLS 1.2+, company_id, UUID request_id, requested_at and HMAC-SHA256 signing over the agreed canonical payload. "
            "Allow a five-minute clock skew. The same request_id returns the original result and never repeats a debit or credit. external_member_id is the cross-system member key.",
            "Confirm HMAC-SHA256, key exchange/rotation, clock-skew window and external_member_id format; nominate OlyLife's API owner.",
        ],
        [
            "3. Sumsub access and data handling",
            "OlyLife owns the Sumsub contract/account and provisions named sandbox and production access to VCCHUB. VCCHUB operates the SDK/server integration. "
            "Sumsub's approved name and DOB are authoritative for cardholder creation. Raw documents/biometrics remain in Sumsub unless an approved exception applies; VCCHUB retains only required verified fields, decision and audit references.",
            "Confirm access approvers, production level, mismatch/manual-review rules, DPA/retention policy and whether any raw evidence may be copied outside Sumsub.",
        ],
        [
            "4. Account-ready mapping",
            "After KYC approval and zero-balance wallet creation, VCCHUB sends a signed wallet_account_ready event containing external_member_id, vcchub_user_id, cardholder_id and wallet_id. "
            "OlyLife stores one active member-to-wallet mapping and uses POST /wallet/status to reconcile uncertain events. Top-up stays disabled until the mapping is unambiguous and ready.",
            "Confirm the webhook URL, event receiver owner, mapping uniqueness rule, retry window and reconciliation schedule.",
        ],
        [
            "5. Top-up approval, debit and reversal",
            "The member submits a request in OlyLife and balances remain unchanged while pending. One authorised Admin/Support reviewer may approve normal requests; maker-checker applies above OlyLife's risk threshold and to manual adjustments/reversals. "
            "Approval atomically rechecks eligibility and commission, debits once, then calls POST /wallet/topup with the same idempotency chain. A permanent credit failure triggers an idempotent commission reversal or controlled repair.",
            "Confirm reviewer roles, maker-checker threshold, per-transaction/daily limits, retry policy, reversal authority and finance-reconciliation owner.",
        ],
        [
            "6. Physical-card delivery",
            "VCCHUB collects no member recipient address. Every V2 physical card is sent to one centrally configured OlyLife office address. Physical issuance is blocked if the address/version is missing or inactive. Address changes require authorised OlyLife approval and an audit trail.",
            "Provide and confirm the legal delivery address, recipient/contact, telephone, operating hours, courier instructions, change approvers and failed-delivery owner.",
        ],
        [
            "7. Support ownership and service levels",
            "OlyLife owns first-line support for member eligibility and commission/top-up approval; VCCHUB owns wallet, card, login/2FA and KYC orchestration, escalating Sumsub/PhotonPay as needed. "
            "Proposed response targets: P1 acknowledgement within 30 minutes (24x7), P2 within four business hours and P3 within one business day. Every handoff includes request_id and the relevant system reference.",
            "Confirm support hours, P1/P2/P3 targets, named operational contacts, escalation channels and member-facing OlyLife support details.",
        ],
        [
            "8. UAT and go-live acceptance",
            "Joint UAT covers matched/not-found/unavailable member checks, address validation, KYC outcomes, ready mapping, approval/rejection/insufficient commission, exactly-once wallet credit, reversal, one card per type, office delivery, card top-up and authenticated cancellation. "
            "Go-live requires all critical/high scenarios passed, no open Severity 1 or 2 defect, production credentials/configuration validated and signed Product, Technology and Operations/Compliance approval from both parties.",
            "Confirm OlyLife's UAT testers and signatories, add required scenarios, agree defect severity definitions and provide target UAT/go-live dates.",
        ],
    ]

    table = doc.add_table(rows=1, cols=3)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = value
    format_table(table, headers, [1900, 5000, 2460], body_size=8.1, first_col_bold=True)
    for row in table.rows:
        set_cant_split(row)
    move_before(table._tbl, next_heading)

    close = add_callout(
        doc,
        "OlyLife confirmation requested",
        "For each row, reply Confirmed or provide amended wording, accountable owner and target date. Silence is not treated as approval. "
        "VCCHUB will incorporate the agreed responses into the final API contracts, delivery backlog and UAT plan.",
        LIGHT_GREEN,
        GREEN,
    )
    move_before(close._p, next_heading)

    update_control_table(doc)
    compact_cover(doc)
    # Keep the 4.7 workflow block on a clean page. Word can otherwise push the
    # heading/table chain into the top-left trim area after the preceding long
    # contract table when exporting to PDF.
    workflow_heading = next(
        p for p in doc.paragraphs if p.text.strip().startswith("4.7 OlyLife internal top-up")
    )
    workflow_heading.paragraph_format.page_break_before = True
    doc.core_properties.title = "OlyLife-VCCHUB Implementation Responsibility Specification - Version 2.1"
    doc.core_properties.subject = "V2 shared-delivery proposals and OlyLife confirmation actions"
    doc.core_properties.author = "Star SaaS Limited"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_v2_1()
