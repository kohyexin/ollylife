from __future__ import annotations

from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT = OUT_DIR / "Ollylife_VCCHUB_Implementation_Responsibility_Specification.docx"
STAR_SAAS_LOGO = ROOT / "assets" / "branding" / "starsaas_logo.png"

# compact_reference_guide preset (resolved tokens)
PAGE_W = Inches(8.5)
PAGE_H = Inches(11)
MARGIN = Inches(1.0)
HEADER_DIST = Inches(0.492)
FOOTER_DIST = Inches(0.492)
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT = "Calibri"
BODY_SIZE = 11
BODY_AFTER = 6
BODY_LINE = 1.25
TABLE_FONT_SIZE = 8.7  # Named override: dense responsibility/API matrices.
SMALL_SIZE = 9.2

INK = "182230"
NAVY = "17365D"
BLUE = "2E74B5"
STAR_BLUE = "0E68BB"
DARK_BLUE = "1F4D78"
PURPLE = "6656E8"
GREEN = "32936F"
LIGHT_BLUE = "E8EEF5"
LIGHT_PURPLE = "F1EFFF"
LIGHT_GREEN = "EAF6F1"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
BORDER = "CBD5E1"
WHITE = "FFFFFF"
RISK = "9B1C1C"
GOLD = "7A5A00"


def set_cell_margins(cell, margins=CELL_MARGINS_DXA):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side in ("top", "start", "bottom", "end"):
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER, size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        el = tcBorders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cant = trPr.find(qn("w:cantSplit"))
    if cant is None:
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_run(run, size=None, color=INK, bold=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_DXA, (widths_dxa, sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(CONTENT_DXA))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        set_cant_split(row)
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_paragraph_border(paragraph, side="left", color=PURPLE, size=18, space=8):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    border = pBdr.find(qn(f"w:{side}"))
    if border is None:
        border = OxmlElement(f"w:{side}")
        pBdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill=LIGHT_PURPLE):
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pPr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instruction
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])
    return run


def add_star_saas_logo(doc, after=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    shape = p.add_run().add_picture(str(STAR_SAAS_LOGO), width=Inches(2.03))
    shape._inline.docPr.set("descr", "Star SaaS logo")
    return p


def configure_star_saas_footer(footer):
    first = footer.paragraphs[0]
    first.clear()
    first.paragraph_format.space_before = Pt(0)
    first.paragraph_format.space_after = Pt(0)
    first.paragraph_format.line_spacing = 1.0
    first.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    set_paragraph_border(first, "top", STAR_BLUE, 10, 3)
    set_run(first.add_run("STAR SAAS LIMITED"), 7.8, INK, True)
    first.add_run("\t")
    set_run(first.add_run("Confidential - Implementation Draft  |  "), 7.5, MID_GRAY)
    set_run(add_field(first, "PAGE"), 7.5, MID_GRAY)

    second = footer.add_paragraph()
    second.paragraph_format.space_before = Pt(0)
    second.paragraph_format.space_after = Pt(0)
    second.paragraph_format.line_spacing = 1.0
    set_run(second.add_run("Flat 1506, 15/F Lucky Center No. 165-171 Wan Chai Road Hong Kong"), 7.3, INK)


def add_star_saas_disclaimer(doc):
    add_star_saas_logo(doc, after=28)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("PROPRIETARY & CONFIDENTIAL"), 18, NAVY, True)

    owner = doc.add_paragraph()
    owner.paragraph_format.space_after = Pt(20)
    set_run(owner.add_run("STAR SAAS LIMITED"), 10.5, STAR_BLUE, True)

    notice = (
        "This document and the information contained in it are owned by and proprietary to Star SaaS Limited "
        "(\"Star SaaS\"). It is provided solely to authorised recipients for evaluation, implementation planning "
        "and delivery of the Ollylife-VCCHUB Wallet Activation & Card Program. The document may not be copied, "
        "distributed, disclosed, published or used for any other purpose without the prior written consent of "
        "Star SaaS. All intellectual property rights in this document, including its structure, analysis, "
        "specifications and associated materials, remain with Star SaaS unless otherwise agreed in writing. "
        "Receipt of this document does not grant any licence or transfer of ownership. If you are not an authorised "
        "recipient, please notify Star SaaS and permanently delete all copies."
    )
    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.space_after = Pt(18)
    body.paragraph_format.line_spacing = 1.3
    set_run(body.add_run(notice), 11, INK)

    add_callout(
        doc,
        "Document control",
        "Owner: Star SaaS Limited. Classification: Confidential. Distribution is limited to authorised Ollylife, VCCHUB and approved delivery stakeholders.",
        LIGHT_BLUE,
        STAR_BLUE,
    )


def add_numbering_definition(doc, bullet=True):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1
    num_id = (max(num_ids) + 1) if num_ids else 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"A11C{abstract_id:04X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"B22D{abstract_id:04X}"[-8:])
    abstract.append(tmpl)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "\uf0b7" if bullet else "%1.")
    lvl.append(lvlText)
    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")
    lvl.append(lvlJc)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    pPr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    pPr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    pPr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    lvl.append(pPr)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    marker_font = "Symbol" if bullet else FONT
    rFonts.set(qn("w:ascii"), marker_font)
    rFonts.set(qn("w:hAnsi"), marker_font)
    if bullet:
        rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), str(num_id))
    numPr.extend([ilvl, numId])


def configure_document(doc):
    sec = doc.sections[0]
    sec.page_width = PAGE_W
    sec.page_height = PAGE_H
    sec.top_margin = MARGIN
    sec.right_margin = MARGIN
    sec.bottom_margin = MARGIN
    sec.left_margin = MARGIN
    sec.header_distance = HEADER_DIST
    sec.footer_distance = FOOTER_DIST
    sec.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER)
    normal.paragraph_format.line_spacing = BODY_LINE

    title = styles["Title"]
    title.font.name = FONT
    title._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    title._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    title.font.size = Pt(29)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_ppr = title._element.get_or_add_pPr()
    title_border = title_ppr.find(qn("w:pBdr"))
    if title_border is not None:
        title_ppr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = FONT
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    subtitle.font.size = Pt(13.5)
    subtitle.font.color.rgb = RGBColor.from_string(MID_GRAY)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(22)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        s = styles[style_name]
        s.font.name = FONT
        s._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        s._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for list_style_name in ("List Bullet", "List Number"):
        ls = styles[list_style_name]
        ls.font.name = FONT
        ls._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        ls._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        ls.font.size = Pt(BODY_SIZE)
        ls.font.color.rgb = RGBColor.from_string(INK)
        ls.paragraph_format.left_indent = Inches(0.375)
        ls.paragraph_format.first_line_indent = Inches(-0.188)
        ls.paragraph_format.space_before = Pt(0)
        ls.paragraph_format.space_after = Pt(4)
        ls.paragraph_format.line_spacing = 1.25

    # Quiet running header with Star SaaS ownership in the footer.
    for header in (sec.header, sec.even_page_header):
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        set_run(hp.add_run("OLLYLIFE × VCCHUB  |  IMPLEMENTATION RESPONSIBILITY"), 8.2, MID_GRAY, True)
    for footer in (sec.footer, sec.even_page_footer):
        configure_star_saas_footer(footer)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(text.upper()), 9.2, PURPLE, True)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    set_keep_with_next(p)
    return p


def add_para(doc, text="", bold_prefix=None, after=None, size=None, color=None, italic=False):
    p = doc.add_paragraph()
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        set_run(p.add_run(bold_prefix), size or BODY_SIZE, color or INK, True)
        set_run(p.add_run(text[len(bold_prefix):]), size or BODY_SIZE, color or INK, None, italic)
    else:
        set_run(p.add_run(text), size or BODY_SIZE, color or INK, None, italic)
    return p


def add_bullet(doc, text, bullet_num_id, lead=None):
    p = doc.add_paragraph()
    apply_num(p, bullet_num_id)
    if lead and text.startswith(lead):
        set_run(p.add_run(lead), BODY_SIZE, INK, True)
        set_run(p.add_run(text[len(lead):]), BODY_SIZE, INK)
    else:
        set_run(p.add_run(text), BODY_SIZE, INK)
    return p


def add_numbered(doc, text, decimal_num_id, lead=None):
    p = doc.add_paragraph()
    apply_num(p, decimal_num_id)
    if lead and text.startswith(lead):
        set_run(p.add_run(lead), BODY_SIZE, INK, True)
        set_run(p.add_run(text[len(lead):]), BODY_SIZE, INK)
    else:
        set_run(p.add_run(text), BODY_SIZE, INK)
    return p


def add_callout(doc, label, text, fill=LIGHT_PURPLE, accent=PURPLE):
    p = doc.add_paragraph()
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, "left", accent, 18, 8)
    set_run(p.add_run(label + "  "), 10.3, accent, True)
    set_run(p.add_run(text), 10.3, INK)
    return p


def format_table(table, headers, widths, body_size=TABLE_FONT_SIZE, header_fill=LIGHT_BLUE,
                 first_col_bold=False, alignments=None):
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(table.rows):
        for cidx, cell in enumerate(row.cells):
            if ridx == 0:
                set_cell_shading(cell, header_fill)
            elif ridx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2.5)
                p.paragraph_format.line_spacing = 1.15
                if alignments:
                    p.alignment = alignments[cidx]
                for run in p.runs:
                    set_run(run, 8.4 if ridx == 0 else body_size,
                            NAVY if ridx == 0 else INK,
                            True if ridx == 0 or (first_col_bold and cidx == 0) else None)
    # Set header text explicitly to avoid inherited styles.
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text), 8.4, NAVY, True)
    return table


def add_table(doc, headers, rows, widths, **kwargs):
    table = doc.add_table(rows=1, cols=len(headers))
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            cells[idx].text = str(value)
    format_table(table, headers, widths, **kwargs)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    return table


def add_owner_block(doc, owner, statement, bullets, bullet_num_id, color, fill):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, "left", color, 18, 8)
    p.paragraph_format.left_indent = Inches(0.14)
    set_run(p.add_run(owner), 12.2, color, True)
    add_para(doc, statement, size=10.5, color=MID_GRAY, after=5)
    for b in bullets:
        add_bullet(doc, b, bullet_num_id)


def build():
    doc = Document()
    configure_document(doc)
    bullet_num_id = add_numbering_definition(doc, bullet=True)
    decimal_num_id = add_numbering_definition(doc, bullet=False)

    # Cover - customer_pack header pattern with Star SaaS ownership.
    add_star_saas_logo(doc, after=10)
    add_kicker(doc, "Partner implementation specification")
    p = doc.add_paragraph(style="Title")
    set_run(p.add_run("Ollylife–VCCHUB Wallet Activation & Card Program"), 29, NAVY, True)
    sp = doc.add_paragraph(style="Subtitle")
    set_run(sp.add_run("Implementation responsibilities, system interfaces, data ownership and delivery acceptance"), 13.5, MID_GRAY)

    meta = doc.add_table(rows=3, cols=4)
    values = [
        ("Prepared for", "Ollylife and VCCHUB delivery teams", "Status", "Draft with live Issuing API mapping"),
        ("Document owner", "Star SaaS Limited", "Classification", "Confidential"),
        ("Version", "0.7", "Date", "2 September 2026"),
    ]
    for r_idx, row_values in enumerate(values):
        for c_idx, value in enumerate(row_values):
            meta.rows[r_idx].cells[c_idx].text = value
    set_table_geometry(meta, [1500, 3300, 1200, 3360])
    for r in meta.rows:
        for c_idx, cell in enumerate(r.cells):
            set_cell_shading(cell, LIGHT_GREEN if c_idx in (0, 2) else WHITE)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                for run in p.runs:
                    set_run(run, 8.8 if c_idx in (0, 2) else 9.3,
                            GREEN if c_idx in (0, 2) else INK,
                            True if c_idx in (0, 2) else None)

    doc.add_paragraph()
    add_callout(doc, "Purpose", "Define who must build, expose, consume, operate and test each component of the journey demonstrated in the current prototype. This is an implementation planning document; final commercial terms, API schemas and regulatory decisions remain subject to joint sign-off.", LIGHT_PURPLE, PURPLE)
    add_heading(doc, "Outcome in one sentence", 2)
    add_para(doc, "Ollylife owns the member entry point and commission ledger; VCCHUB owns wallet onboarding, identity-verification orchestration, wallet/card operations and the APIs/webhooks connecting the two platforms; Sumsub supplies the KYC service used by VCCHUB.")
    add_heading(doc, "Document scope", 2)
    for text in [
        "Wallet invitation, registration, Terms & Conditions, identity verification and account binding.",
        "SSO-style launch from Ollylife, direct VCCHUB login and two-factor authentication.",
        "Commission-funded wallet top-up, cardholder creation, virtual/physical card issuance and physical delivery details.",
        "Wallet-to-card top-up, balances, transaction records, security controls, failure handling and UAT.",
        "Mapping of the currently published Star SaaS Issuing APIs to the journey, including reuse, extension and net-new ownership.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_page_break(doc)
    add_star_saas_disclaimer(doc)

    add_page_break(doc)
    add_heading(doc, "1. Recommended ownership model", 1)
    add_callout(doc, "Core rule", "The party that owns the system of record should expose the authoritative interface. VCCHUB keeps wallet/card APIs internal and provides partner-facing invitation, binding and SSO interfaces. Ollylife provides commission balance and atomic commission withdrawal/debit interfaces consumed by VCCHUB.", LIGHT_GREEN, GREEN)
    rows = [
        ("Member account & commission", "Ollylife", "Logged-in member, external member ID, commission balance and commission debit ledger."),
        ("Wallet onboarding & login", "VCCHUB", "Invitation, registration, Terms acceptance, credentials, 2FA, sessions and wallet portal."),
        ("KYC workflow", "Ollylife + VCCHUB", "Ollylife owns the Sumsub contract/account; VCCHUB implements and operates the integration through access provisioned by Ollylife; Sumsub performs verification."),
        ("Wallet & card ledger", "VCCHUB", "Cardholder, wallet, cards, wallet/card balances, card top-ups and transaction history."),
        ("Account binding", "Joint", "VCCHUB creates the binding and notifies Ollylife; both parties persist their identifiers and current status."),
        ("Card processing / fulfilment", "VCCHUB", "VCCHUB reuses its existing PhotonPay issuer integration and delivery-provider capability; programme fees are configured in VCCHUB before go-live."),
    ]
    add_table(doc, ["Domain", "Accountable party", "Source of truth / duty"], rows, [2200, 1800, 5360], first_col_bold=True)

    add_heading(doc, "1.1 System boundary", 2)
    add_owner_block(doc, "Ollylife", "Owns the member-facing entry and value source.", [
        "Display the logged-in member’s available commission balance and wallet activation/status action.",
        "Confirm the invitation email and call the VCCHUB invitation API from the Ollylife backend.",
        "Provide secure commission balance, debit and reversal APIs consumed by VCCHUB.",
        "Own the Sumsub contract/account and securely provision VCCHUB with the access and credentials required for implementation, testing and production operation.",
        "Store the VCCHUB binding identifiers/status and initiate the logged-in VCCHUB launch flow.",
    ], bullet_num_id, GREEN, LIGHT_GREEN)
    add_owner_block(doc, "VCCHUB", "Owns the wallet product and integration orchestration.", [
        "Provide the invitation, status/binding and SSO/session APIs, plus signed lifecycle webhooks to Ollylife.",
        "Provide registration, Terms acceptance, 2FA, wallet, card, address and transaction user experiences.",
        "Implement, test and operate Sumsub server-to-server and WebSDK flows using access provisioned by Ollylife, then create the cardholder and zero-balance wallet after approval.",
        "Orchestrate commission-funded wallet top-ups and all wallet-to-card/card-issuance operations.",
    ], bullet_num_id, PURPLE, LIGHT_PURPLE)
    add_owner_block(doc, "Sumsub", "Supplies the identity-verification platform under Ollylife’s contract/account, with VCCHUB acting as the delegated technical implementer and operator.", [
        "Capture identity documents and liveness, determine supported document country/type in the verification flow, and perform review.",
        "Return signed events/status and verified applicant data, including first name, last name and date of birth where available.",
        "Provide dashboard controls, level configuration, supported-country/document rules and audit evidence.",
    ], bullet_num_id, BLUE, LIGHT_BLUE)

    add_heading(doc, "2. End-to-end implementation flow", 1)
    flow = [
        ("1", "Ollylife entry", "Ollylife", "Show commission balance and Activate wallet with cards. Confirm default logged-in email or another email."),
        ("2", "Invitation", "Ollylife → VCCHUB", "Ollylife backend calls VCCHUB Invite Wallet User. VCCHUB creates an expiring invite and sends the email."),
        ("3", "Registration", "VCCHUB", "Invite link opens registration. Username equals invitation email; user supplies phone/password and accepts Terms/Privacy."),
        ("4", "Identity verification", "VCCHUB + Sumsub", "VCCHUB creates a server-side access token and embeds/launches Sumsub. Sumsub handles document country/type, capture, liveness and review."),
        ("5", "Cardholder & wallet", "VCCHUB", "On authoritative KYC approval, retrieve verified names/DOB, use registered email/phone, create cardholder and a SGD 0.00 wallet with no card."),
        ("6", "Binding completion", "VCCHUB → Ollylife", "Send ready email and signed webhook. Both platforms store Ollylife member ID ↔ VCCHUB user/cardholder/wallet identifiers."),
        ("7", "Login & 2FA", "VCCHUB", "Ollylife launch uses a short-lived SSO session; direct access uses VCCHUB username/password. Both require 2FA before wallet access."),
        ("8", "Wallet top-up", "VCCHUB ↔ Ollylife", "The user initiates top-up in VCCHUB. VCCHUB shows Ollylife commission availability, then calls Ollylife's atomic commission withdrawal/debit API. On approval, VCCHUB credits the wallet through its internal wallet API and records reconciliation IDs."),
        ("9", "Card creation", "VCCHUB", "Only a funded wallet may create a card. User chooses virtual or physical, provides billing address and, for physical, selects/adds/edits and confirms recipient delivery details."),
        ("10", "Card top-up", "VCCHUB", "From the card action, the user transfers funds from VCCHUB Wallet balance to card balance after policy checks. VCCHUB also deducts applicable card fees from Wallet balance, excluding issuer transaction fees."),
    ]
    add_table(doc, ["#", "Stage", "Lead", "Implementation result"], flow, [480, 1680, 1740, 5460], first_col_bold=True)
    add_callout(doc, "Important control", "The balance display may use a separate enquiry call, but the final commission top-up must use one atomic authorization/debit operation. Do not rely on a prior balance check alone; the available commission may change between calls.", "FFF8E8", GOLD)

    add_heading(doc, "3. Detailed responsibility matrix", 1)
    add_heading(doc, "3.1 Activation, invitation and registration", 2)
    rows = [
        ("Activation entry", "Build button/status UI; show commission balance", "Provide wallet activation status API if live status is required", "N/A", "Ollylife"),
        ("Email confirmation", "Default to logged-in email; allow alternate; validate and record consent", "Accept confirmed email in invitation request", "N/A", "Ollylife"),
        ("Invite API", "Securely consume; send member/correlation/idempotency IDs", "Design, document, secure, host and operate API", "N/A", "VCCHUB"),
        ("Invite email", "Supply approved Ollylife/partner copy if required", "Create expiring token, send email, support resend/expiry", "N/A", "VCCHUB"),
        ("Registration", "No credential handling", "Build page; username=email; phone/password validation; persist account", "N/A", "VCCHUB"),
        ("Terms & Privacy", "Provide Ollylife disclosures/links if applicable", "Capture VCCHUB policy versions, timestamp and consent evidence", "N/A", "VCCHUB / Legal"),
    ]
    add_table(doc, ["Item", "Ollylife", "VCCHUB", "Sumsub", "Accountable"], rows, [1150, 2150, 2600, 1700, 1760], body_size=8.1, first_col_bold=True)

    add_heading(doc, "3.2 KYC, binding and access", 2)
    rows = [
        ("Sumsub setup", "Own contract/account; approve configuration; securely provision VCCHUB access and required credentials", "Implement, test and operate SDK/server integration using provisioned access; secure all credentials", "Provide sandbox/production service", "Ollylife (account) / VCCHUB (integration)"),
        ("Country/document", "Do not ask user to preselect unless required by policy", "Launch generic eligible flow and pass locale/context", "Offer supported country/document selection and validation", "VCCHUB + Sumsub"),
        ("KYC decision", "Receive only status/binding needed", "Use signed webhook/status API as authority; handle retries/review", "Review and return decision/reasons", "VCCHUB"),
        ("Verified data", "Do not receive raw document images by default", "Map approved first/last name, DOB; retain minimum necessary fields", "Extract and return data/evidence", "VCCHUB"),
        ("Cardholder/wallet", "Persist binding reference/status", "Create cardholder and zero-balance wallet; no card", "N/A", "VCCHUB"),
        ("Completion webhook", "Verify signature; process idempotently; update member", "Publish schema, sign, retry and monitor delivery", "N/A", "VCCHUB"),
        ("SSO launch", "Call VCCHUB session API from authenticated backend; show Go to VCCHUB", "Issue short-lived one-time session/URL; validate member binding", "N/A", "VCCHUB"),
        ("Direct login & 2FA", "N/A", "Own password, lockout, recovery, MFA enrollment/challenge and session security", "N/A", "VCCHUB"),
    ]
    add_table(doc, ["Item", "Ollylife", "VCCHUB", "Sumsub", "Accountable"], rows, [1150, 2150, 2600, 1700, 1760], body_size=8.0, first_col_bold=True)

    add_heading(doc, "3.3 Money movement, card creation and servicing", 2)
    rows = [
        ("Commission display", "Expose authoritative available amount/currency/as-of", "Call API and label as Ollylife commission balance", "Ollylife"),
        ("Commission debit", "Provide atomic debit API, idempotency, ledger reference and insufficient-funds response", "Call only after user confirms; never trust browser balance", "Ollylife"),
        ("Wallet credit", "Reconcile Ollylife transaction reference", "Credit wallet exactly once; maintain double-entry/audit trail", "VCCHUB"),
        ("Reversal", "Provide reversal/adjustment control", "Reverse or queue repair if debit succeeds but wallet credit fails", "Joint"),
        ("Card eligibility", "N/A", "Require approved KYC, active cardholder and funded wallet; apply programme limits and a lifetime-issued card quota unless another rule is approved", "VCCHUB"),
        ("Virtual card", "N/A", "Collect billing address; call issuer; store/display masked card and status", "VCCHUB"),
        ("Physical card", "N/A", "Collect billing address and recipient; saved/new/edit address; confirm delivery; fulfilment tracking", "VCCHUB"),
        ("Card top-up", "N/A", "Show and check wallet balance; debit the wallet for the top-up and applicable VCCHUB card fees; credit the card atomically", "VCCHUB"),
        ("Card servicing", "N/A", "Reuse status, cancel and withdrawal functions subject to product policy; cancellation does not restore card quota; any enabled card withdrawal returns funds to the wallet only", "VCCHUB"),
        ("Balances/history", "Show commission debit in Ollylife ledger", "Show wallet/card balances and wallet/card transaction histories", "Each ledger owner"),
    ]
    add_table(doc, ["Item", "Ollylife implementation", "VCCHUB implementation", "Accountable"], rows, [1320, 2880, 3700, 1460], body_size=8.4, first_col_bold=True)

    add_heading(doc, "3.4 Shared delivery responsibilities", 2)
    for text in [
        "Agree API schemas, authentication, environments, IP/network controls, SLAs, rate limits, idempotency and versioning before development freeze.",
        "Define correlation IDs and reconciliation references that appear in both systems’ logs and operational reports.",
        "Complete privacy impact assessment, data-processing terms, consent wording, KYC retention rules and cross-border-data decisions.",
        "Create end-to-end UAT test identities/scenarios and a joint production cutover, support and incident-escalation plan.",
        "Agree ownership for customer support cases: invite/email, KYC review, commission debit, wallet/card ledger, and physical delivery.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "4. External interface catalogue", 1)
    add_para(doc, "This section separates interfaces already published in the Star SaaS Issuing documentation from partner interfaces that must be added for the Ollylife journey. Existing endpoint paths are reproduced for planning only and must be confirmed against the selected API version and production environment before implementation.", italic=True, color=MID_GRAY)
    add_callout(doc, "Mapping conclusion", "Reuse VCCHUB's existing cardholder, wallet and virtual/physical card capabilities internally. Ollylife integrates only the VCCHUB partner-facing invitation, binding/status and SSO interfaces, while VCCHUB consumes new Ollylife commission balance, withdrawal/debit and reversal interfaces. Group APIs are out of scope. Extend VCCHUB so applicable card fees debit the member wallet instead of Merchant balance.", LIGHT_GREEN, GREEN)

    add_heading(doc, "4.1 Journey-to-existing-API coverage", 2)
    coverage = [
        ("Invite / register / bind", "No matching Issuing API", "NEW — VCCHUB", "Build invitation, expiry/resend, activation status and binding lifecycle services."),
        ("Sumsub KYC", "No matching Issuing API", "NEW — JOINT", "Ollylife owns the Sumsub contract/account and provisions access; VCCHUB owns applicant/token/webhook/data retrieval and cardholder mapping."),
        ("Create cardholder", "Create cardholder; Update Cardholder; Query cardholder", "REUSE — VCCHUB INTERNAL", "VCCHUB creates the cardholder internally after authoritative KYC approval using verified first/last name and DOB plus registered email/mobile and billing address. Ollylife does not integrate these APIs."),
        ("Create zero-balance wallet", "Existing VCCHUB wallet creation API", "REUSE — VCCHUB INTERNAL", "VCCHUB creates the SGD 0.00 wallet internally after KYC approval. Ollylife does not call or integrate this API."),
        ("Ollylife commission top-up", "No matching Issuing API", "NEW — OLLYLIFE API", "The member initiates top-up in VCCHUB. VCCHUB calls an Ollylife API that atomically validates and withdraws/debits commission, then credits the VCCHUB wallet internally after approval."),
        ("Create virtual card", "Existing VCCHUB virtual-card function", "REUSE — VCCHUB INTERNAL", "VCCHUB performs issuance internally after funded-wallet and programme checks. Ollylife has no card-creation API integration."),
        ("Create physical card", "Existing VCCHUB physical-card function", "REUSE — VCCHUB INTERNAL", "VCCHUB performs issuance internally, including recipient and delivery-address flow. Ollylife has no card-creation API integration."),
        ("Wallet-to-card top-up", "Existing VCCHUB wallet and card functions", "REUSE + EXTEND — VCCHUB", "VCCHUB checks wallet balance and posts the card top-up internally. Extend fee posting so applicable card fees debit Wallet balance instead of Merchant balance; issuer transaction fees remain issuer/card-side."),
        ("Card servicing", "Update card status; Cancel a card; Withdraw from a card", "REUSE + POLICY CONFIRMATION", "Reuse internally. Confirm whether users may cancel cards and withdraw card balance to wallet. A cancelled card does not free the lifetime card quota; wallet withdrawal back to Ollylife commission is disabled."),
        ("Transactions / balances", "Existing VCCHUB wallet query plus card detail/transaction APIs", "REUSE — VCCHUB INTERNAL", "VCCHUB displays authoritative wallet/card balances and history internally. Ollylife does not integrate these APIs and retains only its own commission ledger."),
        ("Ollylife launch / 2FA", "No matching Issuing API", "NEW — VCCHUB", "Build one-time SSO launch plus VCCHUB authentication and production MFA services."),
        ("Lifecycle webhooks", "General webhook guidance exists; no Issuing binding/card event catalogue evidenced", "NEW / EXTEND — VCCHUB", "Publish signed, versioned wallet/binding/top-up/card events with retries and replay support."),
    ]
    add_table(doc, ["Journey capability", "Published Issuing coverage", "Classification", "Implementation decision"], coverage,
              [1900, 2400, 1700, 3360], body_size=7.8, first_col_bold=True)

    add_heading(doc, "4.2 Applicable VCCHUB wallet and Issuing capability inventory", 2)
    existing_core = [
        ("Cardholder", "Create cardholder", "POST /api/cardholder/create", "Direct reuse after KYC approval; request already contains name, DOB, email, mobile and billing address."),
        ("Cardholder", "Update / Query cardholder", "Published under Issuing", "Reuse for permitted profile correction, status lookup and support. Do not overwrite KYC-approved identity without policy control."),
        ("Cardholder", "Delete cardholder", "Published under Issuing", "Administrative closure only; not part of normal onboarding. Confirm regulatory retention and cards/wallet preconditions."),
        ("Wallet", "Create wallet", "Existing VCCHUB internal API", "Reuse after KYC approval to create a zero-balance member wallet. VCCHUB invokes it internally; no Ollylife integration."),
        ("Wallet", "Query wallet / balance", "Existing VCCHUB internal API", "Reuse as the authoritative source for VCCHUB wallet display, funding eligibility and card-top-up checks."),
        ("Wallet", "Recharge wallet", "Existing VCCHUB internal API", "Reuse internally only after the Ollylife commission withdrawal/debit is approved. Ollylife does not call the wallet API."),
        ("Wallet", "Card-fee debit", "Existing wallet ledger requires extension", "EXTEND so card creation, top-up and servicing fees charged by VCCHUB debit the member Wallet balance instead of Merchant balance. Issuer transaction fees are excluded."),
        ("Wallet", "Withdraw wallet to commission", "Not exposed in this version", "Do not permit wallet withdrawal back to Ollylife commission. Funds credited to the wallet must be spent through the card programme."),
        ("Card", "Query card BIN", "POST /api/card/bin_query", "Reuse internally for available BIN/program metadata and physical/virtual card configuration."),
        ("Card", "Create a card", "POST /api/card/create", "Reuse internally for both virtual and physical issuance after funded-wallet checks. Ollylife does not call this API."),
        ("Card", "Query card details", "POST /api/card/query", "Reuse for card status and card balance. Sensitive PAN/CVV returned by this API require PCI-scoped handling and masking."),
        ("Card", "Query card transaction", "POST /api/card/transaction", "Reuse for card transaction history and reconciliation."),
        ("Card", "Recharge card", "POST /api/card/charge", "Reuse internally for wallet-to-card top-up; first check Wallet balance, then pair request_id with VCCHUB idempotency, wallet debit and fee posting."),
        ("Card", "Query card operation", "POST /api/card/query_operate_result", "Reuse to resolve asynchronous/uncertain create, recharge, withdrawal and cancellation results by order_id."),
        ("Card", "Withdraw / status / cancel", "POST /api/card/withdraw plus published status/cancel APIs", "Reuse internally. If enabled, card withdrawal returns funds to Wallet only. Confirm user access to cancel/withdraw; a cancelled card remains counted against the user card quota."),
        ("Card", "Update shared card limit", "Published under Issuing", "Optional reuse if the selected shared-card programme applies; not required for the base member journey."),
        ("Account", "Query Account Balance", "Published under Issuing", "Merchant/provider operations only. It is not the member Wallet balance and must not be used for the Ollylife member journey."),
    ]
    add_table(doc, ["Domain", "Published API", "Published path / status", "Mapping to Ollylife journey"], existing_core,
              [1050, 1850, 2550, 3910], body_size=7.5, first_col_bold=True)
    add_callout(doc, "Integration boundary", "Ollylife does not consume VCCHUB wallet, cardholder or card APIs. VCCHUB calls these capabilities internally. The partner-facing money movement is VCCHUB → Ollylife commission withdrawal/debit; after approval, VCCHUB posts its own wallet credit.", LIGHT_BLUE, BLUE)
    add_callout(doc, "API-version caveat", "The documentation contains examples from more than one dated API version and some earlier examples use different field casing/signature descriptions. VCCHUB must select one supported version, publish canonical schemas and perform sandbox contract tests before Ollylife integration begins.", "FFF8E8", GOLD)

    add_heading(doc, "4.3 Net-new partner interface catalogue", 2)
    add_para(doc, "The VCCHUB wallet paths below follow the existing Issuing resource/action concept: wallet functions sit under /wallet and use POST. They are logical paths until VCCHUB publishes the final base URL and OpenAPI specification. The Ollylife commission paths remain provisional for Ollylife to confirm or rename.", italic=True, color=MID_GRAY)
    interfaces = [
        ("VCC-01", "Invite wallet user", "VCCHUB", "Ollylife", "POST /wallet/invite", "Create/resend an expiring invitation; return invite ID/status."),
        ("VCC-02", "Activation/binding status", "VCCHUB", "Ollylife", "POST /wallet/status", "Return the current activation, KYC, binding and wallet-ready state."),
        ("VCC-WH-01", "Wallet activation event", "VCCHUB", "Ollylife", "POST {webhook_url}", "Completion callback. event_type is wallet_binding_completed."),
        ("VCC-WH-02", "Lifecycle event", "VCCHUB", "Ollylife", "POST {webhook_url}", "Lifecycle callback. event_type is wallet_binding_updated."),
        ("VCC-03", "SSO launch session", "VCCHUB", "Ollylife", "POST /wallet/session", "Issue a short-lived, single-use VCCHUB launch URL/token."),
        ("OLY-01", "Commission balance", "Ollylife", "VCCHUB", "GET /v1/members/{id}/commission-balance", "Return available commission, currency and as-of time."),
        ("OLY-02", "Commission withdrawal", "Ollylife", "VCCHUB", "POST /v1/commission-withdrawals", "Atomically validate and debit commission for a VCCHUB wallet top-up."),
        ("OLY-03", "Commission reversal", "Ollylife", "VCCHUB", "POST /v1/commission-withdrawals/{id}/reversals", "Reverse a debit if VCCHUB wallet credit cannot complete."),
    ]
    add_table(doc, ["ID", "Interface", "Provider", "Consumer", "Logical endpoint/event", "Purpose"], interfaces,
              [700, 1550, 1000, 1000, 2350, 2760], body_size=8.1, first_col_bold=True)
    add_callout(doc, "VCCHUB wallet API convention", "Use HTTPS POST with application/json, snake_case field names, company_id, a unique request_id and server-generated sign. Follow the Issuing signature process: exclude sign, sort parameters by key, serialize to JSON, append the signkey and hash with SHA-256. Return the standard code, message and result envelope. Use request_id as the idempotency and cross-system correlation key.", LIGHT_GREEN, GREEN)

    add_heading(doc, "4.4 VCC-01 — Invite wallet user (POST /wallet/invite)", 2)
    rows = [
        ("Request body", "company_id, external_member_id, email, locale (optional), return_url (optional and allowlisted), webhook_url (optional and HTTPS), request_id and sign."),
        ("Success response", "code, message and result. result contains request_id, invite_id, status and expires_at. Do not return secrets or reusable login tokens."),
        ("Signing", "Generate sign server-side using the selected VCCHUB Issuing convention. Never expose the signkey in browser code, email, logs or source control."),
        ("Idempotency / lifecycle", "Treat request_id as the idempotency key. The same request_id must return the original outcome. Enforce one active invitation policy, expiry, throttled resend and an audit record."),
        ("Errors", "Use an HTTP status plus code and message; include source where a request field caused the error. Define codes for invalid email, already bound, invite already active, unsupported state, rate limit and service unavailable."),
    ]
    add_table(doc, ["Contract element", "Requirement"], rows, [1800, 7560], body_size=9.2, first_col_bold=True)

    add_heading(doc, "4.5 VCC-WH-01 — Wallet binding completed (provided by VCCHUB)", 2)
    rows = [
        ("Payload", "company_id, request_id, event_id, event_type, event_version, occurred_at, external_member_id, invite_id, vcchub_user_id, cardholder_id, wallet_id, kyc_status, binding_status and sign."),
        ("Delivery", "POST application/json to the HTTPS webhook_url supplied in /wallet/invite; retry with backoff, documented timeout, dead-letter/manual replay and stable event_id."),
        ("Signing", "Sign the webhook payload using the same selected Issuing convention. Exclude sign from the sorted payload; occurred_at supports replay-window validation."),
        ("Ollylife handling", "Verify sign and occurred_at, reject replays, process event_id idempotently, update binding/status, and return a 2xx acknowledgement."),
        ("Privacy", "Do not send document images or unnecessary identity fields. Ollylife needs status and binding identifiers, not the KYC evidence set."),
    ]
    add_table(doc, ["Contract element", "Requirement"], rows, [1800, 7560], body_size=9.2, first_col_bold=True)

    add_heading(doc, "4.6 VCC-03 — Logged-in launch session (POST /wallet/session)", 2)
    rows = [
        ("Precondition", "Authenticated Ollylife member with active binding and permitted wallet status."),
        ("Request body", "company_id, external_member_id, wallet_id or binding_id, return_url (allowlisted), request_id and sign."),
        ("Success response", "code, message and result. result contains request_id, session_url or exchange_token, expires_at and status; no password exposure."),
        ("Signing", "Use the same server-side Issuing sign convention and request_id idempotency rule as /wallet/invite."),
        ("VCCHUB control", "Validate binding, issuer/audience, expiry and nonce; create session only after token exchange; enforce 2FA policy before wallet page."),
        ("Ollylife control", "Call from backend only; never create/assert VCCHUB identity solely from browser parameters; prevent open redirects."),
    ]
    add_table(doc, ["Contract element", "Requirement"], rows, [1800, 7560], body_size=9.2, first_col_bold=True)

    add_heading(doc, "4.7 OLY-01/02/03 — Commission services (provided by Ollylife)", 2)
    rows = [
        ("Balance enquiry", "Member ID → availableAmount, currency, asOf and balanceVersion. Display-only; not an authorization to debit."),
        ("Atomic withdrawal/debit", "withdrawalId/idempotencyKey, member ID, VCCHUB wallet ID, amount, currency and correlation ID → approved/declined/pending, Ollylife transaction ID and new balance."),
        ("Validation", "Positive amount, supported currency, active member, sufficient available commission, limits/risk rules and duplicate-safe processing."),
        ("Reversal", "Original Ollylife transaction ID, amount/reason and idempotency key → reversal status and reference. Restrict to eligible unsettled/failed VCCHUB credit cases."),
        ("Reconciliation", "Both parties store withdrawalId, Ollylife transaction ID, VCCHUB wallet transaction ID, amount, currency, timestamps and final status."),
    ]
    add_table(doc, ["Contract element", "Requirement"], rows, [1800, 7560], body_size=9.2, first_col_bold=True)
    add_callout(doc, "Recommended transaction pattern", "The member initiates top-up in VCCHUB. VCCHUB obtains a current display balance, then submits one idempotent withdrawal request to Ollylife. Ollylife rechecks eligibility and debits commission atomically. VCCHUB credits its wallet exactly once through an internal wallet API. If credit fails after debit, VCCHUB initiates an idempotent reversal or raises a repair item.", LIGHT_GREEN, GREEN)

    add_heading(doc, "4.8 VCCHUB-internal/provider integrations", 2)
    for text in [
        "Sumsub: Ollylife owns the contract/account and securely provisions named access/API credentials; VCCHUB implements access-token creation, WebSDK configuration, signed webhook verification, review-status retrieval and applicant-data retrieval.",
        "VCCHUB wallet and Star SaaS Issuing: use the existing wallet, cardholder and card capabilities behind VCCHUB internal services with canonical field mapping, signing, idempotency, retries and version isolation. Group APIs are out of scope and none of these internal APIs are called by Ollylife.",
        "PhotonPay issuer: VCCHUB reuses its existing virtual/physical issuance and fulfilment integrations; configure wallet-creation, card-creation and top-up fees in VCCHUB before go-live. Issuer transaction fees remain under the existing issuer arrangement.",
        "Notification provider: invitation, resend, account-ready, KYC action-required and card/delivery notifications.",
        "Operations: event replay, manual review, reconciliation, support tooling and audit export.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "4.9 Design references (not Ollylife integration scope)", 2)
    add_callout(doc, "Why this section remains", "Keep a concise reference trail because it explains the conventions used for the new VCCHUB wallet APIs and the basis for classifying VCCHUB cardholder/card functions as internal reuse. It is not an Ollylife implementation checklist and does not require Ollylife to integrate the listed Issuing APIs.", LIGHT_BLUE, BLUE)
    sources = [
        ("Interface / endpoints", "https://docs.star-saas.com/interface and /endpoints", "Server-to-server HTTPS integration and environment conventions."),
        ("Issuing signing", "https://docs.star-saas.com/signing", "snake_case examples, request_id/sign fields, sorted JSON parameters, signkey append and SHA-256 signing."),
        ("Request / response style", "https://docs.star-saas.com/createcardholder", "POST JSON requests and the code, message and result success envelope."),
        ("Internal Issuing reuse", "https://docs.star-saas.com/createcard and /queryoperation", "Card create/recharge/withdraw/cancel operation patterns called only within VCCHUB."),
        ("Webhook guidance", "https://docs.star-saas.com/webhooks", "HTTPS webhook_url request parameter, POST delivery and HTTP acknowledgement convention."),
        ("VCCHUB wallet capability", "Internal VCCHUB specification / implementation", "Wallet create, query, recharge and member-wallet ledger behaviour; canonical internal contracts to be documented by VCCHUB."),
    ]
    add_table(doc, ["Evidence", "Documentation URL", "Used for"], sources, [1600, 3600, 4160], body_size=8.1, first_col_bold=True)

    add_heading(doc, "5. Data ownership and minimum exchange", 1)
    rows = [
        ("Ollylife member ID", "Ollylife", "Ollylife → VCCHUB", "Stable external key for invitation, binding, SSO and commission withdrawal/debit."),
        ("Commission balance/ledger", "Ollylife", "Amount/status only", "VCCHUB must not calculate or persist an authoritative Ollylife balance."),
        ("Confirmed invitation email", "VCCHUB after request; Ollylife records confirmation", "Ollylife → VCCHUB", "Becomes VCCHUB username; define re-invite/email-change policy."),
        ("Registered mobile/password", "VCCHUB", "No password exchange", "Phone supports account/cardholder and MFA; password remains VCCHUB-only."),
        ("KYC evidence/decision", "Sumsub account contracted by Ollylife; operated by VCCHUB", "Status and minimal binding to Ollylife business systems", "Raw images, biometrics and reason codes remain in the approved Sumsub/VCCHUB access and retention model."),
        ("Verified first/last name, DOB", "VCCHUB derived from Sumsub", "Normally VCCHUB/card processor only", "Used to create cardholder; define transliteration/card-name rules."),
        ("Wallet/card balances", "VCCHUB / processor", "Optional status summary to Ollylife", "Authoritative wallet and card ledgers remain VCCHUB-side."),
        ("Billing/delivery address", "VCCHUB / fulfilment provider", "Not sent to Ollylife by default", "Physical recipient may differ from cardholder; retain confirmation evidence."),
        ("Binding map", "Both", "Two-way identifiers/status", "Ollylife member ID ↔ VCCHUB user/cardholder/wallet IDs."),
    ]
    add_table(doc, ["Data", "Source of truth", "Exchange", "Implementation rule"], rows, [1700, 2000, 1800, 3860], body_size=8.4, first_col_bold=True)

    add_heading(doc, "5.1 KYC-to-cardholder mapping", 2)
    rows = [
        ("First name", "Sumsub approved applicant data", "Prefer approved Latin/transliterated field according to card-program rules; never use demo fallback names."),
        ("Last name", "Sumsub approved applicant data", "Preserve legal family name; define compound/multi-part and no-surname handling."),
        ("Date of birth", "Sumsub approved applicant data", "Normalize to ISO date; validate age/program eligibility."),
        ("Email", "VCCHUB registration / invitation", "Use confirmed invitation email; username is the same email."),
        ("Mobile", "VCCHUB registration", "Store normalized E.164 value and verification status."),
        ("Country/document", "Sumsub verification", "Use actual accepted document and country returned by Sumsub; do not trust a browser-only selection."),
    ]
    add_table(doc, ["Cardholder field", "Authoritative source", "Mapping rule"], rows, [1900, 2500, 4960], body_size=9.0, first_col_bold=True)

    add_heading(doc, "5.2 Retention principle", 2)
    add_para(doc, "Each party should retain only the data needed for its contractual, operational and regulatory role. Ollylife owns the Sumsub contract/account but its Ollylife member platform generally requires only binding and lifecycle status, not copies of Sumsub document images or biometric material. Ollylife and VCCHUB must jointly define Sumsub roles, permissions, retention and operational access with legal and compliance teams.")

    add_heading(doc, "6. Security, reliability and operational controls", 1)
    controls = [
        ("API authentication", "Use mutually agreed mTLS, OAuth 2.0 client credentials or signed requests; rotate credentials; separate sandbox and production."),
        ("Webhook security", "HMAC/asymmetric signature over raw body plus timestamp; replay window; event ID idempotency; HTTPS allowlisting as appropriate."),
        ("Secret handling", "Ollylife provisions Sumsub access through approved secure channels; VCCHUB stores delegated API secrets only in server-side secret storage. Never expose credentials in browser code, Git, email or logs; use named access where supported and rotate on role changes."),
        ("Identity/session", "Username=email; strong password policy; account lockout/rate limiting; recovery controls; real MFA in production; short-lived sessions and secure cookies."),
        ("Authorization", "Bind all wallet/card actions to the authenticated VCCHUB subject; never trust member ID, balance, KYC result or amount supplied only by the browser."),
        ("Idempotency", "Required for invitation, webhook processing, commission withdrawal/reversal, wallet credit, card issue, fee posting and card top-up."),
        ("Auditability", "Immutable actor/action/time/request/result records with correlation ID across Ollylife, VCCHUB and providers; redact sensitive fields."),
        ("Money integrity", "Atomic ledger postings, currency precision rules, velocity/amount limits, dual control for manual adjustments and daily reconciliation."),
        ("PII protection", "TLS in transit; encryption at rest; least privilege; purpose-based access; masked card data; retention/deletion/legal-hold processes."),
        ("Monitoring", "Availability, latency, error-rate, webhook backlog, KYC pending age, reconciliation breaks, issuer failures and email delivery alerts."),
    ]
    add_table(doc, ["Control area", "Minimum production requirement"], controls, [1900, 7460], body_size=9.0, first_col_bold=True)
    add_callout(doc, "Demo limitation", "The current demonstration uses Sumsub Sandbox, browser-held demo state and a mock 2FA code (123456). Demo API paths and UI behaviour illustrate the journey but must not be treated as production security, persistence, ledger or regulatory controls.", "FDECEC", RISK)

    add_heading(doc, "6.1 Service-management ownership", 2)
    rows = [
        ("Invite/email issue", "VCCHUB L1/L2; Ollylife confirms submitted address/member context", "VCCHUB email and invite logs"),
        ("KYC pending/rejected", "VCCHUB KYC operations with Sumsub escalation", "Sumsub applicant/review IDs and reason codes"),
        ("Commission mismatch/debit", "Ollylife finance/operations with VCCHUB reconciliation", "Shared withdrawal/correlation IDs"),
        ("Wallet/card balance", "VCCHUB operations / processor", "VCCHUB ledger and processor reference"),
        ("Physical delivery", "VCCHUB / fulfilment provider", "Recipient confirmation, fulfilment and tracking references"),
    ]
    add_table(doc, ["Case", "Lead support owner", "Evidence"], rows, [2200, 4200, 2960], body_size=8.8, first_col_bold=True)

    add_page_break(doc)
    add_heading(doc, "7. Failure and exception handling", 1)
    exceptions = [
        ("Duplicate/already-bound invite", "VCCHUB returns current state; Ollylife shows Go to VCCHUB or resume, not a second account."),
        ("Invite expired or email wrong", "VCCHUB supports controlled resend/reissue; define whether email can change before KYC and how old invite is revoked."),
        ("KYC pending", "VCCHUB shows pending state and accepts webhook/status refresh; do not create cardholder/wallet until authoritative approval."),
        ("KYC rejected/resubmission", "VCCHUB displays permitted user action and routes manual-review/support cases; Ollylife receives only necessary lifecycle status."),
        ("Missing/ambiguous name data", "Do not substitute mock values. Route to exception handling/manual correction permitted by KYC/card-program policy."),
        ("Completion webhook fails", "VCCHUB retries and supports replay; Ollylife can query VCC-02 to reconcile status."),
        ("Commission insufficient", "Ollylife declines atomically without debit; VCCHUB shows current available commission and no wallet credit."),
        ("Debit succeeds, wallet credit fails", "VCCHUB retries exactly-once credit; otherwise triggers reversal/manual repair with shared references."),
        ("Card issue fails", "Post neither the card amount nor VCCHUB card fees, or reverse any reservation atomically; surface the reason safely and allow retry where valid."),
        ("Wallet-to-card top-up fails", "Post neither the card credit nor Wallet debit/fees, or compensate atomically; preserve idempotency and expose transaction status."),
        ("Card cancelled", "Mark the card closed but keep it counted against the user's lifetime-issued card quota; do not automatically restore an issuance slot."),
        ("Card withdrawal", "If user access is approved, return eligible card balance to VCCHUB Wallet only. Do not route card funds or Wallet balance back to Ollylife commission."),
        ("Wallet withdrawal requested", "Reject as unsupported in this version. Wallet funds funded from commission must be spent through the card programme."),
        ("Invalid physical address", "Prevent submission; validate supported delivery country/postcode; require explicit recipient/address confirmation."),
    ]
    add_table(doc, ["Scenario", "Required behaviour"], exceptions, [2500, 6860], body_size=9.0, first_col_bold=True)

    add_heading(doc, "7.1 User-visible status model", 2)
    for text in [
        "Ollylife: Not activated → Invitation sent → Verification in progress → Wallet ready → Action required / unavailable.",
        "VCCHUB onboarding: Invited → Registered → KYC pending → Approved / Rejected / Resubmission required → Wallet created.",
        "Top-up: Initiated → Commission authorized/debited → Wallet credited → Completed, or Reversal pending → Reversed / Manual review.",
        "Card: Not created → Requested → Active / Delivery pending → Blocked / Closed / Failed. Closed/cancelled cards remain counted against the lifetime card quota.",
    ]:
        add_bullet(doc, text, bullet_num_id)

    add_heading(doc, "8. Delivery plan and acceptance", 1)
    phases = [
        ("1. Contract & compliance", "Joint", "Signed responsibility matrix; approved journeys; API/event schemas; auth model; privacy/KYC/card-program decisions; test plan."),
        ("2. Invite & registration", "VCCHUB + Ollylife", "Activation button/email confirmation works; invite API is idempotent; expiring email link; username=email; consent evidence."),
        ("3. Sumsub & binding", "Ollylife + VCCHUB", "Ollylife account/access ready; production-grade VCCHUB token/webhook flow; correct verified name/DOB mapping; zero-balance wallet; signed completion webhook; status query."),
        ("4. Access security", "VCCHUB + Ollylife", "Ollylife launch session and direct login both work; real 2FA, recovery, session and audit controls pass security test."),
        ("5. Commission top-up", "Ollylife + VCCHUB", "Balance enquiry, atomic debit, exactly-once wallet credit, reversal and reconciliation pass concurrency/failure tests."),
        ("6. Card services", "VCCHUB", "Funding gate; virtual/physical selection; billing/recipient address; issue status; wallet-to-card top-up; Wallet-based card fees; servicing and transaction display."),
        ("7. UAT & production", "Joint", "End-to-end happy/negative paths, monitoring, operations, DR, support, cutover and rollback accepted."),
    ]
    add_table(doc, ["Phase", "Lead", "Exit criteria"], phases, [1800, 1900, 5660], body_size=8.8, first_col_bold=True)

    add_heading(doc, "8.1 Minimum end-to-end acceptance scenarios", 2)
    scenarios = [
        "Existing Ollylife member activates using logged-in email; invite received; registration username matches email; Sumsub approves; correct passport/ID name and DOB create cardholder; SGD 0.00 wallet created; binding appears in Ollylife.",
        "Member chooses another email; duplicate/re-invite behaviour is correct; expired link cannot be reused.",
        "Singapore and China document examples complete through the configured program, followed by representative supported global-country/document tests agreed with compliance.",
        "Ollylife launch and direct VCCHUB login both require production 2FA before wallet access; logout/session expiry/recovery work.",
        "Commission top-up succeeds exactly once; insufficient funds declines; repeated request does not double debit; forced VCCHUB failure reverses or raises a controlled repair item.",
        "Wallet starts at zero with no card; card creation is blocked until funded; virtual and physical paths work; physical recipient can use saved, edited or new delivery address.",
        "Card action Top up shows Wallet balance, transfers wallet-to-card exactly once and records both balances/transactions; applicable VCCHUB card fees debit Wallet balance rather than Merchant balance.",
        "Cancelling a card does not restore the user's card quota. If card withdrawal is enabled, funds return only to Wallet; Wallet withdrawal back to Ollylife commission remains unavailable.",
    ]
    for s in scenarios:
        add_numbered(doc, s, decimal_num_id)

    add_heading(doc, "9. Confirmed decisions and remaining build sign-off", 1)
    add_callout(doc, "Confirmed programme decisions", "PhotonPay will remain the card issuer. VCCHUB will configure the applicable wallet-creation, card-creation and top-up fees before go-live. No Ollylife-specific jurisdiction restriction is required; the programme is intended to be open to all users, subject to the standard VCCHUB, issuer, Sumsub and regulatory controls.", LIGHT_GREEN, GREEN)
    decisions = [
        ("KYC", "Sumsub production level(s); verification rules; manual review; transliteration/cardholder-name policy; resubmission; retention and access."),
        ("Identity/access", "SSO protocol and trust model; session expiry; MFA method(s); remembered devices; recovery; account/email change process."),
        ("Integration", "Base URLs; API versioning; auth; certificates/keys; rate limits; webhook signing/retry; idempotency retention; status codes and SLAs."),
        ("Money movement", "Atomic commission withdrawal/debit design; settlement timing; pending states; reversal window; reconciliation files/API; dispute/manual adjustment controls; prohibition on Wallet withdrawal back to commission."),
        ("Card servicing", "Whether cancel and card-to-Wallet withdrawal are user-accessible; cancelled-card quota treatment; eligible withdrawal balance, fees and controls."),
        ("Physical cards", "Recipient rules; address validation; delivery fee; courier; tracking; failed delivery; replacement; allowed alternate recipients."),
        ("Operations", "Support tiers; escalation contacts; monitoring ownership; incident severity/notification; DR/RTO/RPO; audit and reporting cadence."),
    ]
    add_table(doc, ["Decision area", "Items to agree"], decisions, [2000, 7360], body_size=9.0, first_col_bold=True)

    add_heading(doc, "9.1 Immediate next actions", 2)
    add_callout(doc, "Definition of ready", "Development should begin only after the responsible party, interface provider/consumer, authoritative data source, security scheme, idempotency rule, error model and acceptance owner are named for every in-scope integration.", LIGHT_PURPLE, PURPLE)
    actions = [
        ("VCCHUB", "Nominate API owner; select the supported Star SaaS Issuing version; document existing wallet create/query/recharge contracts; configure the PhotonPay programme and Wallet-based fees before go-live; confirm cancelled-card quota and user access to card withdrawal/cancel; provide VCC-01/02/03 and VCC-WH-01 specs; nominate Sumsub technical operators."),
        ("Ollylife", "Nominate backend/frontend and Sumsub account owners; approve the Sumsub level/configuration and securely provision VCCHUB access; provide member/binding data model and draft OLY-01/02/03 commission service contracts."),
        ("Joint", "Run a 90-minute interface workshop to resolve identifiers, auth, idempotency, error/status models, reconciliation and support ownership."),
    ]
    add_table(doc, ["Owner", "Next action"], actions, [1800, 7560], body_size=9.2, first_col_bold=True, header_fill=LIGHT_GREEN)

    add_heading(doc, "9.2 Governance and readiness actions", 2)
    governance_actions = [
        ("Compliance/Legal", "Confirm KYC rules, data processing/retention, Terms/Privacy content and applicable card-program controls."),
        ("QA/Operations", "Convert Section 8 acceptance scenarios into environment-specific test cases, monitoring checks and operational runbooks."),
    ]
    add_table(doc, ["Owner", "Next action"], governance_actions, [1800, 7560], body_size=9.2, first_col_bold=True, header_fill=LIGHT_GREEN)
    add_callout(doc, "Handoff checkpoint", "The first technical workshop should close the highlighted uncertainties: the supported Issuing API version, implementation readiness for Wallet-based fee charging, cancelled-card quota/user access, card-to-Wallet withdrawal policy, and the physical-card fulfilment contract. Once closed, each NEW or EXTEND item in Section 4 can be converted into an owned backlog item and contract test.", LIGHT_BLUE, BLUE)

    # Core properties and update fields on open.
    doc.core_properties.title = "Ollylife–VCCHUB Implementation Responsibility Specification"
    doc.core_properties.subject = "Wallet activation, KYC, commission top-up, card program responsibilities and Star SaaS Issuing API mapping"
    doc.core_properties.author = "Star SaaS Limited"
    doc.core_properties.keywords = "Ollylife, VCCHUB, Sumsub, wallet, card, KYC, API, webhook, RACI, confidential, proprietary"
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    doc.save(OUTPUT)
    print(str(OUTPUT))


if __name__ == "__main__":
    build()
