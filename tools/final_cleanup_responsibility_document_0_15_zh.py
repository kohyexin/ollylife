from pathlib import Path

from docx import Document


target = Path(__file__).resolve().parents[1] / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15_ZH-CN.docx"
doc = Document(target)


def paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from paragraphs(cell)


exact = {
    "card_id, card_type, card_status{\\fn黑体\\fs22\\bord1\\shad0\\3aHBE\\4aH00\\fscx67\\fscy66\\2cHFFFFFF\\3cH808080}被罚refunded_amount, wallet_balance, card_type_slot_released": "card_id、card_type、card_status=CANCELLED、refunded_amount、wallet_balance、card_type_slot_released",
    "交付就绪,wallet充值,制卡,制卡,制卡,制卡等活动正常取消,重复和失序.": "以正常、重复及乱序方式交付账户就绪、钱包充值、卡片创建、卡片充值及卡片取消事件。",
    "提供作为 wallet 用户的签名,呼叫 OlyLife 成员-验证 API,并处理匹配,未找到和无法获取的答复.": "提供钱包用户注册入口，调用 OlyLife 会员验证 API，并处理匹配、未找到及服务不可用的响应。",
    "提供注册、术语接受、地址捕获、Sumsub 核查、直接登录、2FA、wallet、贺卡和交易经历。": "提供账户注册、条款接受、地址收集、Sumsub 验证、直接登录、2FA、钱包、卡片及交易体验。",
    "利用在 API 用户注册期间输入的电子邮件,提供由 VCCHUB 消耗的安全的会员认证 wallet.": "提供安全的会员验证 API，由 VCCHUB 使用用户注册时输入的电子邮箱进行调用。",
    "如果没有活跃成员匹配,则返回一个未找到的明确结果,这样 VCCHUB 就可以提供重试或 OlyLife 支持指导.": "若未匹配到有效会员，应返回明确的未找到结果，使 VCCHUB 可提示用户重试或联系 OlyLife 客服。",
    "安全规定 VCCHUB,与 Sumsub 执行,测试和生产操作所需的访问和证书.": "安全地向 VCCHUB 提供实施、测试及生产运行 Sumsub 所需的访问权限及凭据。",
    "已创建的牌": "卡片已创建",
    "刷卡成功": "卡片充值成功",
    "Wallet 补上成功": "钱包充值成功",
}

replacements = {
    "莱克": "必填",
    "请求栏": "请求字段",
    "情况要求栏": "请求字段",
    "反应领域": "响应字段",
    "现状响应字段": "状态响应字段",
    "世界协调时 3339": "UTC RFC 3339",
    "64 码六分": "64 位十六进制字符串",
    "谜题": "枚举",
    "登机": "开户",
    "刷卡结果": "卡片充值结果",
    "信贷 wallet": "钱包入账",
    "借贷 wallet": "钱包入账",
    "发牌": "发卡",
    "牌资格": "卡片额度",
    "注销": "取消",
}

for paragraph in paragraphs(doc):
    revised = exact.get(paragraph.text, paragraph.text)
    for old, new in replacements.items():
        revised = revised.replace(old, new)
    if paragraph.text == "对":
        revised = "是"
    elif paragraph.text == "总是":
        revised = "始终"
    elif paragraph.text == "有条件的":
        revised = "条件必填"
    if revised != paragraph.text:
        if paragraph.runs:
            paragraph.runs[0].text = revised
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(revised)

doc.save(target)
print(target)
