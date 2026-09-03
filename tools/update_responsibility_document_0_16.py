from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.16.docx"


REPLACEMENTS = {
    "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 0.15": "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 0.16",
    "Purpose  Define who must build, expose, consume, operate and test each component of the programme journey, where an OlyLife member starts in VCCHUB. This 0.15 edition includes the proposed mandatory end-to-end UAT suite, including Korean and Japanese non-Latin-script identity and PhotonPay cardholder-name compatibility tests. The preceding document revision remains retained. Final production configuration and contractual approvals remain subject to joint sign-off.": "Purpose  Define who must build, expose, consume, operate and test each component of the programme journey, where an OlyLife member starts in VCCHUB. This 0.16 edition adds VCCHUB email-OTP verification before OlyLife membership validation, OlyLife Full Name carry-over, and OlyLife-configured programme fees charged atomically to the member Wallet. The preceding document revision remains retained. Final production configuration and contractual approvals remain subject to joint sign-off.",
    "Card creation with one active virtual and one active physical card, wallet-to-card top-up, authenticated cancellation, transaction history and Issuing capability mapping.": "Card creation with one active virtual and one active physical card, OlyLife-configured Wallet fee charging, wallet-to-card top-up, authenticated cancellation, transaction history and Issuing capability mapping.",
    "Core rule  VCCHUB is the wallet onboarding entry point. VCCHUB calls OlyLife for authoritative member eligibility and basic member data, while OlyLife initiates and approves commission-funded wallet top-ups. VCCHUB keeps wallet/card APIs internal except for the partner-facing wallet credit/status interfaces needed by OlyLife.": "Core rule  VCCHUB is the wallet onboarding entry point. VCCHUB verifies control of the entered email before calling OlyLife for authoritative member eligibility and basic member data. OlyLife initiates and approves commission-funded wallet top-ups and owns programme-fee decisions; VCCHUB configures the approved fee schedule and posts triggered fees to the member Wallet. VCCHUB keeps wallet/card APIs internal except for the partner-facing wallet credit/status and lifecycle-event interfaces needed by OlyLife.",
    "Provide a secure member-validation API consumed by VCCHUB using the email entered during wallet-user sign-up.": "Provide a secure member-validation API consumed by VCCHUB only after VCCHUB has verified control of the email entered during wallet-user sign-up.",
    "Return an active-member result with Member ID, First Name, Last Name and the matched email, plus phone data where approved.": "Return an active-member result with Member ID, one Full Name value and the matched email, plus phone data where approved.",
    "Provide Sign up as a wallet user, call the OlyLife member-validation API and handle matched, not-found and unavailable responses.": "Provide Sign up as a wallet user; send and verify a one-time email code before calling the OlyLife member-validation API; and handle invalid/expired OTP, matched, not-found and unavailable responses.",
    "Carry OlyLife Member ID, First Name and Last Name into registration; keep the email as username and collect phone, password and full registered address.": "Carry OlyLife Member ID and Full Name into registration; keep the verified email as username and collect phone, password and full registered address. Sumsub-approved First Name, Last Name and DOB remain authoritative for cardholder creation.",
    "Create the cardholder and zero-balance wallet after Sumsub approval; show member/cardholder/address data read-only during card creation, enforce one active virtual plus one active physical card, show every card as a separate actionable record, use the configured OlyLife office address for physical fulfilment and provide authenticated card cancellation.": "Create the cardholder and zero-balance wallet after Sumsub approval; show member/cardholder/address data read-only during card creation; enforce one active virtual plus one active physical card; calculate and display any OlyLife-configured card-creation fee; atomically debit that fee from Wallet only when issuance succeeds; show every card separately; use the configured OlyLife office address for physical fulfilment; and provide authenticated cancellation.",
    "Mapping conclusion  Reuse VCCHUB’s cardholder, wallet and virtual/physical card capabilities internally. Add an OlyLife member-validation API consumed by VCCHUB and a VCCHUB wallet-credit/status interface consumed by OlyLife after top-up approval. Card cancellation remains internal to VCCHUB. Invitation and SSO APIs from V1 are not required for the programme journey.": "Mapping conclusion  Reuse VCCHUB’s cardholder, wallet and virtual/physical card capabilities internally. Add VCCHUB email-OTP verification before the OlyLife member-validation call, extend VCCHUB fee processing so configured programme fees post to Wallet rather than Merchant balance, and retain the VCCHUB wallet-credit/status interfaces consumed by OlyLife after top-up approval. Card creation, fee posting and cancellation remain internal to VCCHUB. Invitation and SSO APIs are not required for the programme journey.",
    "Integration boundary  VCCHUB calls OlyLife only for member validation during onboarding. OlyLife calls VCCHUB only after an approved commission top-up or when account/wallet status must be reconciled. All cardholder, card and wallet-to-card operations remain internal to VCCHUB.": "Integration boundary  VCCHUB verifies the email internally, then calls OlyLife for member validation. OlyLife calls VCCHUB only after an approved commission top-up or when account/wallet status must be reconciled. Cardholder, card, Wallet fee and wallet-to-card operations remain internal to VCCHUB; OlyLife supplies the approved fee schedule through controlled configuration rather than a member-facing API.",
    "VCCHUB wallet and Star SaaS Issuing: reuse existing wallet, cardholder and card capabilities behind VCCHUB internal services. Carry the registered address into the cardholder/card flow, but do not expose editable address fields on card creation.": "VCCHUB wallet and Star SaaS Issuing: reuse existing wallet, cardholder and card capabilities behind VCCHUB internal services. Extend fee accounting because the current platform posts fees to Merchant balance: programme fees must instead debit the member Wallet, produce a fee-ledger entry and reverse atomically when the triggering operation fails. Carry the registered address into the cardholder/card flow, but do not expose editable address fields on card creation.",
    "Demo limitation  The demonstration uses Sumsub Sandbox, browser-held demo state, a mock 2FA code (123456) for login and card cancellation, simulated OlyLife approval and a configurable placeholder office address. It illustrates the journey but is not a production identity, ledger, approval or fulfilment implementation.": "Demo limitation  The demonstration uses Sumsub Sandbox, browser-held demo state, a mock email OTP (123456), a mock 2FA code (123456) for login and card cancellation, simulated OlyLife approval, a SGD 10.00 sample card-creation fee and a configurable placeholder office address. It illustrates the journey but is not a production email-delivery, identity, ledger, approval, fee-engine or fulfilment implementation.",
    "Card creation shows member/cardholder and registered-address information read-only. The programme permits one active virtual card and one active physical card; both appear as separate rows with independent action menus. Physical creation collects no recipient and uses the configured OlyLife office address.": "Card creation shows member/cardholder and registered-address information read-only. The programme permits one active virtual card and one active physical card; both appear as separate rows with independent action menus. Before confirmation, VCCHUB shows the applicable OlyLife-configured creation fee and resulting Wallet balance. Successful creation posts the fee and card atomically. Physical creation collects no recipient and uses the configured OlyLife office address.",
    "Confirmed programme decisions  PhotonPay remains the issuer. The programme starts in VCCHUB, validates OlyLife membership before registration, collects the member's full address, creates a zero-balance Wallet with no card after KYC, and funds Wallets only through OlyLife's approved top-up process. Each Wallet is assigned a policy allowing one active virtual and one active physical card, with one default BIN per type and no user BIN selection. Policy changes apply prospectively unless a separate existing-wallet migration is approved. Physical cards are sent to the configured OlyLife office. Card cancellation releases the cancelled type slot. OlyLife handles programme/card fee charging directly with PhotonPay; VCCHUB does not deduct such fees from Wallet balance in this release.": "Confirmed programme decisions  PhotonPay remains the issuer. The programme starts in VCCHUB, verifies email control, validates OlyLife membership before registration, collects the member's full address, creates a zero-balance Wallet with no card after KYC, and funds Wallets only through OlyLife's approved top-up process. Each Wallet is assigned a policy allowing one active virtual and one active physical card, with one default BIN per type and no user BIN selection. Policy changes apply prospectively unless a separate existing-wallet migration is approved. Physical cards are sent to the configured OlyLife office. Card cancellation releases the cancelled type slot. Issuer/card-transaction fees remain governed by PhotonPay. Additional programme fees requested by OlyLife are configured in VCCHUB and deducted from Wallet when their trigger succeeds; this requires a VCCHUB extension because current fee postings use Merchant balance.",
    "Definition of ready  Programme development is ready when OLY-01 member validation, account mapping, OlyLife approval/debit, VCCHUB Wallet credit/status, Sumsub access, international-name mapping and PhotonPay compatibility evidence, office-address configuration, authenticated card cancellation and acceptance owners are documented and testable.": "Definition of ready  Programme development is ready when email OTP, OLY-01 member validation, account mapping, OlyLife approval/debit, VCCHUB Wallet credit/status, Wallet fee configuration/posting/reversal, Sumsub access, international-name mapping and PhotonPay compatibility evidence, office-address configuration, authenticated card cancellation and acceptance owners are documented and testable.",
    "Existing OlyLife member signs up in VCCHUB; OLY-01 returns active status, Member ID, First Name, Last Name and matched email; registration continues with those fields shown read-only.": "Existing OlyLife member signs up in VCCHUB; VCCHUB verifies the email OTP before calling OLY-01; OLY-01 returns active status, Member ID, one Full Name value and matched email; registration continues with Member ID and Full Name shown read-only.",
    "Handoff checkpoint  The first programme workshop should close the member-status eligibility rules, OlyLife-to-Sumsub name precedence, account-ready mapping, wallet-credit idempotency/reversal, real OlyLife office address, physical fulfilment controls, default BINs, prospective card-policy versioning and cancellation authentication/audit details. Each NEW or EXTEND item can then become an owned backlog item and contract test.": "Handoff checkpoint  The first programme workshop should close email-OTP controls, member-status eligibility, the OLY-01 Full Name contract, OlyLife-to-Sumsub name precedence, account-ready mapping, Wallet-credit idempotency/reversal, the programme-fee catalogue and approval process, Wallet fee posting/reversal and Merchant-balance exclusion, the real OlyLife office address, physical fulfilment controls, default BINs, prospective card-policy versioning and cancellation authentication/audit details. Each NEW or EXTEND item can then become an owned backlog item and contract test.",
}


TABLE_REPLACEMENTS = {
    "Draft - mandatory UAT suite": "Draft - email OTP and Wallet fee extension",
    "0.15": "0.16",
    "Secure API decision on whether the email belongs to an active eligible member; returns Member ID, First Name and Last Name.": "Secure API decision after VCCHUB email-OTP verification; returns Member ID, Full Name and matched email for an active eligible member.",
    "Cardholder, zero-balance wallet, registration address, cards, balances, top-ups, cancellations and transaction history.": "Cardholder, zero-balance wallet, registration address, cards, balances, programme-fee configuration/postings, top-ups, cancellations and transaction history.",
    "VCCHUB sends the entered email to OLY-01. OlyLife returns active-member status and approved member fields.": "VCCHUB sends an email OTP and verifies control of the address. Only then does VCCHUB call OLY-01; OlyLife returns active-member status and approved member fields.",
    "For a matched member, show Member ID, First Name and Last Name; username=email; collect phone, password, Terms/Privacy and full registered address.": "For a matched member, show Member ID and one Full Name value; username=verified email; collect phone, password, Terms/Privacy and full registered address.",
    "Show cardholder and registered address read-only. Permit one active virtual plus one active physical card and list both separately; physical uses the configured OlyLife office delivery address.": "Show cardholder and registered address read-only. Display the OlyLife-configured card-creation fee and resulting Wallet balance; on success, atomically create the selected card and debit the fee from Wallet. Permit one active virtual plus one active physical card; physical uses the configured OlyLife office address.",
    "Build Sign up as a wallet user and email-entry flow": "Build Sign up as a wallet user, email-OTP send/verify and email-entry flow",
    "Return Member ID, First Name, Last Name, matched email and approved phone data": "Return Member ID, one Full Name value, matched email and approved phone data",
    "Carry values into registration and audit the response": "Carry Member ID and Full Name into registration and audit the response",
    "No credential handling": "No VCCHUB credential or OTP handling",
    "Username=email; collect phone/password and full registered address": "Username=verified email; collect phone/password and full registered address",
    "Require approved KYC, an active cardholder and a funded Wallet; allow one active virtual and one active physical card under the policy version assigned at wallet creation": "Require approved KYC, an active cardholder and a Wallet balance sufficient for the configured fee; allow one active virtual and one active physical card under the policy version assigned at wallet creation",
    "Reuse issuance; show cardholder and registered address read-only": "Reuse issuance; show cardholder/address read-only; preview and atomically debit the configured Wallet fee on success",
    "Check Wallet balance, debit Wallet and credit the selected card atomically. No VCCHUB-configured programme/card fee is deducted from Wallet in this release": "Check Wallet balance for the transfer principal plus any configured fee; atomically debit principal and fee from Wallet, credit only the principal to the selected card and create separate ledger entries",
    "Only an active and programme-eligible OlyLife member may register. Match a normalized email case-insensitively. A successful OLY-01 response returns external_member_id, first_name, last_name, matched_email and approved phone fields. Not-found, inactive or incomplete records create no VCCHUB account.": "VCCHUB first proves email control using a short-lived, single-use OTP with resend, attempt and rate limits. Only then may VCCHUB call OLY-01. An active and programme-eligible match returns external_member_id, full_name, matched_email and approved phone fields. Not-found, inactive or incomplete records create no VCCHUB account.",
    "Joint UAT covers matched/not-found/unavailable member checks, address validation, KYC outcomes, ready mapping, approval/rejection/insufficient commission, exactly-once wallet credit, reversal, one card per type, office delivery, card top-up and authenticated cancellation. Go-live requires all critical/high scenarios passed, no open Severity 1 or 2 defect, production credentials/configuration validated and signed Product, Technology and Operations/Compliance approval from both parties.": "Joint UAT covers email OTP, matched/not-found/unavailable member checks, address validation, KYC outcomes, ready mapping, approval/rejection/insufficient commission, exactly-once wallet credit, programme-fee preview/posting/reversal, one card per type, office delivery, card top-up and authenticated cancellation. Go-live requires all critical/high scenarios passed, no open Severity 1 or 2 defect, production credentials/configuration validated and signed Product, Technology and Operations/Compliance approval from both parties.",
    "OlyLife requires programme/card fees to be charged directly against the created card under the PhotonPay arrangement.": "Issuer/card-transaction fees remain governed by PhotonPay. Additional programme fees requested by OlyLife are charged from the member Wallet when the configured trigger succeeds.",
    "OlyLife agrees the fee schedule and charging treatment directly with PhotonPay. VCCHUB will not configure or deduct these fees from the member Wallet balance in this release. Issuer transaction fees and their card-ledger treatment remain governed by PhotonPay.": "OlyLife owns each programme-fee decision and approves fee code, trigger, amount/rate, currency, tax, effective dates and refundability. VCCHUB implements a controlled fee configuration and extends its fee engine/ledger to debit Wallet rather than Merchant balance. The UI shows the fee and resulting Wallet balance before confirmation. The trigger action and fee posting commit atomically; failure produces no retained fee or an exactly-once reversal. Example: a USD 10.00 card-creation fee deducts USD 10.00 from a USD Wallet when the card is successfully created.",
    "wallet onboarding starts in VCCHUB. OlyLife validates membership by email and returns Member ID, First Name, Last Name and the matched email for an eligible member.": "Wallet onboarding starts in VCCHUB. VCCHUB verifies the entered email with an OTP before calling OlyLife. OlyLife validates membership and returns Member ID, one Full Name value and the matched email for an eligible member.",
    "VCCHUB carries the returned member values into registration, uses the matched email as username and collects mobile, password, consent and full registered address. This release does not use an invitation email or OlyLife-to-VCCHUB SSO launch.": "VCCHUB carries Member ID and Full Name into registration, uses the verified matched email as username and collects mobile, password, consent and full registered address. Sumsub supplies authoritative First Name, Last Name and DOB for cardholder creation. This release does not use an invitation email or OlyLife-to-VCCHUB SSO launch.",
    "Provide OLY-01 for email-based active-member validation and return Member ID, First Name and Last Name.": "Provide OLY-01 for active-member validation after VCCHUB email verification and return Member ID, Full Name and matched email.",
    "Build wallet onboarding, username=email, member fields, Terms, phone/password and registered-address capture.": "Build email-OTP verification, wallet onboarding, username=verified email, member fields, Terms, phone/password and registered-address capture.",
    "Issue after funded-wallet checks only when no active virtual card exists; registered data remains read-only.": "Issue only when no active virtual card exists and Wallet covers the configured creation fee; registered data remains read-only; post fee and issuance atomically.",
    "Issue only when no active physical card exists, using the configured OlyLife office address; no member recipient/address flow.": "Issue only when no active physical card exists and Wallet covers the configured creation fee; use the OlyLife office address; post fee and issuance atomically; collect no member recipient.",
    "Post internally after Wallet-balance checks. No VCCHUB programme/card fee is deducted from Wallet; OlyLife handles the agreed fee arrangement directly with PhotonPay.": "Post internally after checking Wallet covers the card-funding principal plus any OlyLife-configured trigger fee. Debit principal and fee separately from Wallet and credit only principal to the card. This is a VCCHUB extension because current fee charging posts to Merchant balance.",
    "Out of VCCHUB configuration scope": "EXTEND — VCCHUB INTERNAL",
    "OlyLife agrees fee charging directly with PhotonPay. No VCCHUB programme/card fee is configured or deducted from member Wallet balance in this release.": "OlyLife approves the programme-fee schedule. VCCHUB must extend configuration and ledger posting so each triggered fee debits the member Wallet, not Merchant balance, with versioning, idempotency, audit and reversal. PhotonPay issuer/card-transaction fees remain separate.",
    "Reuse internally after funded-Wallet checks. Permit one active card of each type under the wallet's assigned policy version and auto-use the configured default BIN. List cards separately by card ID. Physical fulfilment uses the OlyLife office address. Cancellation releases only the cancelled type slot.": "Reuse internally after confirming Wallet covers the configured creation fee. Permit one active card of each type, auto-use the default BIN, and atomically post the Wallet fee with successful issuance. List cards separately by card ID. Physical fulfilment uses the OlyLife office address. Cancellation releases only the cancelled type slot.",
    "Reuse internally for Wallet-to-card top-up; first check Wallet balance, then pair request_id with VCCHUB idempotency, Wallet debit and selected-card credit.": "Reuse internally for Wallet-to-card top-up; first check Wallet covers principal plus any configured fee, then pair request_id with VCCHUB idempotency, separate Wallet principal/fee debits and selected-card principal credit.",
    "Confirm active/eligible member by email and return approved member fields.": "After VCCHUB email-OTP verification, confirm active/eligible member by email and return approved member fields.",
    "result.first_name / last_name": "result.full_name",
    "Member-reference names shown read-only during registration; KYC-approved names remain authoritative for cardholder creation.": "Single OlyLife member-reference name shown read-only during registration; Sumsub-approved First Name and Last Name remain authoritative for cardholder creation.",
    "Approved principal credited to the Wallet. No VCCHUB programme/card fee is deducted from this amount in this release.": "Approved principal credited to the Wallet. Any programme fee triggered by the top-up is posted as a separate Wallet fee entry and is never netted from the approved principal.",
    "OlyLife First/Last Name": "OlyLife Full Name",
    "Displayed read-only during registration as member-reference data; do not silently overwrite the OlyLife source record.": "Displayed read-only during registration as one member-reference value; do not parse it into cardholder First/Last Name or silently overwrite the OlyLife source record.",
    "First / Last Name at registration": "Full Name at registration",
    "Carry into registration read-only as the matched member profile.": "Carry the single value into registration read-only as the matched member profile; do not assume name-part order.",
    "Required for member validation, account-ready/status processing, top-up approval/debit, Wallet credit, reversal, card issue, card top-up and card cancellation.": "Required for email-OTP verification, member validation, account-ready/status processing, top-up approval/debit, Wallet credit, programme-fee posting/reversal, card issue, card top-up and card cancellation.",
    "Atomic ledger postings, currency precision rules, velocity/amount limits, dual control for manual adjustments and daily reconciliation.": "Atomic principal and fee ledger postings, pre/post balances, currency precision rules, velocity/amount limits, no Merchant-balance leakage, dual control for manual adjustments and daily reconciliation.",
    "Do not continue if required Member ID/name/email is missing; return a support-safe correction path.": "Do not continue if required Member ID/Full Name/matched email is missing; return a support-safe correction path.",
    "Reject a second active card of the same type and show one row with independent actions per card. Auto-use the configured default BIN. If issue/top-up fails, post neither final Wallet debit nor card credit, or reverse reservations atomically.": "Reject a second active card of the same type and show independent actions per card. Auto-use the default BIN. If issue/top-up fails, retain neither the principal nor programme-fee Wallet debit; resolve or reverse every reservation atomically and idempotently.",
    "OLY-01 passes matched/not-found/unavailable tests; Member ID/name carry-over; username=email; phone/password/address and consent evidence.": "Email OTP and OLY-01 pass valid/invalid/expired/rate-limited, matched/not-found/unavailable tests; Member ID/Full Name carry-over; username=verified email; phone/password/address and consent evidence.",
    "Funding gate; one active virtual plus one active physical card; prospective card-policy versioning; default BIN per type; PhotonPay charset/name and embossing compatibility; separate card rows/actions; read-only identity/address; OlyLife-office delivery; Wallet-to-card top-up; authenticated cancellation with Wallet refund and type-slot release; no VCCHUB Wallet-fee deduction.": "Funding and fee gate; OlyLife-controlled fee schedule; Wallet rather than Merchant fee posting; fee preview, atomic debit, idempotency and reversal; one active virtual plus one active physical card; default BIN per type; PhotonPay charset/name compatibility; separate card actions; read-only identity/address; OlyLife-office delivery; Wallet-to-card top-up; authenticated cancellation with Wallet refund/type-slot release.",
    "Enter an email matching one active, eligible OlyLife member and call OLY-01.": "Enter an email matching one active, eligible OlyLife member; verify the VCCHUB email OTP; then call OLY-01.",
    "VCCHUB receives one unambiguous match with Member ID, First Name, Last Name and matched email; registration opens and the request_id is traceable.": "OLY-01 is not called before successful OTP verification. VCCHUB receives one unambiguous match with Member ID, Full Name and matched email; registration opens and both OTP/member request IDs are traceable.",
    "Test unknown, inactive and ineligible emails without disclosing account details.": "Verify email control, then test unknown, inactive and ineligible OlyLife matches without disclosing account details.",
    "Return timeout/503, invalid signature, duplicate records and a matched record missing a required member field.": "Test incorrect, expired, replayed and rate-limited email OTPs; then return OLY-01 timeout/503, invalid signature, duplicate records and a matched record missing a required field.",
    "VCCHUB never bypasses OLY-01. It shows a recoverable or correction-safe message, preserves request_id and creates no account.": "VCCHUB calls neither OLY-01 nor registration before valid OTP. It never bypasses OLY-01, shows a recoverable or correction-safe message, preserves correlation IDs and creates no account.",
    "Member ID and OlyLife name are carried read-only; username equals matched email; phone/address validation and consent evidence persist.": "Member ID and the single OlyLife Full Name are carried read-only; username equals the verified matched email; VCCHUB does not parse that value into KYC name parts; phone/address validation and consent evidence persist.",
    "Execute wallet creation, card creation and funding operations under the agreed PhotonPay fee setup.": "Configure a SGD 10.00 card-creation fee (and representative zero/percentage/future-effective fees), then create cards with sufficient and insufficient Wallet balances. Confirm the equivalent USD 10.00 case in a USD Wallet.",
    "VCCHUB deducts no programme/card fee from Wallet. PhotonPay/card-level charges, if applicable, match the commercial configuration and are separately evidenced.": "The applicable fee and post-action balance are shown before confirmation. On success, VCCHUB debits the fee exactly once from Wallet—not Merchant balance—and records fee code/version, trigger, request/card IDs and pre/post balances. Insufficient balance blocks the action with no posting. PhotonPay issuer/card-transaction fees remain separate.",
    "Attempt virtual and physical card creation while Wallet balance is zero, then repeat after an approved Wallet top-up.": "Attempt virtual and physical card creation with zero balance, with less than the configured fee, and with sufficient Wallet balance; then repeat after an approved Wallet top-up.",
    "Unfunded creation is blocked with no partial reservation. Funded creation proceeds only after authoritative Wallet checks.": "Creation is blocked unless Wallet covers the full configured fee. Funded creation proceeds only after authoritative Wallet checks; successful issuance deducts the fee exactly once and the card starts at zero card balance.",
    "Top up a selected card within Wallet balance, then attempt an amount above available Wallet balance and retry a prior request.": "Top up a selected card where Wallet covers principal plus any configured trigger fee; then test insufficient total balance and retry a prior request.",
    "Successful transfer debits Wallet and credits only the selected card exactly once. Insufficient and duplicate attempts create no incorrect posting.": "Successful transfer debits principal plus fee from Wallet, credits only principal to the selected card and records separate postings exactly once. Insufficient and duplicate attempts create no incorrect posting.",
    "No unsupported active card or unmatched final debit remains. Retry/query/reversal uses stable IDs and reconciles to one final state.": "No unsupported active card, retained programme fee or unmatched Wallet debit remains. Retry/query/reversal uses stable IDs and reconciles card issuance and fee posting to one final state.",
    "Deliver ready, wallet top-up, card-created, card-top-up and card-cancelled events normally, duplicated and out of order.": "Deliver ready, wallet top-up, Wallet-fee-charged, card-created, card-top-up and card-cancelled events normally, duplicated and out of order.",
    "OlyLife commission, VCCHUB Wallet, selected card balances and lifecycle states reconcile with no orphan, duplicate or unexplained amount.": "OlyLife commission, VCCHUB Wallet principal/fee entries, Merchant balance, selected card balances and lifecycle states reconcile with no orphan, duplicate, Merchant-fee leakage or unexplained amount.",
    "Confirmed: one active virtual and one active physical card per Wallet policy version; one default BIN per type; separate card rows/actions; cancellation with fresh 2FA, Wallet refund, permanent Cancelled status and release of only that type slot. Policy changes are prospective; existing-wallet migration is outside this release. PhotonPay must confirm supported cardholder charset/length, transliteration and physical-card embossing rules per configured BIN.": "Confirmed: one active virtual and one active physical card per Wallet policy version; one default BIN per type; separate card rows/actions; cancellation with fresh 2FA, Wallet refund, permanent Cancelled status and release of only that type slot. Policy changes are prospective; existing-wallet migration is outside this release. PhotonPay must confirm supported cardholder charset/length, transliteration and physical-card embossing rules per configured BIN. Programme-fee refundability on later cancellation follows the approved fee configuration and does not alter the released card slot.",
    "Actual OlyLife office address, authorised change process, courier, tracking, failed delivery, replacement and office receiving owner. OlyLife agrees card/programme fee charging directly with PhotonPay; VCCHUB does not deduct those fees from Wallet in this release.": "Actual OlyLife office address, authorised change process, courier, tracking, failed delivery, replacement and office receiving owner. Issuer/card-transaction fees remain under the PhotonPay arrangement; OlyLife-requested additional programme fees follow the separate Wallet-fee control.",
    "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH schemas; implement wallet onboarding, member carry-over, registration address, Sumsub, direct 2FA, zero-balance Wallet creation, read-only card data, international-name mapping, PhotonPay charset validation, prospective card-policy versioning, one default BIN per type, separate per-card actions, OlyLife-office physical fulfilment and authenticated cancellation with Wallet refund/type-slot release.": "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH schemas; implement email OTP, wallet onboarding, Full Name carry-over, registration address, Sumsub, direct 2FA, zero-balance Wallet creation, read-only card data, international-name mapping, PhotonPay charset validation, prospective card-policy versioning, default BINs, separate card actions, OlyLife-office fulfilment, authenticated cancellation and the Wallet-fee configuration/ledger/reversal extension.",
    "Publish OLY-01 with Member ID/First Name/Last Name; provision Sumsub access and authorised Korean/Japanese test data; approve name precedence/transliteration and manual-review policy; implement top-up request/review/debit/reversal; provide the actual office address; and agree all programme/card fee charging directly with PhotonPay.": "Publish OLY-01 with Member ID, Full Name and matched email; provision Sumsub access and authorised Korean/Japanese test data; approve name precedence/transliteration and manual-review policy; implement top-up request/review/debit/reversal; provide the office address; retain issuer/card-transaction fee arrangements with PhotonPay; and approve every additional programme-fee configuration supplied to VCCHUB.",
    "Run an interface/UAT workshop to close identifiers, authentication, status/error models, account mapping, Wallet-credit idempotency, reversal/reconciliation and all UAT-01 to UAT-36 evidence. Obtain PhotonPay confirmation for original/Latin name charset, length and embossing rules; confirm initial policy and default physical/virtual BINs before go-live.": "Run an interface/UAT workshop to close identifiers, email OTP, authentication, status/error models, account mapping, Wallet-credit and Wallet-fee idempotency, fee reversal/reconciliation and all UAT-01 to UAT-36 evidence. Obtain PhotonPay confirmation for original/Latin name charset, length and embossing rules; confirm initial card policy, default BINs and programme-fee schedule before go-live.",
}


def all_paragraphs(doc):
    seen = set()
    for paragraph in doc.paragraphs:
        if paragraph._p not in seen:
            seen.add(paragraph._p)
            yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph._p not in seen:
                        seen.add(paragraph._p)
                        yield paragraph


def set_paragraph_text(paragraph, value):
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def replace_exact(doc):
    found = set()
    for paragraph in all_paragraphs(doc):
        current = paragraph.text
        replacement = REPLACEMENTS.get(current)
        if replacement is None:
            replacement = TABLE_REPLACEMENTS.get(current)
        if replacement is not None:
            set_paragraph_text(paragraph, replacement)
            found.add(current)
    expected = set(REPLACEMENTS) | set(TABLE_REPLACEMENTS)
    missing = sorted(expected - found)
    if missing:
        raise RuntimeError("Expected text not found:\n" + "\n".join(missing))


def set_cell_text(cell, value):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, value)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def append_styled_row(table, values):
    new_tr = deepcopy(table.rows[-1]._tr)
    table._tbl.append(new_tr)
    row = table.rows[-1]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)


def add_new_controls(doc):
    append_styled_row(doc.tables[6], [
        "8. Programme-fee configuration and Wallet charging",
        "OlyLife approves fee code, trigger, amount/rate, currency, tax, effective dates and refundability. VCCHUB versions the configuration and shows the applicable fee before confirmation. Wallet must cover principal plus fee. Trigger action and fee posting commit atomically; an unsuccessful action retains no fee. Default: a later card cancellation does not refund a successfully charged creation fee unless that fee version explicitly says refundable. Cross-currency fee conversion is outside this release; fee currency must equal Wallet currency.",
        "Confirm or amend the initial fee catalogue, amounts, currencies, tax treatment, effective times, refund rules, display wording and authorised configuration approvers. Confirm that same-currency charging and the proposed non-refundable-after-success default are acceptable.",
    ])
    append_styled_row(doc.tables[10], [
        "VCC-WH-06", "Wallet fee charged", "VCCHUB", "OlyLife", "POST {webhook_url}",
        "Synchronize a programme-fee Wallet debit and its trigger/reference without changing Merchant balance.",
    ])
    append_styled_row(doc.tables[21], [
        "wallet.fee.charged", "A configured programme fee commits with its triggering action",
        "fee_transaction_id, fee_code, fee_version, trigger_type, trigger_resource_id, amount, currency, wallet_balance_before, wallet_balance_after",
        "Store the fee result for member history/reconciliation; deduplicate by event_id and never post another debit.",
    ])
    append_styled_row(doc.tables[24], [
        "Programme-fee schedule and fee ledger", "OlyLife (commercial decision) / VCCHUB (configured schedule and ledger)",
        "Approved configuration plus fee status/references", "Version every fee. VCCHUB Wallet ledger is authoritative for debit/reversal; Merchant balance must not be used.",
    ])
    append_styled_row(doc.tables[28], [
        "Programme fee cannot complete", "If Wallet cannot cover the full fee, block the trigger with no posting. If the trigger fails after reservation/posting, release or reverse the fee exactly once. If final status is uncertain, query/reconcile using stable request, trigger and fee-transaction IDs; do not charge Merchant balance.",
    ])
    append_styled_row(doc.tables[36], [
        "Programme fees", "Initial fee catalogue; fee code/trigger; fixed or percentage calculation; Wallet currency; tax; effective dates; member disclosure; refundability (including cancellation); configuration approvers; Merchant-balance exclusion; ledger/event fields; idempotency; reversal and reconciliation. Proposed current-release default is same-currency charging and no refund after a successfully delivered service unless the fee version explicitly says otherwise.",
    ])


def update_json_examples(doc):
    for paragraph in doc.paragraphs:
        if '"first_name": "Olivia"' in paragraph.text and '"last_name": "Chen"' in paragraph.text:
            set_paragraph_text(paragraph, paragraph.text.replace(
                '"first_name": "Olivia",\n    "last_name": "Chen",',
                '"full_name": "Olivia Chen",',
            ))
        if paragraph.text.startswith('{\n  "event_id": "EVT-20260903-000060"'):
            updated = paragraph.text.replace(
                '"card_balance": "0.00",\n    "delivery_destination": "OLYLIFE_OFFICE"',
                '"card_balance": "0.00",\n    "wallet_balance": "490.00",\n    "delivery_destination": "OLYLIFE_OFFICE",\n    "programme_fee": {\n      "fee_transaction_id": "FEE-81000060",\n      "fee_code": "CARD_CREATE",\n      "fee_version": 3,\n      "amount": "10.00",\n      "currency": "SGD",\n      "balance_source": "WALLET"\n    }',
            )
            set_paragraph_text(paragraph, updated)


def add_fee_webhook_example(doc):
    heading_source = next(p for p in doc.paragraphs if p.text == "Card created")
    json_source = next(p for p in doc.paragraphs if p.text.startswith('{\n  "event_id": "EVT-20260903-000060"'))
    heading_xml = deepcopy(heading_source._p)
    json_xml = deepcopy(json_source._p)
    json_source._p.addnext(heading_xml)
    heading_xml.addnext(json_xml)

    heading = Paragraph(heading_xml, heading_source._parent)
    example = Paragraph(json_xml, json_source._parent)
    set_paragraph_text(heading, "Wallet programme fee charged")
    set_paragraph_text(example, '''{
  "event_id": "EVT-20260903-000061",
  "event_type": "wallet.fee.charged",
  "event_version": "1.0",
  "occurred_at": "2026-09-03T03:20:01Z",
  "company_id": 3201,
  "request_id": "card-create-20260903-0060",
  "external_member_id": "OLY-10002345",
  "vcchub_user_id": "USR-50001120",
  "wallet_id": "WLT-90007812",
  "resource_type": "wallet_fee",
  "resource_id": "FEE-81000060",
  "resource_version": 1,
  "status": "COMPLETED",
  "data": {
    "fee_transaction_id": "FEE-81000060",
    "fee_code": "CARD_CREATE",
    "fee_version": 3,
    "trigger_type": "CARD_CREATION",
    "trigger_resource_id": "CRD-30001002",
    "amount": "10.00",
    "currency": "SGD",
    "balance_source": "WALLET",
    "wallet_balance_before": "500.00",
    "wallet_balance_after": "490.00",
    "refundable": false
  },
  "sign": "<64-character-sha256>"
}''')


def main():
    doc = Document(SOURCE)
    replace_exact(doc)
    add_new_controls(doc)
    update_json_examples(doc)
    add_fee_webhook_example(doc)

    # Let these sections flow naturally instead of leaving a mostly blank page
    # when the preceding content has sufficient room.
    for heading_prefix in (
        "4.4 OLY-01",
        "8.2 Proposed mandatory end-to-end UAT test suite",
    ):
        heading = next(p for p in doc.paragraphs if p.text.startswith(heading_prefix))
        heading.paragraph_format.page_break_before = False

    props = doc.core_properties
    props.title = "OlyLife-VCCHUB Wallet Activation & Card Program"
    props.subject = "Implementation Responsibility Specification - Revision 0.16"
    props.author = "Star SaaS Limited"
    props.last_modified_by = "Star SaaS Limited"
    props.comments = "Adds email OTP, OlyLife Full Name and Wallet-funded programme fee controls."
    props.keywords = "OlyLife, VCCHUB, revision 0.16, email OTP, Full Name, Wallet fees, fee reversal, PhotonPay"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
