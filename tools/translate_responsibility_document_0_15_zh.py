from __future__ import annotations

import importlib.util
import os
import re
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15_ZH-CN.docx"

LOCAL_RUNTIME = ROOT / "qa" / "zh_0_15_translation" / "python_packages"
LOCAL_MODELS = ROOT / "qa" / "zh_0_15_translation" / "argos_packages"
sys.path.insert(0, str(LOCAL_RUNTIME))
os.environ["ARGOS_PACKAGES_DIR"] = str(LOCAL_MODELS)
os.environ["ARGOS_CHUNK_TYPE"] = "MINISBD"
os.environ["ARGOS_COMPUTE_TYPE"] = "int8"
os.environ["ARGOS_INTER_THREADS"] = "2"
os.environ["ARGOS_INTRA_THREADS"] = "8"

from argostranslate import translate as argos_translate

INSTALLED_LANGUAGES = argos_translate.get_installed_languages()
ENGLISH = next(language for language in INSTALLED_LANGUAGES if language.code == "en")
CHINESE = next(language for language in INSTALLED_LANGUAGES if language.code == "zh")
LOCAL_TRANSLATOR = ENGLISH.get_translation(CHINESE)

PROTECTED_TERMS = re.compile(
    r"https?://\S+|"
    r"/(?:wallet|v1|api)(?:/[A-Za-z0-9_.{}-]+)+|"
    r"\b(?:OlyLife|VCCHUB|Sumsub|PhotonPay|Star SaaS|OpenAPI|WebSDK|Webhook|Wallet|"
    r"API|KYC|UAT|BIN|MFA|2FA|PII|PAN|CVV|MRZ|HTTPS|HTTP|JSON|SHA-256|UTF-8|"
    r"ISO 4217|SGD|RTO|RPO|SLA|DOB)\b|"
    r"\b(?:OLY|VCC|WH|UAT)-[A-Z0-9/.-]+\b|"
    r"\b[a-z][a-z0-9]*_[a-z0-9_]+\b",
    re.IGNORECASE,
)

TECHNICAL_ONLY = re.compile(
    r"(?:[a-z][a-z0-9_./:+-]*|[A-Z0-9_.:/+<>-]+|\d+(?:\.\d+)?|"
    r"(?:POST|GET|PUT|PATCH|DELETE)\s+/\S+)$"
)


MANUAL_TRANSLATIONS = {
    "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 0.15": "合作伙伴实施规范 · 版本 0.15",
    "VCCHUB-first member validation, onboarding, wallet funding and card services": "以 VCCHUB 为入口的会员验证、开户、钱包入金及卡片服务",
    "Purpose  Define who must build, expose, consume, operate and test each component of the programme journey, where an OlyLife member starts in VCCHUB. This 0.15 edition includes the proposed mandatory end-to-end UAT suite, including Korean and Japanese non-Latin-script identity and PhotonPay cardholder-name compatibility tests. The preceding document revision remains retained. Final production configuration and contractual approvals remain subject to joint sign-off.": "目的  明确当 OlyLife 会员从 VCCHUB 开始项目流程时，各组成部分应由谁建设、提供、接入、运营和测试。本 0.15 版包含建议的强制端到端 UAT 测试套件，包括韩国及日本非拉丁文字身份资料，以及 PhotonPay 持卡人姓名兼容性测试。上一修订版继续保留。最终生产配置及合同批准仍须双方共同签署确认。",
    "VCCHUB owns the member-facing wallet sign-up, registration, Sumsub orchestration and wallet/card experience; OlyLife owns the authoritative member-validation service, Sumsub contract/account, commission ledger and top-up approval workflow; both parties integrate the wallet funding and account-mapping controls.": "VCCHUB 负责面向会员的钱包用户注册、账户注册、Sumsub 编排及钱包/卡片体验；OlyLife 负责权威会员验证服务、Sumsub 合同/账户、佣金账本及充值审批流程；双方共同接入钱包入金及账户映射控制。",
    "VCCHUB-first sign-up, OlyLife membership validation, registration, Terms & Conditions and identity verification.": "从 VCCHUB 发起钱包用户注册、OlyLife 会员验证、账户注册、条款与条件接受及身份验证。",
    "Direct VCCHUB sign-in and two-factor authentication before wallet access; no OlyLife-to-VCCHUB SSO launch is required in this release.": "直接登录 VCCHUB，并在访问钱包前完成双因素认证；本版本不需要从 OlyLife 跳转至 VCCHUB 的 SSO。",
    "OlyLife-initiated commission top-up with Admin/Support approval, plus virtual/physical card issuance.": "由 OlyLife 发起并经管理员/运营人员审批的佣金充值，以及虚拟卡/实体卡发行。",
    "Card creation with one active virtual and one active physical card, wallet-to-card top-up, authenticated cancellation, transaction history and Issuing capability mapping.": "创建卡片（每个钱包限一张有效虚拟卡及一张有效实体卡）、钱包向卡片充值、经认证的卡片取消、交易记录及发卡能力映射。",
    "Core rule  VCCHUB is the wallet onboarding entry point. VCCHUB calls OlyLife for authoritative member eligibility and basic member data, while OlyLife initiates and approves commission-funded wallet top-ups. VCCHUB keeps wallet/card APIs internal except for the partner-facing wallet credit/status interfaces needed by OlyLife.": "核心原则  VCCHUB 是钱包开户入口。VCCHUB 调用 OlyLife 获取权威会员资格及基本会员资料；OlyLife 发起并批准以佣金资金进行的钱包充值。除 OlyLife 所需的合作伙伴钱包入账/状态接口外，VCCHUB 的钱包及卡片 API 均保持内部使用。",
    "Important control  In current release, wallet funding begins in OlyLife, not VCCHUB. OlyLife must approve the request and debit commission atomically before calling VCCHUB to credit the wallet exactly once. The member’s commission and wallet balances must remain unchanged while approval is pending.": "重要控制  本版本的钱包入金从 OlyLife 发起，而非从 VCCHUB 发起。OlyLife 必须先批准请求并以原子方式扣减佣金，之后才可调用 VCCHUB 将资金恰好一次记入钱包。审批待处理期间，会员的佣金余额及钱包余额均不得变动。",
    "The following positions are VCCHUB's proposed implementation defaults. They are written as decisions rather than open questions so OlyLife can approve them quickly. OlyLife should mark each row Confirmed or return replacement wording, an owner and a target date.": "以下内容为 VCCHUB 建议的默认实施方案。为便于 OlyLife 快速审批，内容以拟定结论而非开放问题呈现。OlyLife 应将每一行标记为“已确认”，或提供替代文字、负责人及目标日期。",
    "OlyLife confirmation requested  For each row, reply Confirmed or provide amended wording, accountable owner and target date. Silence is not treated as approval. VCCHUB will incorporate the agreed responses into the final API contracts, delivery backlog and UAT plan.": "请 OlyLife 确认  每一行请回复“已确认”，或提供修订文字、责任人及目标日期。未回复不视为批准。VCCHUB 将把双方同意的答复纳入最终 API 合约、交付待办事项及 UAT 计划。",
    "The following positions are treated as the agreed programme delivery baseline. A later change request may alter a future release, but it does not silently change existing wallets, cards, balances or ledger history.": "以下内容视为双方同意的项目交付基线。后续变更请求可调整未来版本，但不得在未明确迁移的情况下改变现有钱包、卡片、余额或账本历史。",
    "Mapping conclusion  Reuse VCCHUB’s cardholder, wallet and virtual/physical card capabilities internally. Add an OlyLife member-validation API consumed by VCCHUB and a VCCHUB wallet-credit/status interface consumed by OlyLife after top-up approval. Card cancellation remains internal to VCCHUB. Invitation and SSO APIs from V1 are not required for the programme journey.": "映射结论  在 VCCHUB 内部复用现有持卡人、钱包及虚拟卡/实体卡能力。新增由 VCCHUB 调用的 OlyLife 会员验证 API，以及在充值获批后由 OlyLife 调用的 VCCHUB 钱包入账/状态接口。卡片取消仍由 VCCHUB 内部处理。本项目流程不需要 V1 的邀请及 SSO API。",
    "Integration boundary  VCCHUB calls OlyLife only for member validation during onboarding. OlyLife calls VCCHUB only after an approved commission top-up or when account/wallet status must be reconciled. All cardholder, card and wallet-to-card operations remain internal to VCCHUB.": "集成边界  VCCHUB 在开户期间仅为会员验证调用 OlyLife。OlyLife 仅在佣金充值获批后，或需要核对账户/钱包状态时调用 VCCHUB。所有持卡人、卡片及钱包向卡片操作均由 VCCHUB 内部处理。",
    "Partner API convention  Use server-to-server HTTPS, application/json, snake_case fields, company_id, request_id, timestamps and the agreed signing method. Treat request_id as both the idempotency key and cross-system correlation key; never expose credentials or signing keys in browser code.": "合作伙伴 API 规范  使用服务器到服务器的 HTTPS、application/json、snake_case 字段、company_id、request_id、时间戳及双方约定的签名方式。request_id 同时作为幂等键及跨系统关联键；不得在浏览器代码中暴露凭证或签名密钥。",
    "Recommended transaction pattern  The member submits a top-up in OlyLife. OlyLife creates a pending request for Admin/Support review. On approval, OlyLife atomically debits commission and calls VCCHUB /wallet/topup with the wallet ID, amount and idempotency key. VCCHUB credits the wallet exactly once and returns the wallet transaction reference. If credit fails after debit, OlyLife/VCCHUB initiate an idempotent reversal or controlled repair.": "建议交易模式  会员在 OlyLife 提交充值申请。OlyLife 创建待处理请求，供管理员/运营人员审核。批准后，OlyLife 原子化扣减佣金，并使用钱包 ID、金额及幂等键调用 VCCHUB /wallet/topup。VCCHUB 将资金恰好一次记入钱包，并返回钱包交易参考号。如佣金扣减后钱包入账失败，OlyLife/VCCHUB 应发起幂等冲正或受控修复。",
    "Demo limitation  The demonstration uses Sumsub Sandbox, browser-held demo state, a mock 2FA code (123456) for login and card cancellation, simulated OlyLife approval and a configurable placeholder office address. It illustrates the journey but is not a production identity, ledger, approval or fulfilment implementation.": "演示限制  本演示使用 Sumsub Sandbox、浏览器保存的演示状态、用于登录及取消卡片的模拟 2FA 验证码（123456）、模拟 OlyLife 审批及可配置的办公室占位地址。该演示仅用于说明流程，并非生产环境的身份、账本、审批或履约实施。",
    "VCCHUB onboarding: Sign up → OlyLife member check → Registration → KYC pending → Approved / action required → Wallet ready.": "VCCHUB 开户：注册成为钱包用户 → OlyLife 会员验证 → 账户注册 → KYC 待处理 → 已批准 / 需要操作 → 钱包已就绪。",
    "OlyLife top-up: Request submitted → Pending Admin/Support approval → Approved and commission debited → VCCHUB wallet credited → Completed / Reversal pending.": "OlyLife 充值：已提交申请 → 待管理员/运营人员审批 → 已批准并扣减佣金 → VCCHUB 钱包已入账 → 已完成 / 待冲正。",
    "VCCHUB access: Direct sign-in → 2FA challenge → Wallet page. This release does not depend on an OlyLife SSO launch.": "访问 VCCHUB：直接登录 → 2FA 验证 → 钱包页面。本版本不依赖从 OlyLife 发起的 SSO。",
    "Card: Not created → One virtual and/or one physical requested → Active / OlyLife-office delivery pending → Cancelled / Blocked / Failed. Each card appears as its own record. Cancellation returns the eligible card balance to Wallet and releases that card-type slot.": "卡片：未创建 → 已申请一张虚拟卡和/或一张实体卡 → 已激活 / 待配送至 OlyLife 办公室 → 已取消 / 已冻结 / 失败。每张卡片均显示为独立记录。取消卡片后，符合条件的卡片余额退回钱包，并释放该卡片类型的名额。",
}


def iter_paragraphs(doc):
    seen = set()

    def emit(paragraph):
        key = paragraph._p
        if key not in seen:
            seen.add(key)
            yield paragraph

    for paragraph in doc.paragraphs:
        yield from emit(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield from emit(paragraph)
    for section in doc.sections:
        for story in (
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ):
            for paragraph in story.paragraphs:
                yield from emit(paragraph)
            for table in story.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield from emit(paragraph)


def should_translate(text: str) -> bool:
    stripped = text.strip()
    if not stripped or not re.search(r"[A-Za-z]", stripped):
        return False
    if stripped.startswith("{") or "\n  \"" in stripped:
        return False
    if TECHNICAL_ONLY.fullmatch(stripped):
        return False
    if stripped in {"OlyLife", "VCCHUB", "Sumsub", "PhotonPay", "STAR SAAS LIMITED"}:
        return False
    return True


def protect_terms(text: str):
    terms = []

    def replace(match):
        token = f"§{len(terms):04d}§"
        terms.append(match.group(0))
        return token

    return PROTECTED_TERMS.sub(replace, text), terms


def restore_terms(text: str, terms):
    marker_pattern = r"(?:§\d+§|第\d+条§|§\d+条|第§?\d+§?条)"
    found = list(re.finditer(marker_pattern, text))
    if len(found) != len(terms):
        raise RuntimeError(
            f"Technical placeholder mismatch: expected {len(terms)}, found {len(found)} in {text!r}"
        )
    term_iter = iter(terms)
    return re.sub(marker_pattern, lambda _match: next(term_iter), text)


def translate_around_terms(source: str) -> str:
    """Fallback that never exposes protected terms to the language model."""
    output = []
    position = 0
    for match in PROTECTED_TERMS.finditer(source):
        fragment = source[position:match.start()]
        output.append(local_translate(fragment) if re.search(r"[A-Za-z]", fragment) else fragment)
        output.append(match.group(0))
        position = match.end()
    fragment = source[position:]
    output.append(local_translate(fragment) if re.search(r"[A-Za-z]", fragment) else fragment)
    return "".join(output)


def normalise(text: str) -> str:
    text = text.replace("OlyLife", "Ollylife")
    text = text.replace("�", "-").replace("–", "-").replace("—", "-")
    return re.sub(r"\s+", " ", text).strip().casefold()


def load_translation_memory():
    module_path = ROOT / "tools" / "translate_responsibility_document_zh.py"
    spec = importlib.util.spec_from_file_location("legacy_translation", module_path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    old_doc = Document(module.SOURCE)
    seen = set()
    old_unique = [
        p.text
        for p in module.collect_paragraphs(old_doc)
        if p.text.strip() and not (p.text in seen or seen.add(p.text))
    ]
    return {
        normalise(old_unique[index]): value.replace("Ollylife", "OlyLife")
        for index, value in module.T.items()
        if index < len(old_unique)
    }


def local_translate(text: str) -> str:
    return LOCAL_TRANSLATOR.translate(text)


def translate_unique_texts(texts):
    memory = load_translation_memory()
    translated = {}
    work = [text for text in texts if should_translate(text)]
    local_work = []
    for source in work:
        if source in MANUAL_TRANSLATIONS:
            value = MANUAL_TRANSLATIONS[source]
        elif normalise(source) in memory:
            value = memory[normalise(source)]
        else:
            safe_text, terms = protect_terms(source)
            local_work.append((source, safe_text, terms))
            continue
        translated[source] = value

    batches = []
    current = []
    current_size = 0
    for item in local_work:
        projected = current_size + len(item[1]) + 16
        if current and projected > 5000:
            batches.append(current)
            current = []
            current_size = 0
        current.append(item)
        current_size += len(item[1]) + 16
    if current:
        batches.append(current)

    for batch_index, batch in enumerate(batches, start=1):
        joined = []
        for item_index, (_, safe_text, _) in enumerate(batch):
            if item_index:
                joined.append(f"\n🔸{item_index:04d}🔸\n")
            joined.append(safe_text)
        pieces = re.split(r"🔸\d{4}🔸", local_translate("".join(joined)))
        if len(pieces) != len(batch):
            pieces = [local_translate(safe_text) for _, safe_text, _ in batch]
        for (source, _, terms), piece in zip(batch, pieces):
            try:
                translated[source] = restore_terms(piece.strip(), terms)
            except RuntimeError:
                translated[source] = translate_around_terms(source)
        print(f"translated local batch {batch_index}/{len(batches)}", flush=True)

    for source, value in list(translated.items()):
        value = value.replace("Ollylife", "OlyLife")
        value = value.replace("网络钩子", "Webhook").replace("网络挂钩", "Webhook")
        value = value.replace("持卡人名字", "持卡人姓名")
        value = value.replace("顶端", "充值") if "钱包" in value or "卡" in value else value
        translated[source] = value
    return translated


def apply_translation(paragraph, translated):
    nonempty = [run for run in paragraph.runs if run.text]
    if not nonempty:
        paragraph.add_run(translated)
        return

    # Preserve the source's two-run callout label/body treatment.
    if "  " in paragraph.text and len(nonempty) >= 2:
        label_source, body_source = paragraph.text.split("  ", 1)
        if translated.startswith(label_source):
            # Defensive only; normal translation should translate both parts.
            label, body = label_source, translated[len(label_source):].lstrip()
        elif "  " in translated:
            label, body = translated.split("  ", 1)
        else:
            label, body = translated.split(" ", 1) if " " in translated else (translated, "")
        nonempty[0].text = label + ("  " if body else "")
        nonempty[1].text = body
        for run in nonempty[2:]:
            run.text = ""
        return

    nonempty[0].text = translated
    for run in nonempty[1:]:
        run.text = ""


def set_chinese_font(doc):
    east_asia = "Microsoft YaHei"
    for style in doc.styles:
        if style.type == 1:
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), east_asia)
    for paragraph in iter_paragraphs(doc):
        for run in paragraph.runs:
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.rFonts
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:eastAsia"), east_asia)


def adjust_cover(doc):
    for paragraph in doc.paragraphs[:20]:
        if "钱包激活" in paragraph.text and "VCCHUB" in paragraph.text:
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                run.font.size = Pt(21)
        elif "会员验证" in paragraph.text and "VCCHUB" in paragraph.text:
            for run in paragraph.runs:
                run.font.size = Pt(11.5)


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(SOURCE)
    paragraphs = list(iter_paragraphs(doc))
    unique = list(OrderedDict.fromkeys(p.text for p in paragraphs if p.text.strip()))
    translations = translate_unique_texts(unique)

    for paragraph in paragraphs:
        if paragraph.text in translations:
            apply_translation(paragraph, translations[paragraph.text])

    # Ensure the central title is idiomatic even if the service returns a literal phrasing.
    for paragraph in doc.paragraphs[:20]:
        if "OlyLife" in paragraph.text and "VCCHUB" in paragraph.text and "钱包" in paragraph.text:
            paragraph.runs[0].text = "OlyLife-VCCHUB 钱包激活与卡片项目"
            for run in paragraph.runs[1:]:
                run.text = ""
            break

    set_chinese_font(doc)
    adjust_cover(doc)

    props = doc.core_properties
    props.title = "OlyLife-VCCHUB 钱包激活与卡片项目"
    props.subject = "实施责任规范 - 版本 0.15（简体中文）"
    props.author = "Star SaaS Limited"
    props.last_modified_by = "Star SaaS Limited"
    props.comments = "English revision 0.15 translated into Simplified Chinese."
    props.keywords = "OlyLife, VCCHUB, Sumsub, PhotonPay, 钱包, 卡片, KYC, API, Webhook, UAT"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
