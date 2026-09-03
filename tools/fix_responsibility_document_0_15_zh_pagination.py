from pathlib import Path

from docx import Document


target = Path(__file__).resolve().parents[1] / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_0.15_ZH-CN.docx"
doc = Document(target)

for index, paragraph in enumerate(doc.paragraphs):
    if paragraph.text == "1. 建议的责任归属模型" and index and not doc.paragraphs[index - 1].text.strip():
        previous = doc.paragraphs[index - 1]
        previous._element.getparent().remove(previous._element)
        break

doc.save(target)
print(target)
