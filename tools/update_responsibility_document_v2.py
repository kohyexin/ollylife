from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "Ollylife_VCCHUB_Implementation_Responsibility_Specification.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.docx"


def iter_all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.even_page_header, section.first_page_header,
                     section.footer, section.even_page_footer, section.first_page_footer):
            yield from part.paragraphs
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def copy_run_format(run):
    if run is None or run._r.rPr is None:
        return None
    return deepcopy(run._r.rPr)


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def set_paragraph_text(paragraph, text):
    template = copy_run_format(paragraph.runs[0] if paragraph.runs else None)
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if template is not None:
        run._r.insert(0, template)


def set_callout(paragraph, label, body):
    first_format = copy_run_format(paragraph.runs[0] if paragraph.runs else None)
    second_format = copy_run_format(paragraph.runs[1] if len(paragraph.runs) > 1 else None)
    clear_runs(paragraph)
    first = paragraph.add_run(label + "  ")
    second = paragraph.add_run(body)
    if first_format is not None:
        first._r.insert(0, first_format)
    if second_format is not None:
        second._r.insert(0, second_format)
    elif first_format is not None:
        second._r.insert(0, deepcopy(first_format))
        second.bold = False


def set_cell_text(cell, text):
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, str(text))
    for extra in list(cell.paragraphs[1:]):
        cell._tc.remove(extra._p)


def replace_table(table, headers, rows):
    desired = len(rows) + 1
    if desired > len(table.rows):
        raise ValueError(f"Table requires {desired} rows but only {len(table.rows)} are available")
    while len(table.rows) > desired:
        table._tbl.remove(table.rows[-1]._tr)
    for col, value in enumerate(headers):
        set_cell_text(table.rows[0].cells[col], value)
    for row_index, values in enumerate(rows, start=1):
        for col, value in enumerate(values):
            set_cell_text(table.rows[row_index].cells[col], value)


def replace_brand_casing(doc):
    replacements = (("OLLYLIFE", "OLYLIFE"), ("Ollylife", "OlyLife"), ("ollylife", "olyLife"))
    for paragraph in iter_all_paragraphs(doc):
        for run in paragraph.runs:
            for old, new in replacements:
                # Assigning run.text reconstructs the run and removes embedded
                # drawings/page breaks. Only touch runs that contain text that
                # actually needs replacement so the template artwork and
                # pagination controls remain intact.
                if old in run.text:
                    run.text = run.text.replace(old, new)


def build_v2():
    doc = Document(SOURCE)
    replace_brand_casing(doc)
    paragraphs = doc.paragraphs

    paragraph_updates = {
        1: "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 2",
        2: "OlyLife–VCCHUB Wallet Activation & Card Program (Version 2)",
        3: "VCCHUB-first member validation, onboarding, wallet funding and card services",
        5: (
            "Purpose  Define who must build, expose, consume, operate and test each component of the Version 2 journey, "
            "where an OlyLife member starts in VCCHUB. This is a new V2 implementation baseline; the prior English "
            "document remains retained as the V1 reference. Final commercial terms, API schemas and regulatory decisions "
            "remain subject to joint sign-off."
        ),
        7: (
            "VCCHUB owns the member-facing wallet sign-up, registration, Sumsub orchestration and wallet/card experience; "
            "OlyLife owns the authoritative member-validation service, Sumsub contract/account, commission ledger and "
            "top-up approval workflow; both parties integrate the wallet funding and account-mapping controls."
        ),
        9: "VCCHUB-first sign-up, OlyLife membership validation, registration, Terms & Conditions and identity verification.",
        10: "Direct VCCHUB sign-in and two-factor authentication before wallet access; no OlyLife-to-VCCHUB SSO launch is required in V2.",
        11: "OlyLife-initiated commission top-up with Admin/Support approval, plus virtual/physical card issuance.",
        12: "Card creation with one active virtual and one active physical card, wallet-to-card top-up, authenticated cancellation, transaction history and Issuing capability mapping.",
        13: "",
        22: (
            "Core rule  VCCHUB is the V2 onboarding entry point. VCCHUB calls OlyLife for authoritative member eligibility "
            "and basic member data, while OlyLife initiates and approves commission-funded wallet top-ups. VCCHUB keeps "
            "wallet/card APIs internal except for the partner-facing wallet credit/status interfaces needed by OlyLife."
        ),
        26: "Owns the authoritative member identity record, Sumsub contract/account, commission ledger and top-up approval process.",
        27: "Provide a secure member-validation API consumed by VCCHUB using the email entered during wallet-user sign-up.",
        28: "Return an active-member result with Member ID, First Name, Last Name and the matched email, plus phone data where approved.",
        29: "If no active member matches, return a clear not-found result so VCCHUB can offer retry or OlyLife-support guidance.",
        30: "Securely provision VCCHUB with the Sumsub access and credentials required for implementation, testing and production operation.",
        31: "Provide the OlyLife top-up request, Admin/Support approval, atomic commission debit/reversal and VCCHUB wallet-credit orchestration.",
        33: "Owns the V2 wallet-user entry point, onboarding journey and wallet/card product orchestration.",
        34: "Provide Sign up as a wallet user, call the OlyLife member-validation API and handle matched, not-found and unavailable responses.",
        35: "Provide registration, Terms acceptance, address capture, Sumsub verification, direct login, 2FA, wallet, card and transaction experiences.",
        36: "Carry OlyLife Member ID, First Name and Last Name into registration; keep the email as username and collect phone, password and full registered address.",
        37: "Create the cardholder and zero-balance wallet after Sumsub approval; show member/cardholder/address data read-only during card creation, enforce one active virtual plus one active physical card, show every card as a separate actionable record, use the configured OlyLife office address for physical fulfilment and provide authenticated card cancellation.",
        39: "Supplies the identity-verification platform under OlyLife’s contract/account, with VCCHUB acting as the delegated technical implementer and operator.",
        45: (
            "Important control  In V2, wallet funding begins in OlyLife, not VCCHUB. OlyLife must approve the request and "
            "debit commission atomically before calling VCCHUB to credit the wallet exactly once. The member’s commission "
            "and wallet balances must remain unchanged while approval is pending."
        ),
        47: "3.1 VCCHUB entry, OlyLife member validation and registration",
        54: "Agree member-validation, wallet-credit/status and account-ready event schemas, authentication, SLAs, idempotency and versioning before development freeze.",
        55: "Define correlation IDs linking OlyLife member validation, top-up approval/debit and VCCHUB wallet-credit records.",
        56: "Complete privacy assessment, data-processing terms, registration-address purpose/retention and Sumsub access rules.",
        57: "Create V2 UAT scenarios for matched/not-found members, registration/address validation, KYC, approval-based top-up and physical-office delivery.",
        58: "Agree support ownership for member lookup, KYC review, top-up approval/debit, wallet/card ledger and OlyLife-office physical delivery.",
        60: "This section separates existing internal Star SaaS Issuing capabilities from the V2 partner interfaces between VCCHUB and OlyLife. Logical paths remain subject to final OpenAPI and environment sign-off.",
        61: (
            "Mapping conclusion  Reuse VCCHUB’s cardholder, wallet and virtual/physical card capabilities internally. Add an "
            "OlyLife member-validation API consumed by VCCHUB and a VCCHUB wallet-credit/status interface consumed by "
            "OlyLife after top-up approval. Card cancellation remains internal to VCCHUB. Invitation and SSO APIs from V1 are not required for the V2 base journey."
        ),
        66: "Integration boundary  VCCHUB calls OlyLife only for member validation during onboarding. OlyLife calls VCCHUB only after an approved commission top-up or when account/wallet status must be reconciled. All cardholder, card and wallet-to-card operations remain internal to VCCHUB.",
        68: "4.3 V2 partner interface catalogue",
        69: "The paths below are logical V2 contracts until the parties publish final base URLs and OpenAPI definitions. VCCHUB wallet functions remain under /wallet; OlyLife may rename its provisional member and top-up paths during interface sign-off.",
        71: "Partner API convention  Use server-to-server HTTPS, application/json, snake_case fields, company_id, request_id, timestamps and the agreed signing method. Treat request_id as both the idempotency key and cross-system correlation key; never expose credentials or signing keys in browser code.",
        72: "4.4 OLY-01 — Validate OlyLife member (POST /v1/members/validate)",
        74: "4.5 VCC-01 — Credit member wallet (POST /wallet/topup)",
        76: "4.6 VCC-WH-01 / VCC-02 — Account-ready event and wallet status",
        78: "4.7 OlyLife internal top-up request and approval workflow",
        80: "Recommended transaction pattern  The member submits a top-up in OlyLife. OlyLife creates a pending request for Admin/Support review. On approval, OlyLife atomically debits commission and calls VCCHUB /wallet/topup with the wallet ID, amount and idempotency key. VCCHUB credits the wallet exactly once and returns the wallet transaction reference. If credit fails after debit, OlyLife/VCCHUB initiate an idempotent reversal or controlled repair.",
        82: "Sumsub: OlyLife owns the contract/account and securely provisions named access/API credentials; VCCHUB implements token creation, WebSDK configuration, signed webhook verification, review-status retrieval and approved applicant-data retrieval.",
        83: "VCCHUB wallet and Star SaaS Issuing: reuse existing wallet, cardholder and card capabilities behind VCCHUB internal services. Carry the registered address into the cardholder/card flow, but do not expose editable address fields on V2 card creation.",
        84: "PhotonPay issuer: VCCHUB reuses existing virtual/physical issuance and cancellation capabilities. For a V2 physical card, VCCHUB uses the configured OlyLife office delivery address and does not collect a member recipient address.",
        85: "Notification provider: KYC action-required, account-ready, top-up outcome and card/delivery notifications. V2 does not require an invitation email.",
        88: "Why this section remains  Keep a concise reference trail for the signing, envelope and internal Issuing conventions used by the V2 member-validation, wallet-credit, card top-up and card-cancellation operations. It is not an OlyLife checklist for VCCHUB-internal card APIs.",
        95: "Each party should retain only the data required for its contractual, operational and regulatory role. OlyLife owns the Sumsub contract/account and authoritative member record; VCCHUB retains the minimum verified identity, registration address and wallet/card data required to operate the programme. Raw Sumsub document images and biometrics remain subject to the approved Sumsub/VCCHUB access and retention model.",
        98: "Demo limitation  The Version 2 demonstration uses Sumsub Sandbox, browser-held demo state, a mock 2FA code (123456) for login and card cancellation, simulated OlyLife approval and a configurable placeholder office address. It illustrates the journey but is not a production identity, ledger, approval or fulfilment implementation.",
        105: "VCCHUB onboarding: Sign up → OlyLife member check → Registration → KYC pending → Approved / action required → Wallet ready.",
        106: "OlyLife top-up: Request submitted → Pending Admin/Support approval → Approved and commission debited → VCCHUB wallet credited → Completed / Reversal pending.",
        107: "VCCHUB access: Direct sign-in → 2FA challenge → Wallet page. V2 does not depend on an OlyLife SSO launch.",
        108: "Card: Not created → One virtual and/or one physical requested → Active / OlyLife-office delivery pending → Cancelled / Blocked / Failed. Each card appears as its own record. Cancellation returns the eligible card balance to Wallet and releases that card-type slot.",
        112: "Existing OlyLife member signs up in VCCHUB; OLY-01 returns active status, Member ID, First Name, Last Name and matched email; registration continues with those fields shown read-only.",
        113: "Unknown/inactive member receives a not-found response and can retry the email or contact OlyLife support; no VCCHUB account is created.",
        114: "Member completes phone, password, Terms/Privacy and full registered address. Username equals the verified email, and the address persists into the wallet/card profile.",
        115: "Sumsub approves a representative supported document; the authoritative approved name and DOB create the cardholder and a zero-balance wallet with no card.",
        116: "Direct VCCHUB login requires production 2FA before wallet access; recovery, rate limiting, logout and session expiry work.",
        117: "OlyLife top-up remains pending with no balance change until Admin/Support approval; approval debits commission and credits VCCHUB exactly once; rejection changes neither balance.",
        118: "Card creation shows member/cardholder and registered-address information read-only. The programme permits one active virtual card and one active physical card; both appear as separate rows with independent action menus. Physical creation collects no recipient and uses the configured OlyLife office address.",
        119: "Wallet-to-card top-up shows Wallet balance, transfers exactly once and records balances/transactions against the selected card. Card cancellation requires fresh 2FA, is irreversible, returns the eligible card balance to Wallet, records the event and releases only the cancelled card-type slot so that type can be replaced.",
        121: "Confirmed programme decisions  PhotonPay remains the issuer. V2 starts in VCCHUB, validates OlyLife membership before registration, collects the member’s full address, and funds wallets only through OlyLife’s approved top-up process. A member may hold one active virtual card and one active physical card, displayed as separate actionable records. Physical cards are sent to the configured OlyLife office address without collecting a member recipient address. Card cancellation is user-accessible in VCCHUB with fresh 2FA; eligible card balance returns to Wallet and only the cancelled card-type slot is released.",
        124: "Definition of ready  V2 development is ready when OLY-01 member validation, account mapping, OlyLife approval/debit, VCCHUB wallet credit/status, Sumsub access, office-address configuration, authenticated card cancellation and acceptance owners are documented and testable.",
        128: "Handoff checkpoint  The first V2 workshop should close the member-status eligibility rules, OlyLife-to-Sumsub name precedence, account-ready mapping, wallet-credit idempotency/reversal, real OlyLife office address, physical fulfilment controls and cancellation authentication/audit details. Each NEW or EXTEND item can then become an owned backlog item and contract test.",
    }
    for index, text in paragraph_updates.items():
        set_paragraph_text(paragraphs[index], text)

    set_callout(paragraphs[5], "Purpose", paragraph_updates[5].split("Purpose  ", 1)[1])
    set_callout(paragraphs[22], "Core rule", paragraph_updates[22].split("Core rule  ", 1)[1])
    set_callout(paragraphs[45], "Important control", paragraph_updates[45].split("Important control  ", 1)[1])
    set_callout(paragraphs[61], "Mapping conclusion", paragraph_updates[61].split("Mapping conclusion  ", 1)[1])
    set_callout(paragraphs[71], "Partner API convention", paragraph_updates[71].split("Partner API convention  ", 1)[1])
    set_callout(paragraphs[88], "Why this section remains", paragraph_updates[88].split("Why this section remains  ", 1)[1])
    set_callout(paragraphs[98], "Demo limitation", paragraph_updates[98].split("Demo limitation  ", 1)[1])
    set_callout(paragraphs[121], "Confirmed programme decisions", paragraph_updates[121].split("Confirmed programme decisions  ", 1)[1])
    set_callout(paragraphs[124], "Definition of ready", paragraph_updates[124].split("Definition of ready  ", 1)[1])
    set_callout(paragraphs[128], "Handoff checkpoint", paragraph_updates[128].split("Handoff checkpoint  ", 1)[1])

    tables = doc.tables
    replace_table(tables[0], ["Prepared for", "", "Status", ""], [
        ["Prepared for", "OlyLife and VCCHUB delivery teams", "Status", "Draft – V2 VCCHUB-first flow"],
        ["Document owner", "Star SaaS Limited", "Classification", "Confidential"],
        ["Version", "0.11 (V2 card-type limits)", "Date", "2 September 2026"],
    ][1:])
    # Restore the control table's label/value arrangement after applying the generic row helper.
    control_values = [
        ["Prepared for", "OlyLife and VCCHUB delivery teams", "Status", "Draft – V2 VCCHUB-first flow"],
        ["Document owner", "Star SaaS Limited", "Classification", "Confidential"],
        ["Version", "0.11 (V2 card-type limits)", "Date", "2 September 2026"],
    ]
    for r_index, values in enumerate(control_values):
        for c_index, value in enumerate(values):
            set_cell_text(tables[0].rows[r_index].cells[c_index], value)

    replace_table(tables[1], ["Domain", "Accountable party", "Source of truth / duty"], [
        ["Member identity & commission", "OlyLife", "Authoritative member status/profile, external Member ID, commission balance and commission debit ledger."],
        ["Membership validation", "OlyLife", "Secure API decision on whether the email belongs to an active eligible member; returns Member ID, First Name and Last Name."],
        ["Wallet onboarding & login", "VCCHUB", "Sign-up entry, registration, Terms acceptance, credentials, 2FA, sessions and wallet portal."],
        ["KYC workflow", "OlyLife + VCCHUB", "OlyLife owns the Sumsub contract/account; VCCHUB implements and operates the integration; Sumsub performs verification."],
        ["Wallet, card & registered address", "VCCHUB", "Cardholder, zero-balance wallet, registration address, cards, balances, top-ups, cancellations and transaction history."],
        ["Top-up approval & funding", "Joint", "OlyLife owns request approval and commission debit; VCCHUB owns exactly-once wallet credit and wallet ledger."],
    ])

    replace_table(tables[2], ["#", "Stage", "Lead", "Implementation result"], [
        ["1", "VCCHUB entry", "VCCHUB", "User opens VCCHUB and selects Sign up as a wallet user."],
        ["2", "OlyLife member check", "VCCHUB → OlyLife", "VCCHUB sends the entered email to OLY-01. OlyLife returns active-member status and approved member fields."],
        ["3", "Registration", "VCCHUB", "For a matched member, show Member ID, First Name and Last Name; username=email; collect phone, password, Terms/Privacy and full registered address."],
        ["4", "Identity verification", "VCCHUB + Sumsub", "VCCHUB launches Sumsub. Sumsub handles document country/type, capture, liveness and review."],
        ["5", "Cardholder & wallet", "VCCHUB", "On KYC approval, retrieve approved identity data, create cardholder and a SGD 0.00 wallet with no card."],
        ["6", "Ready, sign-in & 2FA", "VCCHUB", "Show wallet ready, return to VCCHUB sign-in and require 2FA before wallet access."],
        ["7", "Top-up request", "OlyLife", "Member requests a wallet top-up from the OlyLife page; balances remain unchanged while pending."],
        ["8", "Approval & wallet credit", "OlyLife → VCCHUB", "Admin/Support approves, OlyLife atomically debits commission and calls VCCHUB to credit the wallet exactly once."],
        ["9", "Card creation", "VCCHUB", "Show cardholder and registered address read-only. Permit one active virtual plus one active physical card and list both separately; physical uses the configured OlyLife office delivery address."],
        ["10", "Card top-up & cancellation", "VCCHUB", "Apply each action to the selected card. Cancellation requires fresh 2FA, refunds eligible card balance to Wallet and releases only that card-type slot."],
    ])

    replace_table(tables[3], ["Item", "OlyLife", "VCCHUB", "Sumsub", "Accountable"], [
        ["VCCHUB sign-up entry", "No onboarding UI required", "Build Sign up as a wallet user and email-entry flow", "N/A", "VCCHUB"],
        ["Member-validation API", "Design, secure, host and operate OLY-01", "Consume server-to-server with correlation/idempotency controls", "N/A", "OlyLife"],
        ["Member not found", "Provide support contact and eligibility rules", "Show retry and Contact OlyLife support; create no account", "N/A", "VCCHUB / OlyLife Support"],
        ["Matched member data", "Return Member ID, First Name, Last Name, matched email and approved phone data", "Carry values into registration and audit the response", "N/A", "OlyLife"],
        ["Registration", "No credential handling", "Username=email; collect phone/password and full registered address", "N/A", "VCCHUB"],
        ["Terms & Privacy", "Provide OlyLife disclosures/links if applicable", "Capture VCCHUB policy versions, timestamp and consent evidence", "N/A", "VCCHUB / Legal"],
    ])

    replace_table(tables[4], ["Item", "OlyLife", "VCCHUB", "Sumsub", "Accountable"], [
        ["Sumsub setup", "Own contract/account; approve configuration; securely provision VCCHUB access", "Implement, test and operate SDK/server integration; secure credentials", "Provide sandbox/production service", "OlyLife (account) / VCCHUB (integration)"],
        ["Country/document", "Do not ask user to preselect unless required by policy", "Launch generic eligible flow and pass locale/context", "Offer supported country/document selection and validation", "VCCHUB + Sumsub"],
        ["KYC decision", "Receive only status/mapping needed", "Use signed webhook/status API as authority; handle retries/review", "Review and return decision/reasons", "VCCHUB"],
        ["Verified data", "OlyLife member names remain registration reference data", "Use Sumsub-approved first/last name and DOB as authoritative cardholder identity", "Extract and return approved data/evidence", "VCCHUB"],
        ["Registered address", "No address collection in OlyLife for V2 onboarding", "Collect full address at registration and carry it into cardholder/card views", "N/A", "VCCHUB"],
        ["Cardholder/wallet", "Retain mapping/status needed for top-up eligibility", "Create cardholder and zero-balance wallet; no card", "N/A", "VCCHUB"],
        ["Account ready", "Process signed ready event or reconcile through status before allowing top-up", "Show ready page; publish event/status without exposing integration details in UI", "N/A", "Joint"],
        ["Direct login & 2FA", "N/A", "Own password, lockout, recovery, MFA challenge and session security", "N/A", "VCCHUB"],
    ])

    replace_table(tables[5], ["Item", "OlyLife implementation", "VCCHUB implementation", "Accountable"], [
        ["Commission display", "Show authoritative available commission on the OlyLife top-up page", "No commission top-up button in VCCHUB V2", "OlyLife"],
        ["Top-up request", "Create pending request linked to member/wallet; make no balance change", "Provide wallet/status reference needed by OlyLife", "OlyLife"],
        ["Admin/Support approval", "Approve/reject with actor, reason, timestamp and dual-control rules", "N/A", "OlyLife"],
        ["Commission debit & wallet credit", "On approval, atomically debit commission and call VCC-01 idempotently", "Credit wallet exactly once and return transaction reference", "Joint"],
        ["Reversal", "Reverse/adjust commission when an approved debit cannot produce a final wallet credit", "Return deterministic status; support retry/reconciliation", "Joint"],
        ["Card eligibility", "N/A", "Require approved KYC, active cardholder and funded wallet; allow one active virtual and one active physical card", "VCCHUB"],
        ["Virtual card", "N/A", "Reuse issuance; show cardholder and registered address read-only", "VCCHUB"],
        ["Physical card", "Provide/approve the office delivery address and change-control process", "Collect no recipient; use configured OlyLife office address and track fulfilment", "VCCHUB"],
        ["Card top-up", "N/A", "Check Wallet balance; debit Wallet and credit card atomically; debit applicable VCCHUB card fees from Wallet", "VCCHUB"],
        ["Card cancellation & history", "No cancellation API integration required", "Expose a per-card action menu; require fresh 2FA; return eligible card balance to Wallet; mark permanently cancelled; release only its virtual/physical type slot and retain the cancelled row/audit history", "VCCHUB"],
    ])

    replace_table(tables[6], ["Journey capability", "Published Issuing coverage", "Classification", "Implementation decision"], [
        ["OlyLife member validation", "No matching Issuing API", "NEW — OLYLIFE", "Provide OLY-01 for email-based active-member validation and return Member ID, First Name and Last Name."],
        ["VCCHUB registration", "No matching Issuing API", "NEW / REUSE — VCCHUB", "Build V2 onboarding, username=email, member fields, Terms, phone/password and registered-address capture."],
        ["Sumsub KYC", "No matching Issuing API", "NEW — JOINT", "OlyLife owns account/access; VCCHUB owns token/SDK/webhook/data retrieval and mapping."],
        ["Create cardholder", "Existing cardholder APIs", "REUSE — VCCHUB INTERNAL", "Use Sumsub-approved identity plus registration email/mobile/address; OlyLife does not call these APIs."],
        ["Create zero-balance wallet", "Existing VCCHUB wallet creation API", "REUSE — VCCHUB INTERNAL", "Create SGD 0.00 wallet internally after KYC approval; no card."],
        ["OlyLife top-up approval", "No matching Issuing API", "NEW — OLYLIFE", "Member requests in OlyLife; Admin/Support approves before commission debit."],
        ["Credit VCCHUB wallet", "Existing internal wallet recharge capability", "NEW PARTNER API + REUSE", "Expose VCC-01 to OlyLife; route approved requests to existing internal wallet credit exactly once."],
        ["Create virtual card", "Existing VCCHUB virtual-card function", "REUSE — VCCHUB INTERNAL", "Issue after funded-wallet checks only when no active virtual card exists; registered data remains read-only."],
        ["Create physical card", "Existing VCCHUB physical-card function", "REUSE + CONFIG — VCCHUB", "Issue only when no active physical card exists, using the configured OlyLife office address; no member recipient/address flow."],
        ["Wallet-to-card top-up", "Existing wallet/card functions", "REUSE + EXTEND — VCCHUB", "Post internally and move applicable VCCHUB fees from Merchant balance to Wallet balance."],
        ["Card servicing / balances", "Existing Issuing and wallet capabilities", "REUSE + CONTROL — VCCHUB INTERNAL", "List each card separately and target actions by card ID. On cancellation, require fresh 2FA, refund eligible balance to Wallet, mark permanently cancelled, release only its type slot and audit the event. Wallet withdrawal remains unavailable."],
        ["Direct login / 2FA", "No matching Issuing API", "NEW / EXTEND — VCCHUB", "Production authentication, MFA, recovery and session controls; V2 base flow has no SSO launch."],
    ])

    set_cell_text(tables[7].rows[1].cells[3], "Reuse after KYC approval using Sumsub-approved identity plus registered email/mobile/address; Member ID remains the external OlyLife key.")
    set_cell_text(tables[7].rows[10].cells[3], "Reuse internally for virtual and physical issuance after funded-wallet checks, with one active card of each type. List both cards separately and target actions by card ID. Physical fulfilment uses the configured OlyLife office address. Cancellation requires fresh 2FA, Wallet refund and release of only that type slot.")
    set_cell_text(tables[7].rows[15].cells[3], "Reuse internally. If enabled, card withdrawal returns funds to Wallet only. Cancellation requires fresh 2FA, marks the selected card permanently cancelled and releases only its virtual/physical type slot; the other active card remains unaffected.")

    replace_table(tables[8], ["ID", "Interface", "Provider", "Consumer", "Logical endpoint/event", "Purpose"], [
        ["OLY-01", "Validate member", "OlyLife", "VCCHUB", "POST /v1/members/validate", "Confirm active/eligible member by email and return approved member fields."],
        ["VCC-01", "Credit wallet", "VCCHUB", "OlyLife", "POST /wallet/topup", "Credit an approved commission top-up exactly once; return wallet transaction/status."],
        ["VCC-02", "Wallet/account status", "VCCHUB", "OlyLife", "POST /wallet/status", "Confirm wallet readiness, wallet ID and top-up eligibility for the member mapping."],
        ["VCC-WH-01", "Wallet account ready", "VCCHUB", "OlyLife", "POST {webhook_url}", "Optional push event to establish/update member-to-wallet mapping without showing integration events in the UI."],
        ["OLY-02", "Top-up request", "OlyLife", "OlyLife Admin/Support", "POST /v1/wallet-topups", "Create a pending member request; no commission debit or wallet credit."],
        ["OLY-03", "Approve/reject top-up", "OlyLife", "OlyLife internal", "POST /v1/wallet-topups/{id}/decision", "Record reviewer decision; on approval, debit commission and call VCC-01."],
        ["OLY-04", "Commission reversal", "OlyLife", "OlyLife internal / VCCHUB ops", "POST /v1/wallet-topups/{id}/reversal", "Reverse an eligible debit if wallet credit cannot complete."],
    ])

    replace_table(tables[9], ["Contract element", "Requirement"], [
        ["Request body", "company_id, email, request_id, requested_at and sign. VCCHUB must not assert active status from browser data."],
        ["Matched response", "code, message and result containing request_id, exists=true, member_status, external_member_id, first_name, last_name, matched_email and approved phone fields."],
        ["Not-found response", "Return a deterministic member-not-found/ineligible code without exposing unrelated member data. VCCHUB shows retry or Contact OlyLife support."],
        ["Security / privacy", "Server-to-server authentication/signing, rate limiting, audit logging and data minimisation. Do not reveal whether arbitrary emails belong to a member outside the controlled VCCHUB flow."],
        ["Idempotency / errors", "Treat request_id as correlation/idempotency key. Define invalid email, member not found, inactive/ineligible, rate limit and service unavailable errors."],
    ])

    replace_table(tables[10], ["Contract element", "Requirement"], [
        ["Precondition", "Approved OlyLife top-up request, successful atomic commission debit, active member-wallet mapping and supported currency."],
        ["Request body", "company_id, request_id/idempotency_key, external_member_id, wallet_id, amount, currency, olylife_transaction_id, approved_at and sign."],
        ["Success response", "code, message and result containing request_id, wallet_transaction_id, status, credited_amount, currency and resulting wallet balance/as-of time."],
        ["Exactly-once control", "The same request_id must return the original result and must never create a second wallet credit. Resolve pending/unknown outcomes before retrying with a new key."],
        ["Failure / reversal", "Return deterministic declined, pending and failed states. If commission was already debited and credit cannot complete, trigger OLY-04 or controlled repair with shared references."],
    ])

    replace_table(tables[11], ["Contract element", "Requirement"], [
        ["Account-ready event", "event_id, event_type=wallet_account_ready, external_member_id, vcchub_user_id, cardholder_id, wallet_id, occurred_at, event_version and sign."],
        ["Status request", "company_id, external_member_id or wallet_id, request_id and sign."],
        ["Status response", "wallet_status, wallet_id, vcchub_user_id, cardholder_id, topup_eligible and updated_at. Do not return credentials or KYC evidence."],
        ["Delivery / reconciliation", "Signed HTTPS event with retry and stable event_id, or idempotent status query. OlyLife stores the mapping before enabling top-up."],
        ["UI rule", "The V2 ready page only shows that the wallet is ready and returns the user to VCCHUB sign-in; integration-event details remain backend-only."],
        ["OlyLife control", "Verify signature/timestamp, process idempotently and do not enable funding for an unverified or ambiguous member-wallet mapping."],
    ])

    replace_table(tables[12], ["Contract element", "Requirement"], [
        ["Request", "Member ID, wallet ID, amount, currency and member confirmation create a PENDING_APPROVAL record. No balance changes."],
        ["Review", "Admin/Support views commission balance, request amount, member/wallet mapping and risk context; approve or reject with actor, timestamp and reason."],
        ["Approval", "Atomically recheck eligibility/sufficient commission and debit once; then call VCC-01 using a stable idempotency key."],
        ["Rejection / failure", "Rejection leaves both balances unchanged. If debit succeeds but wallet credit fails, retry safely or initiate OLY-04 reversal/repair."],
        ["Reconciliation", "Store top-up request ID, reviewer, OlyLife transaction ID, VCC-01 request ID, VCCHUB wallet transaction ID, amount, currency, timestamps and final status."],
    ])

    replace_table(tables[14], ["Data", "Source of truth", "Exchange", "Implementation rule"], [
        ["OlyLife Member ID", "OlyLife", "OlyLife → VCCHUB", "Stable external key returned by OLY-01 and used for mapping, support and top-up."],
        ["Member eligibility", "OlyLife", "OlyLife → VCCHUB", "Only an active/eligible matched result may proceed to registration."],
        ["OlyLife First/Last Name", "OlyLife", "OlyLife → VCCHUB", "Displayed read-only during registration as member-reference data; do not silently overwrite the OlyLife source record."],
        ["Email, mobile, password", "VCCHUB registration", "No password exchange", "Username equals matched email; normalize mobile; password remains VCCHUB-only."],
        ["Registered address", "VCCHUB registration", "VCCHUB internal / processor as needed", "Collected during V2 registration and shown read-only at card creation."],
        ["KYC evidence & decision", "Sumsub account contracted by OlyLife; operated by VCCHUB", "Minimal status/mapping", "Raw images/biometrics follow approved Sumsub/VCCHUB access and retention."],
        ["Verified First/Last Name & DOB", "Sumsub approved applicant", "VCCHUB/card processor", "Authoritative cardholder identity; define mismatch/manual-review rules against OlyLife member names."],
        ["Commission / wallet / card ledgers", "OlyLife / VCCHUB", "References and status only", "Each party remains authoritative for its own ledger; reconcile using shared IDs."],
        ["Physical office address", "OlyLife approved; configured by VCCHUB", "VCCHUB / fulfilment provider", "No member recipient is collected; version, secure and audit office-address changes."],
    ])

    replace_table(tables[15], ["Field", "Authoritative source", "V2 mapping rule"], [
        ["Member ID", "OlyLife OLY-01 response", "Persist as the external member key; display read-only where operationally useful."],
        ["First / Last Name at registration", "OlyLife OLY-01 response", "Carry into registration read-only as the matched member profile."],
        ["Cardholder First / Last Name", "Sumsub approved applicant", "Use approved legal/transliterated fields; define mismatch and no-surname handling."],
        ["Date of birth", "Sumsub approved applicant", "Normalize to ISO date and validate programme eligibility."],
        ["Email / Mobile", "Matched OlyLife email / VCCHUB registration", "Username equals matched email; store normalized E.164 mobile and verification status."],
        ["Registered address", "VCCHUB registration", "Persist full address and show read-only on card creation; it is not the physical delivery address in V2."],
    ])

    set_cell_text(tables[16].rows[6].cells[1], "Required for member validation, account-ready/status processing, top-up approval/debit, wallet credit, reversal, card issue, fee posting, card top-up and card cancellation.")

    replace_table(tables[17], ["Case", "Lead support owner", "Evidence"], [
        ["Member not found/ineligible", "OlyLife member support; VCCHUB shows retry/contact route", "OLY-01 request/correlation ID and eligibility result"],
        ["KYC pending/rejected", "VCCHUB KYC operations with Sumsub escalation", "Sumsub applicant/review IDs and reason codes"],
        ["Top-up approval/debit", "OlyLife Admin/Support and finance operations", "Request, reviewer, commission ledger and decision audit"],
        ["Wallet/card balance", "VCCHUB operations / processor", "VCCHUB ledger and processor reference"],
        ["Physical office delivery", "VCCHUB / fulfilment provider with OlyLife facilities contact", "Configured office-address version, fulfilment and tracking references"],
    ])

    replace_table(tables[18], ["Scenario", "Required behaviour"], [
        ["Member not found/ineligible", "VCCHUB creates no account and prompts the user to retry the email or contact OlyLife support."],
        ["Member-validation unavailable", "Show a recoverable service message; do not bypass OLY-01 or accept browser-supplied eligibility."],
        ["Incomplete OlyLife member data", "Do not continue if required Member ID/name/email is missing; return a support-safe correction path."],
        ["Registration address incomplete", "Prevent submission until country, state/province, city, address line 1 and postal code pass validation."],
        ["KYC pending", "Show pending state and refresh from authoritative Sumsub webhook/status; create no cardholder/wallet yet."],
        ["KYC rejected/resubmission", "Display permitted action and route manual-review/support cases; exchange only necessary status."],
        ["OlyLife/Sumsub name mismatch", "Do not use mock names or silently alter sources; apply agreed cardholder precedence and manual-review policy."],
        ["Account mapping unavailable", "Do not enable OlyLife top-up until account-ready event/status establishes one unambiguous member-wallet mapping."],
        ["Top-up rejected", "Record reviewer/reason; leave commission and wallet balances unchanged."],
        ["Commission insufficient", "Recheck at approval, decline atomically without debit and create no wallet credit."],
        ["Debit succeeds, wallet credit fails", "Retry VCC-01 with the same idempotency key; otherwise trigger reversal/manual repair with shared references."],
        ["Card type, issue or top-up control", "Reject a second active card of the same type and show one row with independent actions per card. If issue/top-up fails, post neither final debit/credit nor fees, or reverse reservations atomically."],
        ["OlyLife office address unavailable", "Block physical issuance; do not fall back to a member-entered recipient. Route to VCCHUB/OlyLife operations."],
        ["Card cancellation", "Require fresh 2FA and explicit irreversible confirmation; refund eligible card balance to Wallet exactly once, mark the selected card permanently cancelled, release only that card-type slot and record the cancellation audit/transaction. The other card type remains occupied. Wallet withdrawal remains unsupported."],
    ])

    replace_table(tables[19], ["Phase", "Lead", "Exit criteria"], [
        ["1. Contract & compliance", "Joint", "Approved V2 journey; responsibilities; member-data/KYC/address purpose; API/event schemas; security and test plan."],
        ["2. Member validation & registration", "OlyLife + VCCHUB", "OLY-01 passes matched/not-found/unavailable tests; Member ID/name carry-over; username=email; phone/password/address and consent evidence."],
        ["3. Sumsub & wallet creation", "OlyLife + VCCHUB", "Provisioned Sumsub access; approved identity mapping; zero-balance wallet; account-ready/status mapping."],
        ["4. Access security", "VCCHUB", "Direct sign-in, real 2FA, recovery, session, rate-limit and audit controls pass security tests."],
        ["5. OlyLife top-up approval", "OlyLife + VCCHUB", "Pending request, Admin/Support decision, atomic debit, exactly-once wallet credit, reversal and reconciliation pass."],
        ["6. Card services", "VCCHUB", "Funding gate; one active virtual plus one active physical card; separate card rows/actions; read-only identity/address; OlyLife-office delivery; wallet-to-card top-up; authenticated cancellation with Wallet refund and type-slot release; Wallet-based fees."],
        ["7. UAT & production", "Joint", "End-to-end happy/negative paths, monitoring, operations, support, cutover and rollback accepted."],
    ])

    replace_table(tables[20], ["Decision area", "Items to agree"], [
        ["KYC", "Sumsub production level(s), verification rules, manual review, transliteration/cardholder-name policy, resubmission, retention and access."],
        ["Member validation & mapping", "Eligible member statuses; required OLY-01 fields; matched-email rules; OlyLife/Sumsub name mismatch handling; account-ready event/status ownership."],
        ["Identity/access", "Password/MFA methods, remembered devices, recovery, lockout/rate limits, session expiry and account/email change process."],
        ["Top-up approval & money movement", "Reviewer roles/dual control, limits, atomic debit, VCC-01 idempotency, settlement, reversal, reconciliation and manual adjustment controls."],
        ["Card servicing", "Confirmed: one active virtual and one active physical card, separate card rows/actions, and cancellation with fresh 2FA, Wallet refund, permanent Cancelled status and release of only that type slot. Agree processor edge cases, fees and whether card-to-Wallet withdrawal will ever be enabled separately."],
        ["Physical cards", "Actual OlyLife office address, authorised change process, delivery fee, courier, tracking, failed delivery, replacement and office receiving owner."],
        ["Operations", "Support tiers, escalation contacts, monitoring ownership, incident notification, DR/RTO/RPO, audit and reporting cadence."],
    ])

    replace_table(tables[21], ["Owner", "Next action"], [
        ["VCCHUB", "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH-01 schemas; implement V2 onboarding, member carry-over, registration address, Sumsub, direct 2FA, read-only card data, one active card per type, separate per-card actions, OlyLife-office physical fulfilment and authenticated cancellation with Wallet refund/type-slot release."],
        ["OlyLife", "Publish OLY-01 with Member ID/First Name/Last Name; approve Sumsub access; implement top-up request/review/debit/reversal; provide the actual office address and facilities/support ownership."],
        ["Joint", "Run an interface workshop to close identifiers, auth, status/error models, account mapping, wallet-credit idempotency, reversal/reconciliation and support ownership."],
    ])

    doc.core_properties.title = "OlyLife–VCCHUB Implementation Responsibility Specification – Version 2"
    doc.core_properties.subject = "VCCHUB-first member validation, registration, Sumsub KYC, OlyLife-approved wallet top-up and authenticated card cancellation"
    doc.core_properties.author = "Star SaaS Limited"
    doc.core_properties.keywords = "OlyLife, VCCHUB, Version 2, member validation, Sumsub, wallet, card, cancellation, 2FA, KYC, top-up approval, confidential"
    # The V2 scope combines the previous fifth bullet into the fourth. Its
    # taller cover can push the template's standalone page-break paragraph to
    # a blank page, so put the page break directly on the disclaimer logo.
    paragraphs[13]._p.getparent().remove(paragraphs[13]._p)
    paragraphs[14]._p.getparent().remove(paragraphs[14]._p)
    disclaimer_logo_ppr = paragraphs[15]._p.get_or_add_pPr()
    if not disclaimer_logo_ppr.xpath("./w:pageBreakBefore"):
        disclaimer_logo_ppr.append(OxmlElement("w:pageBreakBefore"))
    # The source template also contains a manual break immediately before
    # Chapter 7. Once the V2 tables grow, that break can land at the top of a
    # page and create an otherwise blank page. Remove only that redundant
    # break; the chapter flows naturally onto the next page.
    live_paragraphs = doc.paragraphs
    for index, paragraph in enumerate(live_paragraphs[:-1]):
        next_text = live_paragraphs[index + 1].text.strip()
        if not paragraph.text.strip() and paragraph._p.xpath(".//w:br") and next_text.startswith("7. Failure and exception handling"):
            paragraph._p.getparent().remove(paragraph._p)
            break
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_v2()
