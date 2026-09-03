from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from build_responsibility_document import (
    BLUE,
    CONTENT_DXA,
    INK,
    LIGHT_BLUE,
    LIGHT_GRAY,
    MID_GRAY,
    NAVY,
    format_table,
    set_cant_split,
    set_cell_margins,
    set_paragraph_border,
    set_paragraph_shading,
    set_run,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.1.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.2.docx"


def find_paragraph(doc: Document, prefix: str):
    return next(p for p in doc.paragraphs if p.text.strip().startswith(prefix))


def delete_between(start_paragraph, end_paragraph) -> None:
    element = start_paragraph._p.getnext()
    while element is not None and element is not end_paragraph._p:
        next_element = element.getnext()
        element.getparent().remove(element)
        element = next_element


def move_before(element, anchor_paragraph) -> None:
    anchor_paragraph._p.addprevious(element)


def insert_paragraph(doc: Document, anchor, text: str = "", style: str | None = None,
                     size: float = 10.2, color: str = INK, bold: bool = False,
                     italic: bool = False, before: float = 0, after: float = 6):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.18
    if text:
        set_run(paragraph.add_run(text), size, color, bold, italic)
    move_before(paragraph._p, anchor)
    return paragraph


def insert_heading(doc: Document, anchor, text: str, level: int = 3, page_break: bool = False):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    move_before(paragraph._p, anchor)
    return paragraph


def insert_table(doc: Document, anchor, headers, rows, widths, body_size: float = 8.2,
                 first_col_bold: bool = True):
    table = doc.add_table(rows=1, cols=len(headers))
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
    format_table(table, headers, widths, body_size=body_size, first_col_bold=first_col_bold)
    for row in table.rows:
        set_cant_split(row)
        for cell in row.cells:
            set_cell_margins(cell, {"top": 90, "bottom": 90, "start": 115, "end": 115})
    move_before(table._tbl, anchor)
    return table


def insert_code(doc: Document, anchor, label: str, payload: dict):
    insert_paragraph(doc, anchor, label, style="Heading 3", size=11.2, color=NAVY, bold=True, before=7, after=4)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(8)
    paragraph.paragraph_format.right_indent = Pt(4)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.0
    # Keep compact JSON examples on one page. Word will still split an example
    # that is genuinely taller than a page, but this avoids orphaned closing
    # lines on a mostly blank page for normal request/response payloads.
    paragraph.paragraph_format.keep_together = True
    set_paragraph_shading(paragraph, LIGHT_GRAY)
    set_paragraph_border(paragraph, "left", BLUE, 14, 5)
    run = paragraph.add_run(json.dumps(payload, indent=2, ensure_ascii=True))
    set_run(run, 7.3, INK)
    run.font.name = "Courier New"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Courier New")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Courier New")
    move_before(paragraph._p, anchor)
    return paragraph


def update_cell(cell, text: str) -> None:
    cell.text = text


def update_control_and_shared_decision(doc: Document) -> None:
    control = doc.tables[0]
    update_cell(control.rows[0].cells[3], "Draft - V2.2 message contracts")
    update_cell(control.rows[2].cells[1], "0.13 (V2.2 API and webhook schemas)")
    update_cell(control.rows[2].cells[3], "3 September 2026")

    shared = next(
        t for t in doc.tables
        if len(t.columns) == 3 and t.rows[0].cells[0].text.strip() == "Decision area"
    )
    for row in shared.rows[1:]:
        if row.cells[0].text.strip().startswith("2. API security"):
            update_cell(
                row.cells[1],
                "Use server-to-server HTTPS, company_id, UUID request_id, requested_at and the existing Star SaaS Issuing signing convention: "
                "sort all JSON parameters alphabetically by key, exclude sign, serialize the canonical JSON, append the pre-shared signkey and hash with SHA-256. "
                "Allow a five-minute clock skew. The same request_id returns the original result and never repeats a debit or credit. external_member_id is the cross-system member key.",
            )
            update_cell(
                row.cells[2],
                "Confirm the SHA-256 signing convention, canonical JSON rules, key exchange/rotation, clock-skew window and external_member_id format; nominate OlyLife's API owner.",
            )
            break
    format_table(shared, ["Decision area", "VCCHUB proposed default", "OlyLife action / response"],
                 [1900, 5000, 2460], body_size=8.1, first_col_bold=True)

    security = next(
        t for t in doc.tables
        if len(t.columns) == 2 and t.rows[0].cells[0].text.strip() == "Control area"
    )
    for row in security.rows[1:]:
        if row.cells[0].text.strip() == "Webhook security":
            update_cell(
                row.cells[1],
                "Use the agreed Star SaaS signing convention over the canonical JSON body: recursively sort object keys, exclude sign, append the environment signkey and calculate lowercase SHA-256 hex. Enforce the timestamp replay window, event_id idempotency, HTTPS and source allowlisting as appropriate.",
            )
            break
    format_table(
        security,
        ["Control area", "Minimum production requirement"],
        [1900, CONTENT_DXA - 1900],
        body_size=8.4,
        first_col_bold=True,
    )


def expand_interface_catalogue(doc: Document) -> None:
    table = next(
        t for t in doc.tables
        if len(t.columns) == 6 and t.rows[0].cells[0].text.strip() == "ID"
    )
    existing = {row.cells[0].text.strip() for row in table.rows[1:]}
    additions = [
        ["VCC-WH-02", "Wallet top-up result", "VCCHUB", "OlyLife", "POST {webhook_url}", "Final or failed asynchronous wallet-credit result."],
        ["VCC-WH-03", "Card created", "VCCHUB", "OlyLife", "POST {webhook_url}", "Synchronize a newly issued virtual or physical card without PAN/CVV."],
        ["VCC-WH-04", "Card top-up result", "VCCHUB", "OlyLife", "POST {webhook_url}", "Synchronize successful or failed wallet-to-card funding."],
        ["VCC-WH-05", "Card cancelled", "VCCHUB", "OlyLife", "POST {webhook_url}", "Synchronize cancellation, eligible Wallet refund and released card-type slot."],
    ]
    for values in additions:
        if values[0] not in existing:
            cells = table.add_row().cells
            for index, value in enumerate(values):
                cells[index].text = value
    format_table(
        table,
        ["ID", "Interface", "Provider", "Consumer", "Logical endpoint/event", "Purpose"],
        [720, 1420, 900, 980, 2100, 3240],
        body_size=7.7,
        first_col_bold=True,
    )
    for row in table.rows:
        set_cant_split(row)


def build_oly_01(doc: Document, anchor) -> None:
    insert_paragraph(
        doc, anchor,
        "OLY-01 is provided and operated by OlyLife and consumed only by the VCCHUB backend. The field definitions below are the proposed implementation contract; OlyLife may confirm or return amended limits before OpenAPI freeze.",
        italic=True, color=MID_GRAY,
    )
    insert_heading(doc, anchor, "4.4.1 Request fields")
    insert_table(doc, anchor, ["Element", "Req.", "Type / format", "Description and validation"], [
        ["company_id", "Yes", "integer", "VCCHUB/OlyLife partner company identifier issued for the environment."],
        ["email", "Yes", "string, max 254", "Trim, lowercase for matching and validate as an email address. Do not use browser-supplied member status."],
        ["request_id", "Yes", "string, max 64", "VCCHUB-generated unique idempotency/correlation key. The same ID and payload return the original result."],
        ["requested_at", "Yes", "RFC 3339 UTC", "Time the request was signed, for example 2026-09-03T02:15:30Z. Proposed replay window: five minutes."],
        ["sign", "Yes", "64-char hex", "SHA-256 signature using the agreed sorted-parameter plus signkey convention; exclude sign from the signing input."],
    ], [1650, 700, 1800, 5210])

    insert_heading(doc, anchor, "4.4.2 Response fields")
    insert_table(doc, anchor, ["Element", "Presence", "Type / values", "Consumer handling"], [
        ["code", "Always", "string", "0000 means success. Any other code must be handled as a business/technical outcome, not inferred from message text."],
        ["message", "Always", "string", "Human-readable summary for logs/support; never parse it for program logic."],
        ["result.request_id", "Always", "string", "Echo the request_id for end-to-end correlation."],
        ["result.exists", "Always", "boolean", "true only when one controlled member match exists."],
        ["result.eligible", "Always", "boolean", "true only when the matched member may enter this wallet programme."],
        ["result.member_status", "Always", "ACTIVE, INACTIVE, SUSPENDED or NOT_FOUND", "VCCHUB continues only for ACTIVE with exists=true and eligible=true."],
        ["result.external_member_id", "Matched only", "string", "Stable OlyLife member identifier persisted as the external mapping key."],
        ["result.first_name / last_name", "Matched only", "UTF-8 strings", "Member-reference names shown read-only during registration; KYC-approved names remain authoritative for cardholder creation."],
        ["result.matched_email", "Matched only", "email string", "Canonical matched address and VCCHUB username. It must equal the controlled normalized match."],
        ["result.mobile_prefix / mobile", "Optional", "E.164 components", "Return only where approved. VCCHUB validates or collects the mobile during registration."],
        ["result.updated_at", "Matched only", "RFC 3339 UTC", "Last relevant OlyLife member-profile update time."],
    ], [1720, 980, 2200, 4460], body_size=7.9)

    insert_heading(doc, anchor, "4.4.3 Proposed HTTP and business outcomes")
    insert_table(doc, anchor, ["HTTP", "code", "Meaning", "VCCHUB action"], [
        ["200", "0000", "Matched and eligible", "Continue to registration using returned member fields."],
        ["200", "OLY2001", "Not found or not eligible", "Create no account; show retry or Contact OlyLife support without revealing private membership detail."],
        ["422", "OLY2002", "Matched record lacks required data", "Create no account; show support-safe correction route and log request_id."],
        ["400", "OLY1001", "Invalid request", "Correct the named source field; do not retry unchanged."],
        ["401", "OLY1002", "Invalid/expired signature", "Stop, alert integration operations and verify key/time configuration."],
        ["409", "OLY1003", "request_id reused with different payload", "Treat as an integration defect; never create a new request ID blindly."],
        ["429/503", "OLY9000", "Rate-limited or temporarily unavailable", "Retry with bounded backoff using the same request_id; do not bypass member validation."],
    ], [800, 1100, 2700, 4760], body_size=7.9)

    insert_code(doc, anchor, "Illustrative request", {
        "company_id": 3201,
        "email": "olivia.chen@example.com",
        "request_id": "oly-member-20260903-000001",
        "requested_at": "2026-09-03T02:15:30Z",
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Illustrative matched response", {
        "code": "0000",
        "message": "Success",
        "result": {
            "request_id": "oly-member-20260903-000001",
            "exists": True,
            "eligible": True,
            "member_status": "ACTIVE",
            "external_member_id": "OLY-10002345",
            "first_name": "Olivia",
            "last_name": "Chen",
            "matched_email": "olivia.chen@example.com",
            "mobile_prefix": "+65",
            "mobile": "81234567",
            "updated_at": "2026-09-03T02:15:31Z"
        }
    })
    insert_code(doc, anchor, "Illustrative privacy-safe not-found response", {
        "code": "OLY2001",
        "message": "Member not found or not eligible",
        "result": {
            "request_id": "oly-member-20260903-000001",
            "exists": False,
            "eligible": False,
            "member_status": "NOT_FOUND"
        }
    })


def build_vcc_01(doc: Document, anchor) -> None:
    insert_paragraph(
        doc, anchor,
        "VCC-01 is provided by VCCHUB and called by OlyLife only after an approved, atomic commission debit. Monetary values use decimal strings and ISO 4217 currency codes. request_id is the exactly-once key across debit, API response, webhook and reconciliation.",
        italic=True, color=MID_GRAY,
    )
    insert_heading(doc, anchor, "4.5.1 Request fields")
    insert_table(doc, anchor, ["Element", "Req.", "Type / format", "Description and validation"], [
        ["company_id", "Yes", "integer", "OlyLife partner/company identifier issued by VCCHUB for the environment."],
        ["request_id", "Yes", "string, max 64", "Stable exactly-once key. OlyLife must reuse it for retry/recovery of the same approved top-up."],
        ["external_member_id", "Yes", "string", "OlyLife member key already mapped to one active VCCHUB wallet."],
        ["wallet_id", "Yes", "string", "Target VCCHUB wallet from the account-ready mapping/status response."],
        ["amount", "Yes", "decimal string", "Positive amount in currency minor-unit precision, for example 500.00. Never use a binary floating-point value."],
        ["currency", "Yes", "ISO 4217", "Must equal the wallet currency; proposed V2 base currency is SGD."],
        ["olylife_transaction_id", "Yes", "string", "Immutable OlyLife commission-ledger debit reference."],
        ["approved_at", "Yes", "RFC 3339 UTC", "Recorded OlyLife Admin/Support approval time."],
        ["requested_at", "Yes", "RFC 3339 UTC", "Time this signed API request was created."],
        ["webhook_url", "Optional", "HTTPS URL", "Override callback for this operation; otherwise use the OlyLife endpoint configured in VCCHUB."],
        ["sign", "Yes", "64-char hex", "SHA-256 signature following the common signing contract."],
    ], [1650, 700, 1800, 5210], body_size=7.9)

    insert_heading(doc, anchor, "4.5.2 Response fields")
    insert_table(doc, anchor, ["Element", "Presence", "Type / values", "OlyLife handling"], [
        ["code / message", "Always", "strings", "Use code for program logic and message for logs/support."],
        ["result.request_id", "Always", "string", "Echoed exactly-once/correlation key."],
        ["result.wallet_transaction_id", "Accepted/final", "string", "VCCHUB wallet-ledger reference persisted by OlyLife."],
        ["result.status", "Always", "COMPLETED, PROCESSING, DECLINED or FAILED", "COMPLETED is final success. PROCESSING is non-final; wait for webhook or query status using the same identifiers."],
        ["result.credited_amount", "Completed", "decimal string", "Net principal credited to the wallet, excluding separately reported fees."],
        ["result.currency", "Accepted/final", "ISO 4217", "Must match the request."],
        ["result.wallet_balance", "Completed", "decimal string", "Authoritative wallet balance immediately after the posting."],
        ["result.balance_as_of", "Completed", "RFC 3339 UTC", "Timestamp for wallet_balance."],
        ["result.external_member_id / wallet_id", "Always", "strings", "Echoed mapping values; reject/reconcile any mismatch."],
        ["result.failure_code / failure_message", "Failed/declined", "strings", "Machine-readable reason plus support-safe description; never infer outcome only from HTTP timeout."],
    ], [1780, 980, 2200, 4400], body_size=7.9)

    insert_heading(doc, anchor, "4.5.3 Proposed HTTP and processing outcomes")
    insert_table(doc, anchor, ["HTTP", "code", "Status", "OlyLife action"], [
        ["200", "0000", "COMPLETED", "Mark top-up completed and reconcile the returned wallet transaction/balance."],
        ["202", "VCC3002", "PROCESSING", "Do not debit again or create a new request ID. Wait for VCC-WH-02 or query /wallet/status."],
        ["400/422", "VCC1001", "DECLINED", "Treat as final input/business rejection; start reversal if commission was already debited."],
        ["401", "VCC1002", "DECLINED", "Stop and resolve signature/key/time configuration."],
        ["404", "VCC2001", "DECLINED", "Member-wallet mapping not found; do not redirect funds."],
        ["409", "VCC3001", "DECLINED", "Same request_id has a different payload; escalate as an integrity defect."],
        ["5xx/timeout", "VCC9000", "UNKNOWN", "Retry the identical payload with the same request_id, then reconcile. Never issue a second commission debit."],
    ], [800, 1100, 1300, 6160], body_size=7.9)

    insert_code(doc, anchor, "Illustrative request", {
        "company_id": 3201,
        "request_id": "topup-20260903-000045",
        "external_member_id": "OLY-10002345",
        "wallet_id": "WLT-90007812",
        "amount": "500.00",
        "currency": "SGD",
        "olylife_transaction_id": "COM-DB-880031",
        "approved_at": "2026-09-03T03:05:00Z",
        "requested_at": "2026-09-03T03:05:02Z",
        "webhook_url": "https://api.olylife.example/v1/vcchub/events",
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Illustrative completed response", {
        "code": "0000",
        "message": "Success",
        "result": {
            "request_id": "topup-20260903-000045",
            "wallet_transaction_id": "WTX-70009111",
            "status": "COMPLETED",
            "external_member_id": "OLY-10002345",
            "wallet_id": "WLT-90007812",
            "credited_amount": "500.00",
            "currency": "SGD",
            "wallet_balance": "500.00",
            "balance_as_of": "2026-09-03T03:05:03Z"
        }
    })
    insert_code(doc, anchor, "Illustrative non-final response", {
        "code": "VCC3002",
        "message": "Wallet credit is processing",
        "result": {
            "request_id": "topup-20260903-000045",
            "wallet_transaction_id": "WTX-70009111",
            "status": "PROCESSING",
            "external_member_id": "OLY-10002345",
            "wallet_id": "WLT-90007812"
        }
    })


def build_vcc_02(doc: Document, anchor) -> None:
    insert_paragraph(
        doc, anchor,
        "VCC-02 is the authoritative reconciliation endpoint when an account-ready or lifecycle webhook is delayed, duplicated or missed. OlyLife may query by external_member_id, wallet_id, or both; if both are supplied they must map to the same wallet.",
        italic=True, color=MID_GRAY,
    )
    insert_heading(doc, anchor, "4.6.1 Status request fields")
    insert_table(doc, anchor, ["Element", "Req.", "Type / format", "Description"], [
        ["company_id", "Yes", "integer", "OlyLife partner/company identifier."],
        ["external_member_id", "Conditional", "string", "Required when wallet_id is not supplied."],
        ["wallet_id", "Conditional", "string", "Required when external_member_id is not supplied."],
        ["include_cards", "No", "boolean", "When true, return non-sensitive card summaries for lifecycle reconciliation."],
        ["request_id", "Yes", "string, max 64", "Unique correlation/idempotency key for the query."],
        ["requested_at", "Yes", "RFC 3339 UTC", "Signed request time."],
        ["sign", "Yes", "64-char hex", "SHA-256 signature following the common signing contract."],
    ], [1700, 900, 1850, 4910], body_size=8.0)

    insert_heading(doc, anchor, "4.6.2 Status response fields")
    insert_table(doc, anchor, ["Element", "Presence", "Type / values", "Meaning / handling"], [
        ["code / message", "Always", "strings", "Standard response envelope."],
        ["result.request_id", "Always", "string", "Echoed correlation key."],
        ["result.external_member_id", "Always", "string", "OlyLife external mapping key."],
        ["result.vcchub_user_id / cardholder_id / wallet_id", "Ready mapping", "strings", "Persist together as one mapping; top-up requires an unambiguous active mapping."],
        ["result.wallet_status", "Always", "PENDING_KYC, READY, SUSPENDED or CLOSED", "Only READY may be top-up eligible."],
        ["result.topup_eligible", "Always", "boolean", "VCCHUB decision after wallet/cardholder/compliance checks."],
        ["result.currency / wallet_balance", "Ready wallet", "ISO 4217 / decimal string", "Authoritative wallet denomination and current balance."],
        ["result.cards[]", "If include_cards=true", "array", "One item per card: card_id, card_type, card_status, last4, currency and card_balance. Never return PAN or CVV."],
        ["result.updated_at", "Always", "RFC 3339 UTC", "As-of time for the returned status snapshot."],
    ], [1850, 980, 2250, 4280], body_size=7.9)

    insert_code(doc, anchor, "Illustrative status request", {
        "company_id": 3201,
        "external_member_id": "OLY-10002345",
        "wallet_id": "WLT-90007812",
        "include_cards": True,
        "request_id": "status-20260903-000090",
        "requested_at": "2026-09-03T04:10:00Z",
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Illustrative status response", {
        "code": "0000",
        "message": "Success",
        "result": {
            "request_id": "status-20260903-000090",
            "external_member_id": "OLY-10002345",
            "vcchub_user_id": "USR-50001120",
            "cardholder_id": "1952564520941649920",
            "wallet_id": "WLT-90007812",
            "wallet_status": "READY",
            "topup_eligible": True,
            "currency": "SGD",
            "wallet_balance": "425.00",
            "cards": [
                {"card_id": "CRD-30001001", "card_type": "VIRTUAL", "card_status": "ACTIVE", "last4": "1965", "currency": "SGD", "card_balance": "75.00"},
                {"card_id": "CRD-30001002", "card_type": "PHYSICAL", "card_status": "DELIVERY_PENDING", "last4": "6366", "currency": "SGD", "card_balance": "0.00"}
            ],
            "updated_at": "2026-09-03T04:10:01Z"
        }
    })


def build_webhooks(doc: Document, anchor) -> None:
    insert_heading(doc, anchor, "4.7 Partner webhook message contract", level=2)
    insert_paragraph(
        doc, anchor,
        "VCCHUB sends lifecycle events to OlyLife so the member record, wallet funding history and optional card summary remain synchronized. The contract intentionally excludes passwords, full PAN, CVV, raw KYC documents and biometrics. All values below are proposed defaults for joint OpenAPI/event-schema sign-off.",
        italic=True, color=MID_GRAY,
    )

    insert_heading(doc, anchor, "4.7.1 Delivery, signing and acknowledgement rules")
    insert_table(doc, anchor, ["Rule", "Proposed contract"], [
        ["Endpoint", "OlyLife provides one HTTPS POST endpoint per environment, for example /v1/vcchub/events. A per-operation webhook_url may override the configured endpoint only where explicitly allowed."],
        ["Content type", "application/json; UTF-8. Maximum proposed payload size: 64 KB."],
        ["Signing", "Use the Star SaaS Issuing convention: recursively sort object keys, exclude sign, serialize canonical JSON without insignificant whitespace, append the environment signkey and calculate lowercase SHA-256 hex."],
        ["Delivery identity", "event_id is globally unique and stable across retries. OlyLife deduplicates by event_id and stores the first valid payload/result."],
        ["Ordering", "Delivery is at least once and may be out of order. resource_version increases for each resource; OlyLife ignores an older version after acknowledging it."],
        ["Acknowledgement", "Return HTTP 200 after durable idempotent acceptance. Duplicate valid events also return 200. The response body echoes event_id using code/message/result."],
        ["Retry", "Any non-2xx or timeout triggers the Star SaaS webhook retry pattern: first retry after about one second, exponential backoff with jitter and a delivery window of up to 36 hours. Exhausted events enter operational reconciliation."],
        ["Receiver errors", "401 for invalid signature/expired timestamp, 400 for invalid schema, 429 for throttling and 5xx for a temporary receiver failure. Do not acknowledge before durable acceptance."],
        ["Reconciliation", "OlyLife may call POST /wallet/status after a missed/uncertain event. VCCHUB exposes replay by event_id to authorised operations."],
    ], [2100, 7260], body_size=8.2)

    insert_heading(doc, anchor, "4.7.2 Common webhook envelope")
    insert_table(doc, anchor, ["Element", "Req.", "Type / format", "Description"], [
        ["event_id", "Yes", "string, max 64", "Globally unique event identifier and OlyLife deduplication key."],
        ["event_type", "Yes", "enum string", "One of the event names in the catalogue below."],
        ["event_version", "Yes", "string", "Schema version, initially 1.0."],
        ["occurred_at", "Yes", "RFC 3339 UTC", "Time the underlying state transition committed."],
        ["company_id", "Yes", "integer", "Partner/company identifier for routing and key selection."],
        ["request_id", "Conditional", "string", "Original API/action correlation key when an initiating request exists."],
        ["external_member_id", "Yes", "string", "OlyLife member mapping key."],
        ["vcchub_user_id / wallet_id", "Yes", "strings", "VCCHUB account and wallet identifiers."],
        ["resource_type / resource_id", "Yes", "strings", "Resource changed: wallet, wallet_topup, card or card_topup, plus its stable identifier."],
        ["resource_version", "Yes", "positive integer", "Monotonic version used for out-of-order delivery handling."],
        ["status", "Yes", "event-specific enum", "New committed or final operation status."],
        ["data", "Yes", "JSON object", "Event-specific fields. Values are part of the signed canonical payload."],
        ["sign", "Yes", "64-char hex", "SHA-256 signature over the full envelope except sign."],
    ], [1750, 700, 1900, 5010], body_size=7.8)

    insert_heading(doc, anchor, "4.7.3 Lifecycle event catalogue")
    insert_table(doc, anchor, ["Event type", "When emitted", "Required data fields", "OlyLife handling"], [
        ["wallet.account_ready", "KYC-approved cardholder and zero-balance wallet are committed", "cardholder_id, wallet_status=READY, currency, wallet_balance", "Persist mapping; enable OlyLife wallet top-up only if topup_eligible=true."],
        ["wallet.topup.succeeded", "Approved commission-funded wallet credit commits", "wallet_transaction_id, olylife_transaction_id, amount, currency, wallet_balance, balance_as_of", "Mark request completed and reconcile the commission debit to the VCCHUB credit."],
        ["wallet.topup.failed", "Wallet credit reaches a final failure", "olylife_transaction_id, failure_code, failure_message, wallet_balance, balance_as_of", "Open reversal/repair using the original IDs; do not debit again."],
        ["card.created", "Virtual card is active or physical issuance is accepted", "card_id, card_type, card_status, last4, currency, card_balance, delivery_destination", "Store a non-sensitive card summary. Never expect PAN/CVV."],
        ["card.topup.succeeded", "Wallet-to-card transfer commits", "card_id, wallet_transaction_id, card_transaction_id, amount, currency, wallet_balance, card_balance", "Update optional balance mirror/history using the selected card_id."],
        ["card.topup.failed", "Wallet-to-card top-up reaches final failure", "card_id, failure_code, failure_message, wallet_balance, card_balance", "Record failure; do not infer or post a successful transfer."],
        ["card.cancelled", "Cancellation, eligible Wallet refund and slot release commit", "card_id, card_type, card_status=CANCELLED, refunded_amount, wallet_balance, card_type_slot_released", "Mark only the selected card cancelled; preserve history and permit replacement of that type."],
    ], [2150, 2450, 3100, 1660], body_size=7.6)

    insert_heading(doc, anchor, "4.7.4 Illustrative webhook messages")
    insert_code(doc, anchor, "Wallet account ready", {
        "event_id": "EVT-20260903-000001", "event_type": "wallet.account_ready", "event_version": "1.0",
        "occurred_at": "2026-09-03T02:45:10Z", "company_id": 3201, "request_id": "onboard-20260903-001",
        "external_member_id": "OLY-10002345", "vcchub_user_id": "USR-50001120", "wallet_id": "WLT-90007812",
        "resource_type": "wallet", "resource_id": "WLT-90007812", "resource_version": 1, "status": "READY",
        "data": {"cardholder_id": "1952564520941649920", "wallet_status": "READY", "topup_eligible": True, "currency": "SGD", "wallet_balance": "0.00"},
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Wallet top-up successful", {
        "event_id": "EVT-20260903-000045", "event_type": "wallet.topup.succeeded", "event_version": "1.0",
        "occurred_at": "2026-09-03T03:05:03Z", "company_id": 3201, "request_id": "topup-20260903-000045",
        "external_member_id": "OLY-10002345", "vcchub_user_id": "USR-50001120", "wallet_id": "WLT-90007812",
        "resource_type": "wallet_topup", "resource_id": "WTX-70009111", "resource_version": 1, "status": "COMPLETED",
        "data": {"wallet_transaction_id": "WTX-70009111", "olylife_transaction_id": "COM-DB-880031", "amount": "500.00", "currency": "SGD", "wallet_balance": "500.00", "balance_as_of": "2026-09-03T03:05:03Z"},
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Card created", {
        "event_id": "EVT-20260903-000060", "event_type": "card.created", "event_version": "1.0",
        "occurred_at": "2026-09-03T03:20:00Z", "company_id": 3201, "request_id": "card-create-20260903-0060",
        "external_member_id": "OLY-10002345", "vcchub_user_id": "USR-50001120", "wallet_id": "WLT-90007812",
        "resource_type": "card", "resource_id": "CRD-30001002", "resource_version": 1, "status": "DELIVERY_PENDING",
        "data": {"card_id": "CRD-30001002", "card_type": "PHYSICAL", "card_status": "DELIVERY_PENDING", "last4": "6366", "currency": "SGD", "card_balance": "0.00", "delivery_destination": "OLYLIFE_OFFICE"},
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Card top-up successful", {
        "event_id": "EVT-20260903-000075", "event_type": "card.topup.succeeded", "event_version": "1.0",
        "occurred_at": "2026-09-03T03:35:00Z", "company_id": 3201, "request_id": "card-topup-20260903-0075",
        "external_member_id": "OLY-10002345", "vcchub_user_id": "USR-50001120", "wallet_id": "WLT-90007812",
        "resource_type": "card_topup", "resource_id": "CTX-60001234", "resource_version": 1, "status": "COMPLETED",
        "data": {"card_id": "CRD-30001001", "wallet_transaction_id": "WTX-70009200", "card_transaction_id": "CTX-60001234", "amount": "75.00", "currency": "SGD", "wallet_balance": "425.00", "card_balance": "75.00"},
        "sign": "<64-character-sha256>"
    })
    insert_code(doc, anchor, "Card cancelled", {
        "event_id": "EVT-20260903-000099", "event_type": "card.cancelled", "event_version": "1.0",
        "occurred_at": "2026-09-03T04:00:00Z", "company_id": 3201, "request_id": "card-cancel-20260903-0099",
        "external_member_id": "OLY-10002345", "vcchub_user_id": "USR-50001120", "wallet_id": "WLT-90007812",
        "resource_type": "card", "resource_id": "CRD-30001001", "resource_version": 4, "status": "CANCELLED",
        "data": {"card_id": "CRD-30001001", "card_type": "VIRTUAL", "card_status": "CANCELLED", "refunded_amount": "75.00", "currency": "SGD", "wallet_balance": "500.00", "card_type_slot_released": True},
        "sign": "<64-character-sha256>"
    })

    insert_heading(doc, anchor, "4.7.5 OlyLife acknowledgement example", page_break=True)
    insert_code(doc, anchor, "HTTP 200 after durable idempotent acceptance", {
        "code": "0000",
        "message": "Success",
        "result": {"event_id": "EVT-20260903-000099", "received_at": "2026-09-03T04:00:01Z"}
    })
    insert_paragraph(
        doc, anchor,
        "Receiver algorithm: verify HTTPS source controls and signature; validate event_version/schema; check event_id; durably store the event and processing state; apply only when resource_version is newer; commit the local update; then return HTTP 200. A duplicate valid event returns the same successful acknowledgement without reapplying the state change.",
        size=9.6, color=INK, after=10,
    )


def build_v2_2() -> None:
    doc = Document(SOURCE)
    update_control_and_shared_decision(doc)
    expand_interface_catalogue(doc)

    heading_44 = find_paragraph(doc, "4.4 OLY-01")
    heading_45 = find_paragraph(doc, "4.5 VCC-01")
    heading_46 = find_paragraph(doc, "4.6 VCC-WH-01")
    heading_47 = find_paragraph(doc, "4.7 OlyLife internal")
    heading_48 = find_paragraph(doc, "4.8 VCCHUB-internal")
    heading_49 = find_paragraph(doc, "4.9 Design references")

    delete_between(heading_44, heading_45)
    delete_between(heading_45, heading_46)
    delete_between(heading_46, heading_47)

    heading_44.paragraph_format.page_break_before = True
    # Let the following interface sections use the remaining page space after
    # their preceding JSON example. Forced breaks produced mostly blank pages.
    heading_45.paragraph_format.page_break_before = False
    heading_46.paragraph_format.page_break_before = False
    heading_47.paragraph_format.page_break_before = False
    heading_46.text = "4.6 VCC-02 - Wallet/account status (POST /wallet/status)"
    heading_47.text = "4.8 OlyLife internal top-up request and approval workflow"
    heading_48.text = "4.9 VCCHUB-internal/provider integrations"
    heading_49.text = "4.10 Design references (not OlyLife integration scope)"

    build_oly_01(doc, heading_45)
    build_vcc_01(doc, heading_46)
    build_vcc_02(doc, heading_47)
    build_webhooks(doc, heading_47)

    doc.core_properties.title = "OlyLife-VCCHUB Implementation Responsibility Specification - Version 2.2"
    doc.core_properties.subject = "V2 implementation ownership with field-level API responses, examples and signed lifecycle webhook contracts"
    doc.core_properties.author = "Star SaaS Limited"
    doc.core_properties.keywords = "OlyLife, VCCHUB, Version 2.2, API contract, response schema, webhook, wallet, card, top-up, cancellation, SHA-256"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_v2_2()
