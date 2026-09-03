from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from update_responsibility_document_v2_2 import (
    insert_heading,
    insert_paragraph,
    insert_table,
    update_cell,
)
from build_responsibility_document import format_table, set_cant_split


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.2.docx"
OUTPUT = ROOT / "deliverables" / "OlyLife_VCCHUB_Implementation_Responsibility_Specification_V2.3.docx"


def find_paragraph(doc: Document, prefix: str):
    return next(p for p in doc.paragraphs if p.text.strip().startswith(prefix))


def find_table(doc: Document, first_header: str, columns: int | None = None):
    return next(
        table
        for table in doc.tables
        if table.rows
        and (columns is None or len(table.columns) == columns)
        and table.rows[0].cells[0].text.strip() == first_header
    )


def find_row(table, first_cell: str):
    return next(row for row in table.rows[1:] if row.cells[0].text.strip() == first_cell)


def remove_row(table, row) -> None:
    table._tbl.remove(row._tr)


def set_row(table, first_cell: str, values) -> None:
    row = find_row(table, first_cell)
    for index, value in enumerate(values):
        update_cell(row.cells[index], value)


def update_version_and_front_matter(doc: Document) -> None:
    control = doc.tables[0]
    update_cell(control.rows[0].cells[3], "Draft - V2.3 confirmed programme assumptions")
    update_cell(control.rows[2].cells[1], "0.14 (V2.3 confirmed assumptions and constraints)")
    update_cell(control.rows[2].cells[3], "3 September 2026")

    find_paragraph(doc, "PARTNER IMPLEMENTATION SPECIFICATION").text = (
        "PARTNER IMPLEMENTATION SPECIFICATION · VERSION 2.3"
    )
    find_paragraph(doc, "OlyLife–VCCHUB Wallet Activation").text = (
        "OlyLife–VCCHUB Wallet Activation & Card Program (Version 2.3)"
    )
    purpose = find_paragraph(doc, "Purpose")
    purpose.text = (
        "Purpose  Define who must build, expose, consume, operate and test each component of the Version 2 journey, "
        "where an OlyLife member starts in VCCHUB. This V2.3 edition records the confirmed programme assumptions, "
        "release constraints and partner message contracts; V2.2 remains retained as the preceding reference. "
        "Final production configuration and contractual approvals remain subject to joint sign-off."
    )


def update_shared_responsibilities(doc: Document) -> None:
    table = find_table(doc, "Decision area", 3)
    support_row = find_row(table, "7. Support ownership and service levels")
    remove_row(table, support_row)
    uat_row = find_row(table, "8. UAT and go-live acceptance")
    update_cell(uat_row.cells[0], "7. UAT and go-live acceptance")
    format_table(
        table,
        ["Decision area", "VCCHUB proposed default", "OlyLife action / response"],
        [1900, 5000, 2460],
        body_size=8.1,
        first_col_bold=True,
    )
    for row in table.rows:
        set_cant_split(row)


def add_confirmed_assumptions(doc: Document) -> None:
    anchor = find_paragraph(doc, "4. External interface catalogue")
    insert_heading(
        doc,
        anchor,
        "3.5 Confirmed V2 programme assumptions and release constraints",
        level=2,
        page_break=True,
    )
    insert_paragraph(
        doc,
        anchor,
        "The following positions are treated as the agreed V2 delivery baseline. A later change request may alter a future release, but it does not silently change existing wallets, cards, balances or ledger history.",
        italic=True,
    )
    insert_table(
        doc,
        anchor,
        ["#", "Confirmed V2 position", "Implementation consequence / boundary"],
        [
            [
                "1",
                "Card entitlement is two cards per wallet account: one physical card and one virtual card.",
                "VCCHUB stores the card policy applied when the wallet is created. A later entitlement amendment applies only to newly created wallet accounts. Migrating existing wallet users requires a separately planned VCCHUB version upgrade/data migration and is excluded from the current go-live scope.",
            ],
            [
                "2",
                "OlyLife requires programme/card fees to be charged directly against the created card under the PhotonPay arrangement.",
                "OlyLife agrees the fee schedule and charging treatment directly with PhotonPay. VCCHUB will not configure or deduct these fees from the member Wallet balance in V2. Issuer transaction fees and their card-ledger treatment remain governed by PhotonPay.",
            ],
            [
                "3",
                "Wallet top-up cannot be initiated from VCCHUB. Every wallet top-up starts in OlyLife and follows OlyLife's internal approval window.",
                "Balances remain unchanged while approval is pending. After approval, OlyLife debits commission once and calls VCCHUB POST /wallet/topup; VCCHUB credits the Wallet exactly once and returns or publishes the final result.",
            ],
            [
                "4",
                "VCCHUB configures one default physical-card BIN and one default virtual-card BIN for the wallet programme.",
                "Card creation automatically uses the default BIN for the selected card type. V2 has no multi-BIN configuration for wallet users and no member BIN-selection screen. Multi-BIN support requires a later version if requested.",
            ],
            [
                "5",
                "A successfully cancelled card releases that card-type entitlement.",
                "Cancellation permanently marks only the selected card as Cancelled, returns any eligible card balance to the Wallet exactly once and restores the physical or virtual slot. The member may create a replacement card of that type provided no other active card of that type exists.",
            ],
            [
                "6",
                "After successful KYC, VCCHUB creates the cardholder and a zero-balance wallet with no card.",
                "Wallet and cardholder creation are VCCHUB-internal operations. A member must first receive an approved OlyLife-funded Wallet top-up before creating a card.",
            ],
            [
                "7",
                "V2 onboarding starts in VCCHUB. OlyLife validates membership by email and returns Member ID, First Name, Last Name and the matched email for an eligible member.",
                "VCCHUB carries the returned member values into registration, uses the matched email as username and collects mobile, password, consent and full registered address. V2 does not use an invitation email or OlyLife-to-VCCHUB SSO launch.",
            ],
            [
                "8",
                "OlyLife owns the Sumsub contract/account and provides controlled sandbox and production access to VCCHUB.",
                "VCCHUB implements and operates the integration. The user is launched into Sumsub's eligible document flow without a VCCHUB country-preselection step; Sumsub handles supported country/document selection and verification.",
            ],
            [
                "9",
                "The member's registered address is collected during registration and displayed read-only during card creation. Physical cards are delivered to the configured OlyLife office address.",
                "VCCHUB does not collect a physical-card recipient address from the member in V2. OlyLife supplies and approves the office address and change-control details used for fulfilment.",
            ],
            [
                "10",
                "VCCHUB direct sign-in requires 2FA before Wallet access. Wallet withdrawal back to OlyLife commission is not available in V2.",
                "The production login/MFA controls belong to VCCHUB. Funds credited to the Wallet are used within the card programme; any future wallet-to-commission withdrawal or SSO flow requires separate scope and approval.",
            ],
        ],
        [620, 4300, 4440],
        body_size=7.9,
        first_col_bold=True,
    )


def align_money_card_and_fee_rules(doc: Document) -> None:
    money = find_table(doc, "Item", 4)
    set_row(
        money,
        "Card eligibility",
        [
            "Card eligibility",
            "N/A",
            "Require approved KYC, an active cardholder and a funded Wallet; allow one active virtual and one active physical card under the policy version assigned at wallet creation",
            "VCCHUB",
        ],
    )
    set_row(
        money,
        "Card top-up",
        [
            "Card top-up",
            "N/A",
            "Check Wallet balance, debit Wallet and credit the selected card atomically. No VCCHUB-configured programme/card fee is deducted from Wallet in V2",
            "VCCHUB",
        ],
    )
    format_table(
        money,
        ["Item", "OlyLife implementation", "VCCHUB implementation", "Accountable"],
        [1320, 2880, 3700, 1460],
        body_size=8.0,
        first_col_bold=True,
    )

    coverage = find_table(doc, "Journey capability", 4)
    set_row(
        coverage,
        "Wallet-to-card top-up",
        [
            "Wallet-to-card top-up",
            "Existing wallet/card functions",
            "REUSE — VCCHUB INTERNAL",
            "Post internally after Wallet-balance checks. No VCCHUB programme/card fee is deducted from Wallet; OlyLife handles the agreed fee arrangement directly with PhotonPay.",
        ],
    )
    format_table(
        coverage,
        ["Journey capability", "Published Issuing coverage", "Classification", "Implementation decision"],
        [1900, 2400, 1700, 3360],
        body_size=7.9,
        first_col_bold=True,
    )

    inventory = find_table(doc, "Domain", 4)
    set_row(
        inventory,
        "Wallet",
        [
            "Wallet",
            "Create wallet",
            "Existing VCCHUB internal API",
            "Reuse after KYC approval to create a zero-balance member Wallet with no card. The wallet receives the current card-policy version; later policy amendments do not update existing wallets automatically.",
        ],
    )
    # The table has several Wallet and Card rows, so target the specific second cell.
    for row in inventory.rows[1:]:
        name = row.cells[1].text.strip()
        if name == "Card-fee debit":
            values = [
                "Wallet",
                "Programme/card fee charging",
                "Out of VCCHUB V2 configuration scope",
                "OlyLife agrees fee charging directly with PhotonPay. No VCCHUB programme/card fee is configured or deducted from member Wallet balance in V2.",
            ]
        elif name == "Query card BIN":
            values = [
                "Card",
                "Query/configure card BIN",
                "POST /api/card/bin_query",
                "Reuse internally to configure one default virtual BIN and one default physical BIN. Card creation auto-selects the matching default; no member selection or multi-BIN wallet configuration in V2.",
            ]
        elif name == "Create a card":
            values = [
                "Card",
                "Create a card",
                "POST /api/card/create",
                "Reuse internally after funded-Wallet checks. Permit one active card of each type under the wallet's assigned policy version and auto-use the configured default BIN. List cards separately by card ID. Physical fulfilment uses the OlyLife office address. Cancellation releases only the cancelled type slot.",
            ]
        elif name == "Recharge card":
            values = [
                "Card",
                "Recharge card",
                "POST /api/card/charge",
                "Reuse internally for Wallet-to-card top-up; first check Wallet balance, then pair request_id with VCCHUB idempotency, Wallet debit and selected-card credit.",
            ]
        elif name == "Update shared card limit":
            values = [
                "Card",
                "Update card entitlement",
                "Not exposed in the V2 base journey",
                "V2 fixes each new Wallet at one virtual plus one physical card. A future policy change applies prospectively; updating existing wallets requires a planned VCCHUB upgrade/migration.",
            ]
        else:
            continue
        for index, value in enumerate(values):
            update_cell(row.cells[index], value)
    format_table(
        inventory,
        ["Domain", "Published API", "Published path / status", "Mapping to OlyLife journey"],
        [1050, 1850, 2550, 3910],
        body_size=7.7,
        first_col_bold=True,
    )

    for table in doc.tables:
        if not table.rows or table.rows[0].cells[0].text.strip() != "Element":
            continue
        for row in table.rows[1:]:
            if row.cells[0].text.strip() == "result.credited_amount":
                update_cell(
                    row.cells[3],
                    "Approved principal credited to the Wallet. No VCCHUB programme/card fee is deducted from this amount in V2.",
                )
                for paragraph in row.cells[3].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7.7)
                break


def align_controls_acceptance_and_actions(doc: Document) -> None:
    controls = find_table(doc, "Control area", 2)
    set_row(
        controls,
        "Idempotency",
        [
            "Idempotency",
            "Required for member validation, account-ready/status processing, top-up approval/debit, Wallet credit, reversal, card issue, card top-up and card cancellation.",
        ],
    )
    format_table(
        controls,
        ["Control area", "Minimum production requirement"],
        [1900, 7460],
        body_size=8.4,
        first_col_bold=True,
    )

    failures = find_table(doc, "Scenario", 2)
    set_row(
        failures,
        "Card type, issue or top-up control",
        [
            "Card type, issue or top-up control",
            "Reject a second active card of the same type and show one row with independent actions per card. Auto-use the configured default BIN. If issue/top-up fails, post neither final Wallet debit nor card credit, or reverse reservations atomically.",
        ],
    )
    format_table(
        failures,
        ["Scenario", "Required behaviour"],
        [2500, 6860],
        body_size=8.4,
        first_col_bold=True,
    )

    delivery = find_table(doc, "Phase", 3)
    set_row(
        delivery,
        "6. Card services",
        [
            "6. Card services",
            "VCCHUB",
            "Funding gate; one active virtual plus one active physical card; prospective card-policy versioning; default BIN per type; separate card rows/actions; read-only identity/address; OlyLife-office delivery; Wallet-to-card top-up; authenticated cancellation with Wallet refund and type-slot release; no VCCHUB Wallet-fee deduction.",
        ],
    )
    format_table(
        delivery,
        ["Phase", "Lead", "Exit criteria"],
        [1800, 1900, 5660],
        body_size=8.2,
        first_col_bold=True,
    )

    decisions = find_table(doc, "Decision area", 2)
    set_row(
        decisions,
        "Card servicing",
        [
            "Card servicing",
            "Confirmed: one active virtual and one active physical card per Wallet policy version; one default BIN per type; separate card rows/actions; cancellation with fresh 2FA, Wallet refund, permanent Cancelled status and release of only that type slot. Policy changes are prospective; existing-wallet migration is outside V2. Confirm processor edge cases only.",
        ],
    )
    set_row(
        decisions,
        "Physical cards",
        [
            "Physical cards",
            "Actual OlyLife office address, authorised change process, courier, tracking, failed delivery, replacement and office receiving owner. OlyLife agrees card/programme fee charging directly with PhotonPay; VCCHUB does not deduct those fees from Wallet in V2.",
        ],
    )
    format_table(
        decisions,
        ["Decision area", "Items to agree"],
        [2000, 7360],
        body_size=8.2,
        first_col_bold=True,
    )

    # Keep the action heading, readiness callout, and first action rows together.
    find_paragraph(doc, "9.1 Immediate next actions").paragraph_format.page_break_before = True

    owners = find_table(doc, "Owner", 2)
    set_row(
        owners,
        "VCCHUB",
        [
            "VCCHUB",
            "Publish OLY-01 consumer requirements and VCC-01/02/VCC-WH schemas; implement V2 onboarding, member carry-over, registration address, Sumsub, direct 2FA, zero-balance Wallet creation, read-only card data, prospective card-policy versioning, one default BIN per type, separate per-card actions, OlyLife-office physical fulfilment and authenticated cancellation with Wallet refund/type-slot release.",
        ],
    )
    set_row(
        owners,
        "OlyLife",
        [
            "OlyLife",
            "Publish OLY-01 with Member ID/First Name/Last Name; provision Sumsub access; implement top-up request/review/debit/reversal; provide the actual office address; and agree all programme/card fee charging directly with PhotonPay.",
        ],
    )
    set_row(
        owners,
        "Joint",
        [
            "Joint",
            "Run an interface workshop to close identifiers, authentication, status/error models, account mapping, Wallet-credit idempotency and reversal/reconciliation. Confirm the initial card-policy version and default physical/virtual BINs before go-live.",
        ],
    )
    format_table(
        owners,
        ["Owner", "Next action"],
        [1800, 7560],
        body_size=8.4,
        first_col_bold=True,
    )

    confirmed = find_paragraph(doc, "Confirmed programme decisions")
    confirmed.text = (
        "Confirmed programme decisions  PhotonPay remains the issuer. V2 starts in VCCHUB, validates OlyLife membership before registration, collects the member's full address, creates a zero-balance Wallet with no card after KYC, and funds Wallets only through OlyLife's approved top-up process. Each Wallet is assigned a policy allowing one active virtual and one active physical card, with one default BIN per type and no user BIN selection. Policy changes apply prospectively unless a separate existing-wallet migration is approved. Physical cards are sent to the configured OlyLife office. Card cancellation releases the cancelled type slot. OlyLife handles programme/card fee charging directly with PhotonPay; VCCHUB does not deduct such fees from Wallet balance in V2."
    )

    # Remove the former support-ownership checkpoint from the closing action note.
    handoff = find_paragraph(doc, "Handoff checkpoint")
    handoff.text = handoff.text.replace(
        ", physical fulfilment controls and cancellation authentication/audit details",
        ", physical fulfilment controls, default BINs, prospective card-policy versioning and cancellation authentication/audit details",
    )


def build_v2_3() -> None:
    doc = Document(SOURCE)
    update_version_and_front_matter(doc)
    update_shared_responsibilities(doc)
    add_confirmed_assumptions(doc)
    align_money_card_and_fee_rules(doc)
    align_controls_acceptance_and_actions(doc)

    doc.core_properties.title = "OlyLife-VCCHUB Implementation Responsibility Specification - Version 2.3"
    doc.core_properties.subject = "V2 implementation ownership, partner contracts, confirmed programme assumptions and release constraints"
    doc.core_properties.author = "Star SaaS Limited"
    doc.core_properties.keywords = "OlyLife, VCCHUB, Version 2.3, assumptions, card policy, default BIN, PhotonPay fees, Wallet top-up, cancellation"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_v2_3()
